# -*- coding: utf-8 -*-
"""生成《第一天报告.docx》—— 网络安全编程课程设计"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, font_name_cn, font_name_en, size, bold=False, color=None):
    """统一设置 run 的中英文字体"""
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = font_name_en
    # 设置中文字体
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_cn)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_custom(doc, text, level, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """添加黑体标题"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(12 if level == 1 else 8)
    pf.space_after = Pt(6 if level == 1 else 4)

    if level == 1:
        size = 22
    elif level == 2:
        size = 16
    else:
        size = 14

    run = p.add_run(text)
    set_run_font(run, '黑体', 'SimHei', size, bold=True)

    # 设置 outline level
    pPr = p._element.get_or_add_pPr()
    outline = parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="{level - 1}"/>')
    pPr.append(outline)

    return p


def add_body_text(doc, text, indent_first_line=True, bold=False, space_after=6):
    """添加宋体正文段落"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(20)
    if indent_first_line:
        pf.first_line_indent = Cm(0.74)  # 约两个字符

    run = p.add_run(text)
    set_run_font(run, '宋体', 'SimSun', 12, bold=bold)
    return p


def add_body_mixed(doc, segments, indent_first_line=True):
    """添加混合格式正文段落，segments 为 (text, bold, font_cn, font_en, size) 的列表"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = Pt(20)
    if indent_first_line:
        pf.first_line_indent = Cm(0.74)

    for seg in segments:
        if len(seg) == 2:
            text, bold = seg
            run = p.add_run(text)
            set_run_font(run, '宋体', 'SimSun', 12, bold=bold)
        else:
            text, bold, font_cn, font_en, size = seg
            run = p.add_run(text)
            set_run_font(run, font_cn, font_en, size, bold=bold)
    return p


def add_code_block(doc, text):
    """添加代码块（灰底 Consolas 9pt）"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.left_indent = Cm(0.5)

    # 灰色底纹
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F2F2F2"/>')
    pPr.append(shd)

    run = p.add_run(text)
    set_run_font(run, 'Consolas', 'Consolas', 9)
    return p


def create_table(doc, headers, rows, col_widths=None):
    """创建带边框的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表格边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, '黑体', 'SimHei', 11, bold=True)
        set_cell_shading(cell, 'D9E2F3')

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            set_run_font(run, '宋体', 'SimSun', 11)

    # 设置列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def main():
    doc = Document()

    # ========== 页面设置 A4 ==========
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ========== 封面标题 ==========
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p_title.paragraph_format
    pf.space_before = Pt(80)
    pf.space_after = Pt(12)
    run = p_title.add_run('网络安全编程课程设计')
    set_run_font(run, '黑体', 'SimHei', 28, bold=True)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf2 = p_sub.paragraph_format
    pf2.space_after = Pt(40)
    run2 = p_sub.add_run('第一天报告')
    set_run_font(run2, '黑体', 'SimHei', 22, bold=True)

    # 分隔线
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf3 = p_line.paragraph_format
    pf3.space_after = Pt(60)
    run3 = p_line.add_run('2026 年 6 月')
    set_run_font(run3, '宋体', 'SimSun', 14)

    # ========== 一、课程设计概述 ==========
    add_heading_custom(doc, '一、课程设计概述', 1)

    # 1.1
    add_heading_custom(doc, '1.1 课程设计题目', 2)
    add_body_text(doc, '智能自适应 C2（Command & Control）框架的设计与实现')

    # 1.2
    add_heading_custom(doc, '1.2 课程设计目的', 2)
    add_body_text(
        doc,
        '通过构建一个完整的 C2 框架，深入理解网络安全攻防对抗中的关键技术，'
        '包括加密通信、环境感知、沙箱检测、隐蔽通信、插件化架构等，'
        '提升网络安全编程能力和安全对抗思维。'
    )

    # 1.3
    add_heading_custom(doc, '1.3 课程设计任务', 2)
    tasks = [
        '实现沙箱/虚拟机/调试器检测模块，通过加权评分算法评估运行环境风险',
        '实现 WMI 环境感知模块，自动识别系统信息和安全软件',
        '实现 RSA-2048 + AES-256-GCM 混合加密通信体系',
        '实现 TCP/UDP/Outlook/PowerShell/Certutil 多通路隐蔽通信',
        '实现插件加密打包与动态加载机制',
        '实现分片传输模块，规避流量监控',
        '实现内网横向探测和窗口活动监控等扩展功能',
        '实现控制端 CLI 交互界面',
    ]
    for i, t in enumerate(tasks, 1):
        add_body_text(doc, f'{i}. {t}', indent_first_line=True)

    # 1.4 开发环境
    add_heading_custom(doc, '1.4 开发环境', 2)
    create_table(
        doc,
        headers=['类别', '说明'],
        rows=[
            ['开发语言', 'Python 3.10+'],
            ['控制端系统', 'Kali Linux'],
            ['受控端系统', 'Windows 10/11'],
            ['加密库', 'PyCryptodome (RSA/AES)'],
            ['系统接口', 'ctypes (Windows API)、WMI'],
            ['通信协议', '自定义 JSON over TCP/UDP'],
            ['测试工具', 'Wireshark、Process Monitor'],
        ],
        col_widths=[4, 10],
    )

    # 1.5 项目进度安排
    add_heading_custom(doc, '1.5 项目进度安排', 2)
    create_table(
        doc,
        headers=['阶段', '时间', '任务内容'],
        rows=[
            ['分析设计准备', '第1天上午', '查找资料，确定选题，需求分析'],
            ['问题分析', '第1天下午', '系统架构设计，模块划分'],
            ['算法设计', '第2天上午', '加密算法、检测算法、分片算法设计'],
            ['算法分析', '第2天下午', '算法复杂性分析与优化'],
            ['编程实现', '第3天', '核心模块编码实现'],
            ['测试调试', '第4天', '测试用例设计、功能测试、运行记录'],
        ],
        col_widths=[3.5, 3, 7.5],
    )

    # ========== 二、问题分析 ==========
    add_heading_custom(doc, '二、问题分析', 1)

    # 2.1
    add_heading_custom(doc, '2.1 项目背景', 2)
    add_body_text(
        doc,
        '随着网络安全攻防对抗的不断升级，传统的安全检测手段面临着越来越大的挑战。'
        '攻击者利用各种技术手段来规避安全软件的检测，包括环境感知、通信隐蔽、数据加密等。'
        '本课程设计旨在通过构建一个智能自适应 C2 框架，深入理解攻防对抗中的关键技术原理，'
        '提升网络安全编程能力。'
    )

    # 2.2
    add_heading_custom(doc, '2.2 需求分析', 2)
    add_body_text(
        doc,
        '本项目需要实现一个完整的 C2 框架，包含以下核心功能需求：'
    )

    # 功能需求
    add_body_mixed(doc, [('功能需求：', True)], indent_first_line=True)

    func_reqs = [
        (
            '沙箱/虚拟机/调试器检测：',
            '在执行恶意行为前判断是否处于分析环境中，避免被安全研究人员逆向分析。'
            '通过调用 Windows API（IsDebuggerPresent、CheckRemoteDebuggerPresent）检测调试器，'
            '通过注册表键值、进程名、MAC 地址前缀检测虚拟机，'
            '通过鼠标/键盘活动、系统运行时间检测沙箱环境。'
        ),
        (
            '环境感知与自适应：',
            '自动识别目标系统信息（主机名、操作系统、架构、内存、磁盘）、'
            '安全软件（火绒、360、Defender 等 8 种）、常用应用（Outlook、微信、Chrome 等 11 种），'
            '为后续通信策略提供决策依据。'
        ),
        (
            '多通路隐蔽通信：',
            '支持 TCP 加密套接字、UDP 分片通信、Outlook COM 邮件外传、'
            'PowerShell 白名单通路、certutil 白名单通路共 6 种通信方式，按优先级自动选择和降级。'
        ),
        (
            '混合加密体系：',
            '采用 RSA-2048 非对称加密交换密钥 + AES-256-GCM 对称加密传输数据的混合方案，'
            '保障通信机密性和完整性。'
        ),
        (
            '插件化模块管理：',
            '支持插件的加密打包（.py → .pyc → AES 加密 → .enc）、'
            '动态加载（内存解密执行）和远程执行。'
        ),
        (
            '分片传输：',
            '大数据包按 4KB 分片，逐片加密发送，设置发送间隔规避流量监控。'
        ),
        (
            '内网横向探测：',
            '自动发现同网段存活主机（TCP 445 + ICMP 双探针）。'
        ),
        (
            '窗口活动监控：',
            '通过 Windows API 采集前台窗口切换记录。'
        ),
    ]

    for i, (title, desc) in enumerate(func_reqs, 1):
        add_body_mixed(
            doc,
            [
                (f'{i}. ', False, '宋体', 'SimSun', 12),
                (title, True, '宋体', 'SimSun', 12),
                (desc, False, '宋体', 'SimSun', 12),
            ],
            indent_first_line=True,
        )

    # 非功能需求
    add_body_mixed(doc, [('非功能需求：', True)], indent_first_line=True)
    non_func_reqs = [
        '安全性：所有通信加密，插件代码加密存储，防重放攻击',
        '隐蔽性：无窗口进程、白名单工具利用、分片传输规避检测',
        '自适应性：根据环境自动选择通信策略',
        '可扩展性：插件化架构，支持动态添加新功能模块',
    ]
    for i, t in enumerate(non_func_reqs, 1):
        add_body_text(doc, f'{i}. {t}', indent_first_line=True)

    # 2.3
    add_heading_custom(doc, '2.3 系统架构分析', 2)
    add_body_text(
        doc,
        '系统采用 C/S 架构，分为控制端（Controller，运行在 Kali Linux）和受控端（Agent，运行在 Windows），'
        '通过共享模块（shared）定义通信协议和加密引擎。'
    )
    add_body_text(doc, '核心模块划分：')

    modules = [
        ('shared/protocol.py', '定义消息类型（cmd/result/heartbeat/register/key_exchange）和指令类型（exec/download/upload/sysinfo/modules/load_mod/run_mod/netprobe/winmon）'),
        ('shared/crypto.py', 'RSA-2048 密钥管理 + AES-256-GCM 加解密引擎'),
        ('controller/main.py', 'CLI 交互界面，会话管理，指令下发'),
        ('controller/comm_server.py', 'TCP 监听，连接管理，密钥交换'),
        ('agent/main.py', 'Agent 7 阶段生命周期管理'),
        ('agent/sandbox_detect.py', '11 项检测指标，加权评分算法'),
        ('agent/env_perception.py', 'WMI 环境感知，安全软件/白名单软件识别'),
        ('agent/comm_channels.py', '6 种通信通道，策略模式实现'),
        ('agent/module_loader.py', '加密模块动态加载，sys.meta_path 钩子'),
        ('agent/frag_transfer.py', '数据分片与重组'),
        ('agent/plugins/', 'recon（网络侦察）、netprobe（内网探测）、winmon（窗口监控）'),
    ]

    for mod, desc in modules:
        add_body_mixed(
            doc,
            [
                (mod, True, 'Consolas', 'Consolas', 10),
                ('：' + desc, False, '宋体', 'SimSun', 12),
            ],
            indent_first_line=True,
        )

    # 2.4
    add_heading_custom(doc, '2.4 关键问题识别', 2)

    issues = [
        ('如何在不引入第三方依赖的情况下检测沙箱/VM/调试器？',
         '通过 ctypes 调用 Windows API'),
        ('如何保证通信不被网络监控发现？',
         '多通路自适应 + 分片传输 + 白名单工具利用'),
        ('如何安全地交换加密密钥？',
         'RSA-2048 公钥交换 + AES-256 会话密钥'),
        ('如何防止插件被逆向分析？',
         '编译加密 + 内存解密执行 + 独立密钥'),
        ('如何处理中文 Windows 命令输出的编码问题？',
         'GBK/UTF-8/GB2312 自动检测解码'),
    ]

    for i, (q, a) in enumerate(issues, 1):
        add_body_mixed(
            doc,
            [
                (f'{i}. ', False, '宋体', 'SimSun', 12),
                (q, False, '宋体', 'SimSun', 12),
                (f' → {a}', False, '宋体', 'SimSun', 12),
            ],
            indent_first_line=True,
        )

    # ========== 保存 ==========
    output_path = r'E:\课设2\第一天报告.docx'
    doc.save(output_path)
    print(f'文档已生成: {output_path}')


if __name__ == '__main__':
    main()
