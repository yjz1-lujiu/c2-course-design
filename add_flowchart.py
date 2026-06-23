# -*- coding: utf-8 -*-
"""生成 Agent 主程序流程图并插入到课程设计报告中"""
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
import os

FLOWCHART_IMG = r"E:\课设2\flowchart.png"
REPORT_PATH   = r"E:\课设2\课程设计报告.docx"

# ═══════════════════════════════════════════════════════════════════════
# Part 1: 生成流程图 PNG
# ═══════════════════════════════════════════════════════════════════════

def draw_flowchart():
    # 中文字体
    plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'STSong']
    plt.rcParams['axes.unicode_minus'] = False

    fig, ax = plt.subplots(1, 1, figsize=(16, 24))
    ax.set_xlim(-1, 11)
    ax.set_ylim(-1, 23)
    ax.set_aspect('equal')
    ax.axis('off')
    fig.patch.set_facecolor('white')

    # ── 样式定义 ──
    W, H = 3.0, 0.65           # 普通节点宽高
    DW, DH = 3.2, 1.2          # 菱形宽高
    SP = 1.3                    # 纵向间距
    CX = 5.0                    # 主流程中心X

    def draw_box(x, y, w, h, text, fc='#EBF5FB', ec='#2980B9', lw=2.0, fs=9.5):
        """圆角矩形节点"""
        box = mpatches.FancyBboxPatch(
            (x - w/2, y - h/2), w, h,
            boxstyle='round,pad=0.08', facecolor=fc, edgecolor=ec, linewidth=lw)
        ax.add_patch(box)
        ax.text(x, y, text, ha='center', va='center', fontsize=fs, fontweight='bold')

    def draw_diamond(x, y, w, h, text, fc='#FFF9C4', ec='#F39C12', lw=2.2, fs=9):
        """菱形决策节点"""
        diamond = plt.Polygon(
            [(x, y + h/2), (x + w/2, y), (x, y - h/2), (x - w/2, y)],
            closed=True, fc=fc, ec=ec, lw=lw, zorder=2)
        ax.add_patch(diamond)
        ax.text(x, y, text, ha='center', va='center', fontsize=fs, fontweight='bold', zorder=3)

    def arrow(x1, y1, x2, y2, color='#555', lw=2.0, ls='-', label='', lx=0, ly=0, fs=9):
        """带箭头的连线"""
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color=color, lw=lw, ls=ls), zorder=1)
        if label:
            mx, my = (x1+x2)/2 + lx, (y1+y2)/2 + ly
            ax.text(mx, my, label, ha='center', va='center', fontsize=fs,
                    fontweight='bold', color=color,
                    bbox=dict(fc='white', ec='none', alpha=0.9, pad=2), zorder=4)

    def hline(x1, x2, y, color='#555', lw=2.0, ls='-'):
        """水平线"""
        ax.plot([x1, x2], [y, y], color=color, lw=lw, ls=ls, zorder=1)

    def vline(x, y1, y2, color='#555', lw=2.0, ls='-'):
        """垂直线"""
        ax.plot([x, x], [y1, y2], color=color, lw=lw, ls=ls, zorder=1)

    # ── Y 坐标（从上到下） ──
    y0 = 22.0   # 启动
    y1 = y0 - SP           # 沙箱检测
    y2 = y1 - SP*1.05      # 风险决策
    y3 = y2 - SP*1.15      # 环境感知
    y4 = y3 - SP           # RSA
    y5 = y4 - SP           # 连接决策
    y6 = y5 - SP           # 通道尝试
    y7 = y6 - SP           # 连接成功?
    y8 = y7 - SP*1.15      # 密钥交换
    y9 = y8 - SP           # 注册
    y10 = y9 - SP          # 命令循环
    y11 = y10 - SP*1.3     # 接收指令
    y12 = y11 - SP         # 执行指令
    y13 = y12 - SP         # 发送结果
    y14 = y13 - SP         # 心跳/断线

    # ══════════════════════════════════════════════════════════════════
    # 主流程节点
    # ══════════════════════════════════════════════════════════════════

    # 启动
    draw_box(CX, y0, W*0.7, H*0.8, 'Agent 启动', fc='#2ECC71', ec='#27AE60', fs=11)

    # Phase 1: 沙箱检测
    draw_box(CX, y1, W, H, 'Phase 1: 沙箱/VM/调试器检测\n(11项加权评分, max_score=19)',
             fc='#FDEBD0', ec='#E67E22')

    # 决策1: 风险判定
    draw_diamond(CX, y2, DW, DH, 'risk_ratio\n>= 0.5 ?')

    # Phase 2: 环境感知
    draw_box(CX, y3, W, H,
             'Phase 2: WMI 环境感知\n(系统信息 + 安全软件 + 白名单应用)',
             fc='#FDEBD0', ec='#E67E22')

    # Phase 3: RSA
    draw_box(CX, y4, W, H*0.9, 'Phase 3: 生成 RSA-2048 密钥对',
             fc='#FDEBD0', ec='#E67E22')

    # 决策2: 通信连接
    draw_diamond(CX, y5, DW*0.9, DH*0.95, '尝试次数\n< 10 ?')

    # Phase 4: 通道选择
    draw_box(CX, y6, W, H*1.05,
             'Phase 4: ChannelManager\n.select_channels() 选择可用通道',
             fc='#FDEBD0', ec='#E67E22')

    # 决策3: 连接成功?
    draw_diamond(CX, y7, DW*0.85, DH*0.95, '通道\n连接成功 ?')

    # Phase 5: 密钥交换
    draw_box(CX, y8, W, H*1.1,
             'Phase 5: RSA 密钥交换\n(接收Controller公钥 -> 发送Agent公钥\n-> 生成AES会话密钥)',
             fc='#FDEBD0', ec='#E67E22')

    # Phase 6: 注册
    draw_box(CX, y9, W, H, 'Phase 6: 发送注册信息\n(系统信息+安全软件列表, AES-256-GCM)',
             fc='#FDEBD0', ec='#E67E22')

    # Phase 7: 命令循环
    draw_box(CX, y10, W, H*0.9, 'Phase 7: 指令执行循环',
             fc='#D5F5E3', ec='#27AE60', lw=2.5, fs=10)

    # 循环内部
    draw_box(CX, y11, W, H, '接收指令 (解密+解析)\n超时检测 + 心跳保活',
             fc='#EBF5FB', ec='#3498DB')
    draw_box(CX, y12, W, H, '执行指令\n(ThreadPoolExecutor, max_workers=4)',
             fc='#EBF5FB', ec='#3498DB')
    draw_box(CX, y13, W, H, '发送结果\n(大数据自动分片, >8KB分片传输)',
             fc='#EBF5FB', ec='#3498DB')

    # 决策4: 心跳/断线
    draw_diamond(CX, y14, DW*0.85, DH, '连接\n断开 ?')

    # 伪退出
    draw_box(CX - 4.0, y2, W*0.75, H*0.8, '伪装退出\n(模拟正常行为)',
             fc='#FADBD8', ec='#E74C3C')

    # 超时退出
    draw_box(CX + 4.0, y5, W*0.65, H*0.8, '退出',
             fc='#FADBD8', ec='#E74C3C')

    # 连接失败退出
    draw_box(CX + 4.0, y7, W*0.65, H*0.8, '退出',
             fc='#FADBD8', ec='#E74C3C')

    # ══════════════════════════════════════════════════════════════════
    # 连线
    # ══════════════════════════════════════════════════════════════════

    # 主流程垂直连线
    arrow(CX, y0 - H*0.4,  CX, y1 + H*0.5)
    arrow(CX, y1 - H*0.5,  CX, y2 + DH/2)
    arrow(CX, y2 - DH/2,   CX, y3 + H*0.5,   label='  < 0.5  继续  ', lx=1.2, fs=8.5)
    arrow(CX, y3 - H*0.5,  CX, y4 + H*0.45)
    arrow(CX, y4 - H*0.45, CX, y5 + DH*0.475)
    arrow(CX, y5 - DH*0.475, CX, y6 + H*0.525)
    arrow(CX, y6 - H*0.525,  CX, y7 + DH*0.475)
    arrow(CX, y7 - DH*0.475, CX, y8 + H*0.55, label='  连接成功  ', lx=1.2, fs=8.5)
    arrow(CX, y8 - H*0.55,  CX, y9 + H*0.5)
    arrow(CX, y9 - H*0.5,  CX, y10 + H*0.45)
    arrow(CX, y10 - H*0.45, CX, y11 + H*0.5)
    arrow(CX, y11 - H*0.5,  CX, y12 + H*0.5)
    arrow(CX, y12 - H*0.5,  CX, y13 + H*0.5)
    arrow(CX, y13 - H*0.5,  CX, y14 + DH/2)
    arrow(CX, y14 - DH/2,   CX, y11 + H*0.5, color='#27AE60', lw=2.5, ls='-',
          label='  正常/保持连接  ', lx=-1.6, ly=0.1, fs=8.5)

    # 退出分支
    # 沙箱检测 -> 伪装退出
    vline(CX, y2 + DH/2, y1 - H*0.5, color='white', lw=0)
    hline(CX, CX - 4.0, y2, lw=2.0)
    arrow(CX - DW/2, y2, CX - 4.0 + W*0.375, y2, color='#E74C3C',
          label='>= 0.5', lx=0, ly=0.3, fs=8.5)

    # 重试超时 -> 退出
    hline(CX + DW*0.45, CX + 4.0 - W*0.325, y5)
    arrow(CX + DW*0.45, y5, CX + 4.0 - W*0.325, y5, color='#E74C3C',
          label='>= 10次', lx=0, ly=0.3, fs=8.5)

    # 连接失败 -> 退出
    hline(CX + DW*0.425, CX + 4.0 - W*0.325, y7)
    arrow(CX + DW*0.425, y7, CX + 4.0 - W*0.325, y7, color='#E74C3C',
          label='全部失败', lx=0, ly=0.3, fs=8.5)

    # 断线 -> 通道切换 (回到通道选择)
    arrow(CX + DW*0.425, y14, CX + DW*0.425, y6 - H*0.525, color='#E67E22',
          lw=2.0, ls='--')
    hline(CX + DW*0.425, CX + W/2, y14)
    vline(CX + DW*0.425, y14, y6 - H*0.525, color='#E67E22', lw=2.0, ls='--')
    arrow(CX + DW*0.425, y6 - H*0.525, CX + W/2, y6 - H*0.525, color='#E67E22',
          lw=2.0, ls='--', label='断线重连\n排除失败通道', lx=1.5, fs=8)

    # 重试循环 (连接失败 -> 回到尝试次数决策)
    hline(CX, CX - 4.5, y7)
    vline(CX - 4.5, y7, y5)
    hline(CX - 4.5, CX - DW*0.45, y5)
    arrow(CX - 4.5, y7, CX - 4.5, y5 + DH*0.475, color='#E67E22', lw=2.0, ls='--')
    arrow(CX - 4.5, y5 + DH*0.475, CX - DW*0.45, y5, color='#E67E22', lw=2.0, ls='--',
          label='连接失败\n等待5秒重试', lx=-1.5, fs=8)

    # ══════════════════════════════════════════════════════════════════
    # 通道优先级侧边栏
    # ══════════════════════════════════════════════════════════════════
    sx, sy = 9.2, y6 + 0.5
    sw, sh_total = 3.0, 4.8
    box = mpatches.FancyBboxPatch(
        (sx - sw/2, sy - sh_total/2), sw, sh_total,
        boxstyle='round,pad=0.1', fc='#FFFDE7', ec='#BDBDBD', lw=1.5, ls='--')
    ax.add_patch(box)

    ax.text(sx, sy + sh_total/2 - 0.35, '通道选择优先级', ha='center', va='center',
            fontsize=10, fontweight='bold', color='#F39C12')

    channels = [
        ('1. TCP 加密套接字',        '#27AE60', '基础通路, 始终可用'),
        ('2. Outlook COM 邮件',      '#8E44AD', '需检测到Outlook'),
        ('3. PowerShell 白名单',     '#2980B9', '系统自带, 低拦截率'),
        ('4. UDP 通道',              '#E67E22', '备选方案'),
        ('5. Certutil 白名单',       '#C0392B', '最后手段'),
    ]
    for i, (ch, color, note) in enumerate(channels):
        cy = sy + sh_total/2 - 0.9 - i * 0.75
        ax.plot(sx - sw/2 + 0.25, cy, 'o', color=color, ms=8, zorder=3)
        ax.text(sx - sw/2 + 0.55, cy + 0.08, ch, va='center', fontsize=8.5, fontweight='bold')
        ax.text(sx - sw/2 + 0.55, cy - 0.18, note, va='center', fontsize=7.5, color='#666',
                style='italic')

    arrow(CX + W/2 + 0.05, y6, sx - sw/2 + 0.05, y6, color='#BDBDBD', lw=1.5, ls='--')

    # ══════════════════════════════════════════════════════════════════
    # 阶段标签
    # ══════════════════════════════════════════════════════════════════
    labels = [
        (0.3, (y0 + y2)/2,           '环境检测阶段',   '#E74C3C'),
        (0.3, (y2 + DH/2 + y4)/2,    '环境感知阶段',   '#E67E22'),
        (0.3, (y4 + y8)/2,           '通信建立阶段',   '#2980B9'),
        (0.3, (y8 + y10)/2,          '注册阶段',       '#27AE60'),
        (0.3, (y10 + y14)/2 - 0.3,   '指令循环阶段',   '#8E44AD'),
    ]
    for lx, ly, lt, lc in labels:
        ax.text(lx, ly, lt, ha='center', va='center', fontsize=10, fontweight='bold',
                color='white', rotation=90,
                bbox=dict(fc=lc, ec=lc, boxstyle='round,pad=0.3', alpha=0.85))

    # ── 标题 ──
    ax.text(CX, 22.9, 'Agent 主程序流程图', ha='center', va='center',
            fontsize=16, fontweight='bold', color='#2C3E50',
            bbox=dict(fc='white', ec='#BDC3C7', boxstyle='round,pad=0.3', lw=1.5))

    fig.tight_layout(rect=[0.02, 0.01, 0.98, 0.99])
    fig.savefig(FLOWCHART_IMG, dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)
    print(f"流程图已保存: {FLOWCHART_IMG} ({os.path.getsize(FLOWCHART_IMG)/1024:.0f} KB)")


# ═══════════════════════════════════════════════════════════════════════
# Part 2: 插入到 Word 文档
# ═══════════════════════════════════════════════════════════════════════

def _replace_para_text(para, new_text):
    """替换段落文本，保留第一个 run 的格式"""
    if not para.runs:
        run = para.add_run(new_text)
        return
    # 保留第一个 run 的格式
    first = para.runs[0]
    from copy import deepcopy
    rPr = deepcopy(first._element.find(qn('w:rPr')))
    # 清除所有 run
    for r in para.runs:
        r._element.getparent().remove(r._element)
    # 添加新 run
    new_run_elem = para._element.makeelement(qn('w:r'), {})
    if rPr is not None:
        new_run_elem.append(rPr)
    t_elem = para._element.makeelement(qn('w:t'), {})
    t_elem.text = new_text
    t_elem.set(qn('xml:space'), 'preserve')
    new_run_elem.append(t_elem)
    # 插入到 pPr 之后（如果有）
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        pPr.addnext(new_run_elem)
    else:
        para._element.insert(0, new_run_elem)


def insert_into_report():
    from docx import Document
    doc = Document(REPORT_PATH)

    # ── 1. 重新编号：七->八, 六->七, 五->六 (逆序避免冲突) ──
    renumber = [
        ("七、 课程设计总结", "八、 课程设计总结"),
        ("六、 编码实现",     "七、 编码实现"),
        ("五、 算法设计",     "六、 算法设计"),
    ]
    for p in doc.paragraphs:
        txt = p.text.strip()
        for old_t, new_t in renumber:
            if txt == old_t:
                _replace_para_text(p, new_t)
                break

    # ── 2. 找到 "六、 算法设计" 的位置 ──
    body = doc.element.body
    all_paras = list(doc.paragraphs)
    insert_before = None
    for p in all_paras:
        if p.text.strip() == "六、 算法设计":
            insert_before = p._element
            break
    if insert_before is None:
        raise RuntimeError("未找到 '六、 算法设计' 段落")

    # 辅助：在指定元素前插入段落（先创建再移动到正确位置）
    def _insert_para(before_elem):
        p = doc.add_paragraph()
        body.insert(list(body).index(before_elem), p._element)
        return p

    # ── 3. 创建五、程序流程图标题 ──
    h1_p = _insert_para(insert_before)
    h1_p.style = doc.styles['Heading 1']
    h1_p.alignment = 0  # LEFT
    run = h1_p.add_run("五、 程序流程图")
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(14)
    run.bold = False

    # ── 4. 图标题 ──
    cap_p = _insert_para(insert_before)
    cap_p.alignment = 1  # CENTER
    cap_p.paragraph_format.space_before = Pt(12)
    cap_p.paragraph_format.space_after = Pt(6)
    run = cap_p.add_run("图5-1  Agent 主程序流程图")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)

    # ── 5. 插入图片 ──
    img_p = _insert_para(insert_before)
    img_p.alignment = 1  # CENTER
    img_p.paragraph_format.space_after = Pt(12)
    run = img_p.add_run()
    run.add_picture(FLOWCHART_IMG, width=Cm(15.5))

    # ── 6. 三段说明文字 ──
    descriptions = [
        (
            "如图5-1所示，Agent 主程序采用严格的 7 阶段生命周期设计。"
            "启动后首先执行沙箱/VM/调试器检测（Phase 1），采用 11 项加权评分算法，"
            "当 risk_ratio >= 0.5 时判定为沙箱环境并伪装退出，避免在分析环境中暴露行为。"
            "通过检测后，执行 WMI 环境感知（Phase 2）采集系统信息、安全软件和白名单应用，"
            "生成 EnvProfile 环境画像。随后生成 RSA-2048 密钥对（Phase 3），"
            "进入通信建立阶段。"
        ),
        (
            "通信建立阶段（Phase 4-5）中，ChannelManager 根据 EnvProfile 按优先级选择通信通道："
            "TCP 加密套接字（基础通路，始终可用）> Outlook COM 邮件（需检测到 Outlook 进程）"
            "> PowerShell 白名单通路（系统自带，低拦截率）> UDP 通道 > Certutil 白名单通路（最后手段）。"
            "connect_best() 方法按优先级依次尝试连接，失败时自动降级到下一通道，最多重试 10 次，"
            "每次间隔 5 秒。连接建立后执行 RSA 公钥互换和 AES 会话密钥协商，"
            "建立 AES-256-GCM 加密通道（Phase 5），随后发送注册信息（Phase 6）并进入指令执行循环（Phase 7）。"
        ),
        (
            "指令执行循环中，Agent 通过 ThreadPoolExecutor（max_workers=4）并发执行指令，"
            "独立线程每 15 秒发送心跳保活。大数据结果（>8KB）自动触发分片传输机制。"
            "当检测到连接断开（心跳超时 45 秒或连接异常）时，"
            "自动调用 _switch_channel() 排除当前失败通道，重新选择备用通道并执行密钥交换和注册，"
            "实现通信链路的自适应容错，保证 Agent 在复杂网络环境下的持续运行能力。"
        ),
    ]
    for desc in descriptions:
        desc_p = _insert_para(insert_before)
        desc_p.paragraph_format.first_line_indent = Cm(0.74)
        desc_p.paragraph_format.line_spacing = 1.25
        desc_p.paragraph_format.space_after = Pt(6)
        run = desc_p.add_run(desc)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)

    # ── 7. 保存 ──
    doc.save(REPORT_PATH)
    print(f"报告已更新: {REPORT_PATH}")


# ═══════════════════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("Step 1: 生成流程图...")
    draw_flowchart()
    print()
    print("Step 2: 插入到课程设计报告...")
    insert_into_report()
    print()
    print("完成!")
