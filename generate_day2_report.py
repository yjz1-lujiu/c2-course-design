# -*- coding: utf-8 -*-
"""生成 网络安全编程课程设计——第二天报告.docx"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── helpers ──────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, font_name, font_name_east_asia=None, size=None, bold=False, color=None):
    """Set run font properties for both ASCII and East Asian."""
    run.font.name = font_name
    run.font.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Set East Asian font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_east_asia or font_name)


def add_heading_styled(doc, text, level=1):
    """Add a heading with SimHei font."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, 'SimHei', 'SimHei', bold=True)
    return h


def add_para(doc, text, font_name='SimSun', size=12, bold=False, indent_first=True,
             space_after=Pt(6), space_before=Pt(0), alignment=None):
    """Add a paragraph with SimSun body text."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    if indent_first:
        pf.first_line_indent = Pt(24)  # ~2 char indent at 12pt
    run = p.add_run(text)
    set_run_font(run, font_name, 'SimSun', size=size, bold=bold)
    return p


def add_bullet(doc, text, font_name='SimSun', size=12, bold=False, level=0):
    """Add a bullet-point paragraph."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    pf = p.paragraph_format
    pf.space_after = Pt(3)
    pf.space_before = Pt(0)
    run = p.add_run(text)
    set_run_font(run, font_name, 'SimSun', size=size, bold=bold)
    return p


def add_code_block(doc, code_text):
    """Add a code block with Consolas 9pt on gray background."""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.first_line_indent = None
        # Gray background on paragraph
        pPr = p._element.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="F0F0F0"/>')
        pPr.append(shd)
        run = p.add_run(line if line else ' ')
        set_run_font(run, 'Consolas', 'Consolas', size=9)


def add_table_with_data(doc, headers, rows, col_widths=None):
    """Add a formatted table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Style: set borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # Header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, 'SimHei', 'SimHei', size=11, bold=True)
        set_cell_shading(cell, 'D9E2F3')

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.cell(r_idx + 1, c_idx)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_text))
            set_run_font(run, 'SimSun', 'SimSun', size=10.5)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    return table


def add_numbered_list(doc, items, font_name='SimSun', size=12):
    """Add numbered items as bullet-style paragraphs."""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(3)
        pf.space_before = Pt(0)
        pf.first_line_indent = Pt(24)
        run = p.add_run(f'{i}. {item}')
        set_run_font(run, font_name, 'SimSun', size=size)


# ── main ─────────────────────────────────────────────────────────────────

def generate_report():
    doc = Document()

    # ── Page setup: A4 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ── Set default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # Also set heading styles to use SimHei
    for i in range(1, 5):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = 'SimHei'
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        hs.font.color.rgb = RGBColor(0, 0, 0)
        if i == 1:
            hs.font.size = Pt(22)
        elif i == 2:
            hs.font.size = Pt(16)
        elif i == 3:
            hs.font.size = Pt(14)
        elif i == 4:
            hs.font.size = Pt(12)

    # =====================================================================
    # TITLE
    # =====================================================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(24)
    title_p.paragraph_format.space_before = Pt(12)
    title_run = title_p.add_run('网络安全编程课程设计')
    set_run_font(title_run, 'SimHei', 'SimHei', size=22, bold=True)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    subtitle_run = subtitle_p.add_run('第二天报告')
    set_run_font(subtitle_run, 'SimHei', 'SimHei', size=18, bold=True)

    # =====================================================================
    # 一、采用的合适的平台
    # =====================================================================
    doc.add_heading('一、采用的合适的平台', level=1)

    # 1.1
    doc.add_heading('1.1 开发语言选择：Python', level=2)
    add_para(doc, '选择理由：')
    add_numbered_list(doc, [
        '丰富的安全工具生态：PyCryptodome 提供 RSA/AES 加密，ctypes 提供 Windows API 调用',
        '跨平台能力：同一套代码可在 Kali Linux（控制端）和 Windows（受控端）运行',
        '快速开发：Python 语法简洁，适合原型开发和课程设计周期',
        '系统级访问：subprocess 模块可调用系统命令，ctypes 可调用底层 API',
        '动态特性：支持 sys.meta_path 动态导入钩子，适合插件化架构',
    ])

    # 1.2
    doc.add_heading('1.2 加密平台：PyCryptodome', level=2)
    add_para(doc, '选择理由：')
    add_numbered_list(doc, [
        '纯 Python 实现，无需编译 C 扩展',
        '支持 RSA-2048（PKCS1_OAEP）、AES-256-GCM 等现代加密算法',
        '提供安全的随机数生成（Crypto.Random）',
        '广泛使用，文档完善，安全可靠',
    ])

    # 1.3
    doc.add_heading('1.3 控制端平台：Kali Linux', level=2)
    add_para(doc, '选择理由：')
    add_numbered_list(doc, [
        'Kali Linux 是专业的渗透测试发行版，预装大量安全工具',
        '适合安全研究人员使用，符合 C2 框架的使用场景',
        'Python 环境完善，可直接运行控制端程序',
    ])

    # 1.4
    doc.add_heading('1.4 受控端平台：Windows 10/11', level=2)
    add_para(doc, '选择理由：')
    add_numbered_list(doc, [
        'Windows 是企业环境中最常见的操作系统，是 C2 框架的主要目标',
        '提供丰富的 API 接口（Win32 API、WMI、COM）',
        '支持 Outlook COM 自动化等隐蔽通信方式',
    ])

    # 1.5
    doc.add_heading('1.5 通信协议设计：自定义 JSON 协议', level=2)
    add_para(doc, '选择理由：')
    add_numbered_list(doc, [
        'JSON 格式可读性强，便于调试',
        'Python 内置 json 库，无需额外依赖',
        '通过 base64 编码支持二进制数据传输',
        '4 字节长度前缀解决 TCP 粘包问题',
    ])

    # 1.6 平台架构总览
    doc.add_heading('1.6 平台架构总览', level=2)
    add_table_with_data(doc,
        headers=['层次', '技术选型', '作用'],
        rows=[
            ['通信层', 'TCP/UDP Socket', '网络数据传输'],
            ['协议层', '自定义 JSON 协议', '消息封装与解析'],
            ['加密层', 'RSA-2048 + AES-256-GCM', '数据加密保护'],
            ['系统层', 'ctypes / WMI / COM', 'Windows 系统接口'],
            ['应用层', 'Python 3.10+', '业务逻辑实现'],
        ],
        col_widths=[3, 5, 6]
    )

    # =====================================================================
    # 二、算法设计
    # =====================================================================
    doc.add_heading('二、算法设计', level=1)

    # 2.1
    doc.add_heading('2.1 混合加密算法设计', level=2)
    doc.add_heading('2.1.1 设计目标', level=3)
    add_para(doc, '解决对称加密密钥分发问题，同时保证通信的机密性、完整性和抗重放性。')

    doc.add_heading('2.1.2 算法描述', level=3)
    add_para(doc, '采用 RSA-2048 + AES-256-GCM 混合加密方案：')

    add_para(doc, '密钥交换阶段（RSA-2048）：', bold=True, indent_first=True)
    add_numbered_list(doc, [
        'Controller 生成 RSA-2048 密钥对，将公钥 PEM 发送给 Agent',
        'Agent 生成 RSA-2048 密钥对，将公钥 PEM 发送给 Controller',
        '双方各自生成 256-bit 随机 AES 会话密钥',
        '用对方的 RSA 公钥（OAEP 填充）加密自己的 AES 会话密钥并发送',
        '双方各自用 RSA 私钥解密得到相同的 AES 会话密钥',
    ])

    add_para(doc, '数据传输阶段（AES-256-GCM）：', bold=True, indent_first=True)
    add_numbered_list(doc, [
        '发送方生成 96-bit 随机 nonce',
        '使用 AES-256-GCM 加密明文数据，得到密文和 128-bit 认证标签',
        '构建消息：header + RSA加密的会话密钥 + nonce + 密文 + tag',
        '序列化为 JSON（二进制字段 base64 编码），添加 4 字节长度前缀',
        '接收方验证时间戳（防重放），解密会话密钥，验证 GCM tag，解密数据',
    ])

    add_para(doc, '消息格式设计：', bold=True, indent_first=True)
    add_code_block(doc, """\
[4字节长度头][JSON消息体]
JSON消息体:
{
    "header": {
        "session_id": "UUID",
        "seq": 序号,
        "timestamp": 时间戳,
        "channel": "tcp/udp/outlook/...",
        "type": "cmd/result/heartbeat/..."
    },
    "encrypted_key": "base64(RSA加密的AES密钥)",
    "nonce": "base64(96-bit随机数)",
    "payload": "base64(AES-GCM密文)",
    "tag": "base64(128-bit认证标签)"
}""")

    # 2.2
    doc.add_heading('2.2 沙箱检测加权评分算法', level=2)
    doc.add_heading('2.2.1 设计目标', level=3)
    add_para(doc, '综合多项指标评估运行环境风险，避免单一指标误判。')

    doc.add_heading('2.2.2 算法描述', level=3)
    add_para(doc, '设检测项集合 C = {c1, c2, ..., c11}，每项对应权重 wᵢ 和检测结果 rᵢ ∈ {0, 1}。')
    add_para(doc, '风险评分计算：', bold=True, indent_first=False)
    add_code_block(doc, """\
risk_score = Σ(rᵢ × wᵢ)  (i=1 到 11)
max_score  = Σ(wᵢ) = 3+3+2+2+2+2+1+1+1+1+1 = 19
risk_ratio = risk_score / max_score""")

    add_para(doc, '检测项与权重设计：', bold=True, indent_first=True)
    add_table_with_data(doc,
        headers=['序号', '检测项', '权重', '检测方法', '权重依据'],
        rows=[
            ['1', 'IsDebuggerPresent', '3', 'kernel32 API', '调试器是最强分析信号'],
            ['2', 'CheckRemoteDebuggerPresent', '3', 'kernel32 API', '远程调试也是强信号'],
            ['3', '调试器进程', '2', 'tasklist 匹配特征进程名', '进程可被伪装'],
            ['4', 'VM 注册表', '2', 'RegOpenKeyExW 检查特征键值', '注册表较为可靠'],
            ['5', 'VM 进程', '2', 'tasklist 匹配 VM 进程名', 'VM 工具进程'],
            ['6', 'VM MAC 前缀', '2', 'getmac 匹配已知前缀', 'MAC 可被修改'],
            ['7', '系统运行时间', '1', 'GetTickCount64 < 10分钟', '短运行时间可疑但不绝对'],
            ['8', '鼠标移动', '1', 'GetCursorInfo 三次采样', '无人操作可疑'],
            ['9', '键盘输入', '1', 'GetAsyncKeyState 全键扫描', '无输入可疑'],
            ['10', '磁盘容量', '1', 'VMware/VB 磁盘 < 60GB', '沙箱磁盘通常较小'],
            ['11', '桌面进程', '1', 'GetForegroundWindow', '无前台窗口可疑'],
        ],
        col_widths=[1.2, 3.5, 1.2, 4.5, 4]
    )

    add_para(doc, '阈值决策：', bold=True, indent_first=True)
    add_bullet(doc, 'risk_ratio >= 0.5（得分 >= 10）→ 沙箱环境，伪装退出')
    add_bullet(doc, '0.25 <= risk_ratio < 0.5（5 <= 得分 < 10）→ 可疑环境')
    add_bullet(doc, 'risk_ratio < 0.25（得分 < 5）→ 正常环境')

    # 2.3
    doc.add_heading('2.3 通信通路选择算法', level=2)
    doc.add_heading('2.3.1 设计目标', level=3)
    add_para(doc, '根据环境感知结果自动选择最优通信通道，失败时自动降级。')

    doc.add_heading('2.3.2 算法描述', level=3)
    add_para(doc, '输入：EnvProfile 环境画像（包含已安装软件列表）')
    add_para(doc, '输出：按优先级排序的可用通道列表', indent_first=False)

    add_para(doc, '算法流程：', bold=True, indent_first=True)
    add_numbered_list(doc, [
        '初始化通道列表 selected = ["tcp"]（TCP 始终可用）',
        '如果 EnvProfile 中检测到 Outlook：selected.append("outlook")',
        'selected.append("powershell")（Windows 自带白名单工具）',
        'selected.append("udp")（备选方案）',
        'selected.append("certutil")（最后手段）',
    ])

    add_para(doc, '连接算法 connect_best(env_profile)：', bold=True, indent_first=True)
    add_code_block(doc, """\
def connect_best(env_profile):
    for ch_type in selected:
        channel = create_channel(ch_type)
        if channel.connect(host, port):
            return channel  # 返回第一个成功的通道
    return None  # 所有通道均失败""")

    # 2.4
    doc.add_heading('2.4 分片传输算法', level=2)
    doc.add_heading('2.4.1 设计目标', level=3)
    add_para(doc, '将大数据包分片传输，规避单包大小检测和流量突发检测。')

    doc.add_heading('2.4.2 算法描述', level=3)
    add_para(doc, '分片算法 Fragmenter：', bold=True, indent_first=True)
    add_para(doc, '输入：原始数据 data，分片大小 chunk_size = 4096', indent_first=False)
    add_para(doc, '输出：分片列表', indent_first=False)
    add_numbered_list(doc, [
        'total_chunks = ceil(len(data) / chunk_size)',
        '生成 transfer_id（8 字符 UUID）',
        'for i in range(total_chunks): 生成分片 fragment = {transfer_id, chunk_index, total_chunks, data, checksum}',
        '每片间隔 0.1 秒发送',
    ])

    add_para(doc, '重组算法 Reassembler：', bold=True, indent_first=True)
    add_para(doc, '输入：分片 fragment', indent_first=False)
    add_para(doc, '输出：完整数据（当所有分片到齐时）', indent_first=False)
    add_numbered_list(doc, [
        '按 transfer_id 查找或创建传输会话',
        '验证 checksum',
        '存储 fragment 到 chunks[chunk_index]',
        'if len(chunks) == total_chunks: 按 index 顺序拼接所有 chunks，返回完整数据；否则返回 None',
    ])

    # 2.5
    doc.add_heading('2.5 内网探测算法', level=2)
    doc.add_heading('2.5.1 设计目标', level=3)
    add_para(doc, '快速发现同网段内存活主机。')

    doc.add_heading('2.5.2 算法描述', level=3)
    add_para(doc, '采用 TCP + ICMP 双探针策略，线程池并发扫描：')
    add_para(doc, '输入：子网前缀 subnet_prefix（如 "192.168.245"），主机范围 [1, 254]', indent_first=False)
    add_para(doc, '输出：存活主机列表', indent_first=False)

    add_code_block(doc, """\
1. 自动检测本机 IP，提取子网前缀
2. 创建信号量 sem = Semaphore(50)    // 限制并发
3. for host_id in range(1, 255):
       创建线程执行 scan_one(host_id)
4. scan_one(host_id):
       sem.acquire()
       ip = f"{subnet_prefix}.{host_id}"
       if tcp_probe(ip, port=445, timeout=0.5s):
           hostname = resolve_hostname(ip)
           记录 {ip, hostname, method="tcp:445"}
       elif ping_host(ip, timeout=0.5s):
           hostname = resolve_hostname(ip)
           记录 {ip, hostname, method="icmp"}
       sem.release()
5. 等待所有线程完成
6. 按 IP 地址排序返回结果""")

    # 2.6
    doc.add_heading('2.6 窗口监控算法', level=2)
    doc.add_heading('2.6.1 设计目标', level=3)
    add_para(doc, '记录用户窗口切换行为。')

    doc.add_heading('2.6.2 算法描述', level=3)
    add_para(doc, '采用定时采样法：')
    add_code_block(doc, """\
1. last_title = ""
2. while 当前时间 < 开始时间 + duration:
       title = GetForegroundWindowTitle()    // user32 API
       if title != last_title:
           记录 {timestamp, title, duration=0}
           if 上一条记录存在:
               上一条记录.duration = 当前时间 - 上一条记录.timestamp
           last_title = title
       sleep(interval)    // 默认 1 秒
3. 返回窗口切换记录列表""")

    # =====================================================================
    # 三、算法分析
    # =====================================================================
    doc.add_heading('三、算法分析', level=1)

    # 3.1
    doc.add_heading('3.1 混合加密算法复杂性分析', level=2)
    add_para(doc, '时间复杂度：', bold=True, indent_first=False)
    add_bullet(doc, 'RSA 密钥生成：O(n³)，n 为密钥位数（2048），耗时约 0.5-1 秒（一次性操作）')
    add_bullet(doc, 'RSA 加密/解密：O(n²)，每次约 1-5ms')
    add_bullet(doc, 'AES-256-GCM 加密/解密：O(m)，m 为数据长度，每次约 0.01-0.1ms（远快于 RSA）')
    add_bullet(doc, '消息序列化（JSON）：O(m)')

    add_para(doc, '空间复杂度：', bold=True, indent_first=False)
    add_bullet(doc, 'RSA 密钥存储：O(n)，约 2KB')
    add_bullet(doc, '消息传输开销：原始数据 + RSA 密文（256 字节）+ nonce（12 字节）+ tag（16 字节）+ JSON 开销')
    add_bullet(doc, '总开销约增加 300-500 字节/消息')

    add_para(doc, '安全性分析：', bold=True, indent_first=False)
    add_bullet(doc, 'RSA-2048：2048 位密钥，等效安全强度 112-bit，抗暴力破解')
    add_bullet(doc, 'AES-256-GCM：256 位密钥 + 96 位 nonce + 128 位认证标签，同时提供加密和完整性保护')
    add_bullet(doc, 'OAEP 填充：抗选择密文攻击（CCA2 安全）')
    add_bullet(doc, '时间戳验证：5 分钟过期窗口，抗重放攻击')

    # 3.2
    doc.add_heading('3.2 沙箱检测算法复杂性分析', level=2)
    add_para(doc, '时间复杂度：O(P)', bold=True, indent_first=False)
    add_bullet(doc, 'P 为进程列表长度，需要遍历进程列表匹配特征进程名')
    add_bullet(doc, 'API 调用（IsDebuggerPresent 等）均为 O(1)')
    add_bullet(doc, '鼠标检测需要 3 次采样，耗时约 1 秒')

    add_para(doc, '空间复杂度：O(P + C)', bold=True, indent_first=False)
    add_bullet(doc, 'P 为进程列表存储空间')
    add_bullet(doc, 'C 为检测结果列表（11 项）')

    add_para(doc, '准确性分析：', bold=True, indent_first=False)
    add_bullet(doc, '加权评分机制降低了单一指标误判的影响')
    add_bullet(doc, '调试器检测（权重 3）准确率高，误报率低')
    add_bullet(doc, '行为检测（权重 1）可能误报，但权重低不影响整体判断')
    add_bullet(doc, '阈值 0.5 意味着至少需要 5 项以上检测为阳性才判定为沙箱')

    # 3.3
    doc.add_heading('3.3 分片传输算法复杂性分析', level=2)
    add_para(doc, '时间复杂度：', bold=True, indent_first=False)
    add_bullet(doc, '分片：O(n/chunk_size)，n 为数据总长度')
    add_bullet(doc, '重组：O(n/chunk_size)（使用字典按 index 存储，拼接为 O(n)）')
    add_bullet(doc, '发送总耗时：O(n/chunk_size × (加密时间 + 间隔时间))')

    add_para(doc, '空间复杂度：', bold=True, indent_first=False)
    add_bullet(doc, '发送端：O(chunk_size)，每次只持有一个分片')
    add_bullet(doc, '接收端：O(n)，需要缓存所有分片直到重组完成')

    add_para(doc, '规避效果分析：', bold=True, indent_first=False)
    add_bullet(doc, '4KB 分片大小接近正常 HTTP 请求大小，不触发大包告警')
    add_bullet(doc, '0.1 秒间隔避免流量突发特征')
    add_bullet(doc, '但总传输时间增加：40KB 数据需要 10 片 × 0.1s = 1 秒额外延迟')

    # 3.4
    doc.add_heading('3.4 内网探测算法复杂性分析', level=2)
    add_para(doc, '时间复杂度：', bold=True, indent_first=False)
    add_bullet(doc, '串行：O(254 × (tcp_timeout + ping_timeout)) = O(254 × 1s) ≈ 254 秒')
    add_bullet(doc, '并发（50 线程）：O(254/50 × 1s) ≈ 5 秒')
    add_bullet(doc, '域名解析：O(k)，k 为存活主机数')

    add_para(doc, '空间复杂度：O(k)，k 为存活主机数', bold=True, indent_first=False)

    add_para(doc, '发现率分析：', bold=True, indent_first=False)
    add_bullet(doc, 'TCP 445 探测：可发现开启 SMB 服务的 Windows 主机')
    add_bullet(doc, 'ICMP 探测：可发现禁用 SMB 但允许 Ping 的主机')
    add_bullet(doc, '双探针策略覆盖率达到 90% 以上（仅防火墙严格过滤的情况可能遗漏）')

    # 3.5 综合评估表
    doc.add_heading('3.5 算法综合评估', level=2)
    add_table_with_data(doc,
        headers=['算法', '时间复杂度', '空间复杂度', '安全性', '实用性'],
        rows=[
            ['混合加密', 'O(n²) RSA + O(m) AES', 'O(n)', '高（112-bit等效）', '高'],
            ['沙箱检测', 'O(P)', 'O(P+C)', '中（加权降低误判）', '高'],
            ['通信选择', 'O(1)', 'O(1)', '高（自动降级）', '高'],
            ['分片传输', 'O(n/chunk)', 'O(n) 接收端', '中（规避检测）', '中'],
            ['内网探测', 'O(254/t) 并发', 'O(k)', '低（主动探测）', '高'],
            ['窗口监控', 'O(duration/interval)', 'O(r) 切换次数', '高（被动采集）', '中'],
        ],
        col_widths=[2.5, 3.5, 2.5, 3, 2.5]
    )

    # ── Save ──
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '第二天报告.docx')
    doc.save(output_path)
    print(f'Report saved to: {output_path}')


if __name__ == '__main__':
    generate_report()
